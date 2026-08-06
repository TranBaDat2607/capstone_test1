#!/usr/bin/env python3
"""
test_evalu_metrics.py — the evaluation layer must MEASURE, never invent.

Why this file exists
--------------------
`evalu/` produced an `evaluation_report.md` whose numbers were not measurements.
Three separate mechanisms manufactured them:

  1. the graph readers used field names no writer in this repo emits
     (`class_name` / `relation` / `source` instead of `class` / `predicate` /
     `subject`), so every graph query returned empty;
  2. an empty query fell through to a hardcoded "benchmark fallback" branch
     (`score: 0.85, aligned_claims: 85, total_claims: 100`) that *looks* like a
     result and is indistinguishable from one in the rendered report;
  3. some metrics never read anything at all (a literal list of "simulated
     cosine similarities"; `precision_at_k = 0.864`).

The failure mode is silent and it points the wrong way — the fabricated numbers
were all FLATTERING (0.85 alignment where the truth is 0.256; a 1.0 exclusion
rate where the truth is 0.098). A capstone report built on them would overstate
the system.

So the assertions here are of two kinds, and both are needed:

  * **agreement** — each metric equals a value this test recomputes for itself,
    straight from the artifact, with its own independent loop (the "oracle"
    pattern the rest of `test/` uses). A metric that reads the wrong field
    cannot agree with an oracle that reads the right one.
  * **non-vacuity** — the oracle values themselves are asserted to be non-empty
    and non-round. Agreement alone is satisfiable by two implementations that
    both return 0; that is exactly the bug being guarded against.

A metric that cannot be measured from artifacts on disk must say so
(`measured=False`) rather than emit a number. `measured=False` with a `reason`
is an honest report; `1.0` with no data behind it is not.

Offline: reads only files already on disk. No LLM, no Neo4j, no network.
Run:  python test/test_evalu_metrics.py
"""

from __future__ import annotations

import json
import re
import sys
from collections import defaultdict
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "evalu"))
sys.path.insert(0, str(REPO_ROOT / "src"))

# Labels below carry Vietnamese number forms ("scaled to triệu"), and the default
# Windows console codepage cannot encode them — without this the test dies inside
# its own print() and reports a crash where the assertion actually passed.
from esg_kg.core.console import ensure_utf8_stdout  # noqa: E402

ensure_utf8_stdout()

RESOLVED = REPO_ROOT / "graph_output" / "resolved" / "resolved_graph.json"
DOSSIERS = REPO_ROOT / "graph_output" / "crosscheck" / "aaa_claim_assessments.json"
SCHEMA = REPO_ROOT / "config" / "schema.json"

failures = []
checks = 0


def check(cond, label, detail=""):
    global checks
    checks += 1
    if cond:
        print(f"  PASS  {label}")
    else:
        print(f"  FAIL  {label}" + (f"\n          {detail}" if detail else ""))
        failures.append(label)


def load(path):
    with open(path, encoding="utf-8") as fh:
        return json.load(fh)


# ---------------------------------------------------------------------------
# Artifact gate. A missing snapshot must SKIP loudly, never pass quietly:
# "all assertions held" over zero artifacts is the vacuous-green failure this
# whole file exists to prevent.
# ---------------------------------------------------------------------------
missing = [p.name for p in (RESOLVED, DOSSIERS, SCHEMA) if not p.exists()]
if missing:
    print(f"SKIP — missing artifacts: {missing}")
    print("Run `python src/esg_kg/core/datasync.py pull` first.")
    sys.exit(2)

graph = load(RESOLVED)
nodes, edges = graph["nodes"], graph["edges"]
dossiers = load(DOSSIERS)
schema = load(SCHEMA)


# ---------------------------------------------------------------------------
# Group 0 — the artifacts really do use these field names.
#
# Pinned because every fabricated number downstream traced back to reading a
# field that does not exist. If a future writer renames `predicate`, this fires
# first and explains the rest.
# ---------------------------------------------------------------------------
print("\n[0] artifact field-name contract")
check(all("class" in n for n in nodes[:100]), "nodes use `class`")
check(not any("class_name" in n for n in nodes[:100]), "nodes do NOT use `class_name`")
check(all("predicate" in e for e in edges[:100]), "edges use `predicate`")
check(all("subject" in e and "object" in e for e in edges[:100]),
      "edges use `subject`/`object` (int indices)")
check(not any("relation" in e or "label" in e for e in edges[:100]),
      "edges do NOT use `relation`/`label`")


# ---------------------------------------------------------------------------
# Group 1 — no fabricated constants survive anywhere in evalu/.
#
# Deliberately a source scan, not a behavioural check: a fallback branch is
# invisible to behaviour right up until the day the artifact is missing, which
# is precisely when someone regenerates the report and believes it.
# ---------------------------------------------------------------------------
print("\n[1] no hardcoded benchmark constants in evalu/")
FABRICATED = [
    ("0.864", "hardcoded Context Precision@5"),
    ("0.912", "hardcoded Context Recall"),
    ("[0.92, 0.88, 0.95, 0.89, 0.94]", "simulated cosine similarities"),
    ("0.8500", "alignment-coverage fallback"),
    ("random_irrelevant = 99", "hardcoded negative control"),
]
sources = {p.name: p.read_text(encoding="utf-8")
           for p in (REPO_ROOT / "evalu").glob("*.py")}
for needle, what in FABRICATED:
    hits = [name for name, text in sources.items() if needle in text]
    check(not hits, f"no {what} ({needle!r})", f"found in {hits}")

# The two modules that fabricated wholesale are gone, not merely patched. Named
# as files rather than matched by a keyword regex: prose describing the removed
# behaviour (this docstring, theirs) would trip a regex forever, and a check that
# cries wolf gets deleted by the next person.
for gone in ("evalu_ragas_metrics.py", "evalu_metamorphic.py"):
    check(not (REPO_ROOT / "evalu" / gone).exists(),
          f"{gone} removed (superseded by evalu_labelfree.py)",
          "it produced numbers from stub adjudicators and literal constants")


# ---------------------------------------------------------------------------
# Oracles — recomputed here, independently, from the artifacts.
# ---------------------------------------------------------------------------
ALIGN_CLASSES = ("SustainabilityClaim", "Goal", "Initiative")

aligned_subjects = {e["subject"] for e in edges if e.get("predicate") == "alignsWithIndicator"}
align_total = sum(1 for n in nodes if n.get("class") in ALIGN_CLASSES)
align_hit = sum(1 for i, n in enumerate(nodes)
                if n.get("class") in ALIGN_CLASSES and i in aligned_subjects)
oracle_alignment = align_hit / align_total

oracle_unverified = sum(1 for d in dossiers
                        if d.get("assessment") == "unverified_insufficient_evidence")
oracle_abstention = oracle_unverified / len(dossiers)

n_flagged = sum(len(d.get("flagged_non_independent_support") or []) for d in dossiers)
n_support = sum(len(d.get("supporting_evidence") or []) for d in dossiers)
oracle_selfverif = n_flagged / (n_flagged + n_support)

zero_penalties = [n for n in nodes
                  if n.get("class") == "Penalty"
                  and n.get("properties", {}).get("amount") == 0]
oracle_zero_tagged = sum(1 for p in zero_penalties
                         if p.get("properties", {}).get("self_reported_zero") is True)

# T1 identity violations read the SCHEMA's identity_keys, not node properties —
# `identity_keys` is a class-level declaration and never appears on a node, so a
# node-level reader silently reports a perfect 0 on every possible input.
TIME_FIELDS = {"valid_from", "valid_to", "is_current", "recorded_at",
               "date", "year", "target_year", "validity_period"}
sys.path.insert(0, str(REPO_ROOT / "src"))
from esg_kg.report.quality import T1_CLASSES  # noqa: E402

schema_classes = schema.get("node_classes", schema.get("classes", {}))
oracle_identity_violations = 0
for cls in T1_CLASSES:
    spec = schema_classes.get(cls) or {}
    if set(spec.get("identity_keys") or []) & TIME_FIELDS:
        oracle_identity_violations += 1

print("\n[oracle] independently recomputed from artifacts")
print(f"         alignment coverage    = {align_hit}/{align_total} = {oracle_alignment:.4f}")
print(f"         abstention rate       = {oracle_unverified}/{len(dossiers)} = {oracle_abstention:.4f}")
print(f"         self-verif exclusion  = {n_flagged}/{n_flagged + n_support} = {oracle_selfverif:.4f}")
print(f"         zero-amount penalties = {len(zero_penalties)} (tagged {oracle_zero_tagged})")
print(f"         T1 identity violations= {oracle_identity_violations}")

print("\n[2] oracles are non-vacuous (agreement on zero proves nothing)")
check(align_total > 1000, f"alignment denominator is real ({align_total})")
check(0 < align_hit < align_total, f"alignment is partial, not 0 or 100% ({align_hit})")
check(len(dossiers) > 1000, f"dossier corpus is real ({len(dossiers)})")
check(n_flagged > 0, f"self-verification guard actually fired ({n_flagged} items)")
check(abs(oracle_alignment - 0.85) > 0.01,
      "true alignment differs from the fabricated 0.85",
      f"oracle={oracle_alignment:.4f} — if this ever equals 0.85 the fallback is back")


# ---------------------------------------------------------------------------
# Group 3 — Tier 1 metrics agree with the oracles.
# ---------------------------------------------------------------------------
print("\n[3] evalu_pipeline_metrics agrees with the oracles")
from evalu_pipeline_metrics import PipelineEvaluator  # noqa: E402

ev = PipelineEvaluator()
res = ev.run_all()
flat = {m["key"]: m for stage in res.values() for m in stage.values()}

REQUIRED_SHAPE = {"key", "metric", "measured", "score", "numerator",
                  "denominator", "source", "note"}
for key, m in flat.items():
    check(REQUIRED_SHAPE <= set(m), f"{key} carries the full result shape",
          f"missing {REQUIRED_SHAPE - set(m)}")

# Scores are rounded to 4 decimals for reporting, so the tolerance is half a
# unit in the last place — tight enough that reading the wrong field cannot pass,
# loose enough that a display choice is not a test failure.
TOL = 5e-5
check(abs(flat["indicator_alignment_coverage"]["score"] - oracle_alignment) < TOL,
      "alignment coverage == oracle",
      f'got {flat["indicator_alignment_coverage"]["score"]}, oracle {oracle_alignment}')
check(abs(flat["evidence_asymmetry_abstention"]["score"] - oracle_abstention) < TOL,
      "abstention rate == oracle",
      f'got {flat["evidence_asymmetry_abstention"]["score"]}, oracle {oracle_abstention}')
check(abs(flat["self_verification_exclusion"]["score"] - oracle_selfverif) < TOL,
      "self-verification exclusion == oracle",
      f'got {flat["self_verification_exclusion"]["score"]}, oracle {oracle_selfverif}')
check(flat["timeless_identity_violation"]["numerator"] == oracle_identity_violations,
      "timeless-identity violations == oracle (schema-level, not node-level)")
check(flat["zero_report_exclusion"]["denominator"] == len(zero_penalties),
      "zero-report denominator == real zero-amount Penalty count",
      f'got {flat["zero_report_exclusion"]["denominator"]}, oracle {len(zero_penalties)}')

# The unmeasurable one must SAY it is unmeasurable. `preserve_property_values`
# runs inside step03 phase 2 and its counter was never persisted, so no artifact
# on disk carries it. A number here could only be invented.
vp = flat["value_preservation_guard"]
check(vp["measured"] is False, "value-preservation guard reports measured=False")
check(vp["score"] is None, "value-preservation guard emits no score",
      f'got {vp["score"]}')
check(bool(vp.get("reason")), "value-preservation guard explains why it is unmeasured")

# Metrics whose truth is already computed by the tested stage implementation
# must be SOURCED from it, not re-derived by a second copy that can drift.
for key in ("temporal_metadata_completeness", "schema_compliance",
            "cluster_conciseness", "timeless_identity_violation"):
    check("quality" in flat[key]["source"],
          f"{key} is sourced from esg_kg.report.quality",
          f'source={flat[key]["source"]!r}')


# ---------------------------------------------------------------------------
# Group 4 — offline label-free metrics (docs/EVALUATION_WITHOUT_LABELS.md).
# ---------------------------------------------------------------------------
print("\n[4] evalu_labelfree — offline, no LLM")
from evalu_labelfree import LabelFreeEvaluator  # noqa: E402

lf = LabelFreeEvaluator()
lf_res = lf.run_all()

# B2 permutation test: shuffling the claim<->evidence mapping must destroy the
# contradiction count. A p-value of exactly 0 is impossible under the (n+1)
# estimator and would mean the null distribution was never built.
b2 = lf_res["b2_permutation_test"]
check(b2["measured"] is True, "B2 permutation test ran")
check(b2["n_permutations"] >= 1000, "B2 used >= 1000 permutations")
check(0 < b2["p_value"] <= 1.0, "B2 p-value is in (0, 1]", f'got {b2["p_value"]}')
check(b2["observed"] == sum(1 for d in dossiers
                            if d.get("assessment") == "appears_contradicted"),
      "B2 observed statistic == real appears_contradicted count")

# Duplicate-claim consistency: the same sentence appearing in two reports must
# get the same verdict. Denominator must be the duplicate GROUPS, not claims.
dup = lf_res["duplicate_claim_consistency"]
by_text = defaultdict(list)
for d in dossiers:
    by_text[re.sub(r"\s+", " ", (d.get("claim_text") or "").strip().lower())].append(d)
oracle_groups = {k: v for k, v in by_text.items() if len(v) > 1 and k}
oracle_incons = sum(1 for v in oracle_groups.values()
                    if len({x.get("assessment") for x in v}) > 1)
check(dup["denominator"] == len(oracle_groups),
      "duplicate-consistency denominator == duplicate groups",
      f'got {dup["denominator"]}, oracle {len(oracle_groups)}')
check(dup["numerator"] == len(oracle_groups) - oracle_incons,
      "duplicate-consistency numerator == consistent groups")
check(len(oracle_groups) > 0, f"there really are duplicate groups ({len(oracle_groups)})")

# B3 structural negative control: retrieval must not pull a DIFFERENT company's
# conduct node into an AAA claim's candidate set.
#
# It cannot run here, and the metric has to SAY so. The ingested corpus is
# single-issuer — every provenance-bearing node traces to an AAA document — so
# there is no second company whose conduct could leak. A "0 leaks" here would be
# guaranteed by the corpus, not earned by the code, and printing it as a pass is
# exactly the flattering-by-construction number this file exists to stop.
b3 = lf_res["b3_structural_negative_control"]
check(b3["measured"] is False, "B3 declares itself unmeasurable on this corpus")
check("issuer" in (b3.get("reason") or "").lower(),
      "B3 names the single-issuer corpus as the reason")
check(b3.get("detail", {}).get("issuer_docs", 0) > 1000,
      "B3 still reports the counts behind that judgement")

# Retrieval yield must come from the dossiers, not from the docs' snapshot.
y = lf_res["retrieval_yield"]
check(y["numerator"] == n_support + n_flagged + sum(
        len(d.get("contradicting_evidence") or []) for d in dossiers),
      "retrieval yield numerator == all retained evidence items")


# ---------------------------------------------------------------------------
# Group 6 — D: anachronism. Evidence dated AFTER the claim it judges.
#
# P8 (docs/TEMPORAL_KG_DESIGN.md §3) says a 2021 report cannot be contradicted
# by information that only surfaced in 2024. This reads that violation straight
# off the verdicts step07 already wrote — no LLM, no perturbation.
# ---------------------------------------------------------------------------
print("\n[6] D — anachronism rate on the existing verdicts")
an = lf_res["d_anachronism"]
check(an["measured"] is True, "D ran")

or_after = {"supports": 0, "contradicts": 0}
or_total = {"supports": 0, "contradicts": 0}
for d in dossiers:
    cy = d.get("year")
    for role, bucket in (("supports", "supporting_evidence"),
                         ("contradicts", "contradicting_evidence")):
        for e in (d.get(bucket) or []):
            ey = e.get("year")
            if isinstance(cy, int) and isinstance(ey, int):
                or_total[role] += 1
                if ey > cy:
                    or_after[role] += 1

for role in ("supports", "contradicts"):
    got = an["by_role"][role]
    check(got["violations"] == or_after[role] and got["comparable"] == or_total[role],
          f"D {role} counts == oracle",
          f'got {got["violations"]}/{got["comparable"]}, oracle {or_after[role]}/{or_total[role]}')

# Split by role deliberately: for a CONTRADICTION, later evidence is a P8
# violation outright; for a SUPPORT it is weaker (a 2016 article may legitimately
# report on 2015). Collapsing them into one rate would overstate the finding.
check(an["by_role"]["contradicts"] is not an["by_role"]["supports"],
      "D reports supports and contradicts separately")
check(or_total["contradicts"] > 0, f'there really are contradictions ({or_total["contradicts"]})')
check(0 < an["by_role"]["contradicts"]["rate"] <= 1.0,
      "D contradiction rate is a real fraction")
# 100% of evidence carries date_uncertain, so the year is often a publish-date
# proxy. The metric must say so or it will be read as harder evidence than it is.
check("date_uncertain" in (an.get("caveat") or ""),
      "D carries the date_uncertain caveat")
check(isinstance(an.get("gap_distribution"), dict) and an["gap_distribution"],
      "D reports the gap distribution, not just a rate")


# ---------------------------------------------------------------------------
# Group 7 — E: time-window ablation.
# ---------------------------------------------------------------------------
print("\n[7] E — time-window ablation")
ab = lf_res["e_time_window_ablation"]
check(ab["measured"] is True, "E ran")
rows = {r["window_after"]: r for r in ab["rows"]}
check(len(rows) >= 4, "E sweeps several window widths")

# The strongest non-vacuity check available: at the window step07 ACTUALLY used,
# the ablation must reproduce the observed evidence count exactly. An ablation
# that cannot reproduce reality at the live setting is not modelling the system.
live = ab["live_window_after"]
check(live in rows, f"E includes the live setting ({live})")
total_ev = sum(len(d.get("supporting_evidence") or []) +
               len(d.get("contradicting_evidence") or []) for d in dossiers)
check(rows[live]["kept_total"] == total_ev,
      "E reproduces the observed evidence count at the live window",
      f'got {rows[live]["kept_total"]}, observed {total_ev}')

# Monotone: a narrower window can never retain more.
widths = sorted(rows)
kept = [rows[w]["kept_total"] for w in widths]
check(all(a <= b for a, b in zip(kept, kept[1:])),
      "E is monotone in the window width", f"kept={kept}")
check(rows[widths[0]]["kept_total"] < rows[live]["kept_total"],
      "E shows a real loss at the narrowest window")


# ---------------------------------------------------------------------------
# Group 8 — A: round-trip grounding of extracted KPI values.
#
# The one metric in this suite that is genuine ACCURACY rather than a proxy: the
# source page is the ground truth for "was this number in the document?", and it
# needs no annotator. Logic is pinned on synthetic pages (Vietnamese number
# forms are where this lives or dies); the real corpus arm only proves the run
# is non-vacuous.
# ---------------------------------------------------------------------------
print("\n[8] A — round-trip grounding")
from evalu_grounding import GroundingEvaluator, number_variants, value_on_page  # noqa: E402

# Vietnamese reports render one quantity many ways. Each of these is a real form
# seen in this corpus; missing any of them turns a correct extraction into a
# false hallucination report, which is worse than not measuring at all.
FORMS = [
    (4500000, "lương bình quân 4,5 triệu đồng/người/tháng", True, "scaled to triệu"),
    (43200, "sản lượng đạt 43.200 tấn", True, "dot as thousands separator"),
    (-843923914, "(Lãi)/lỗ từ hoạt động đầu tư (843.923.914)", True, "accounting negative"),
    (1300000000000, "doanh thu 1.300 tỷ đồng", True, "scaled to tỷ"),
    (99, "hoàn thành 87% kế hoạch năm", False, "absent number must not match"),
]
for value, page, expected, why in FORMS:
    check(value_on_page(value, page) is expected, f"A: {why}",
          f"value={value!r} page={page!r} expected={expected}")

# A 1-digit needle matches almost any page, so such a value must be excluded
# from the denominator outright — not merely fail to match on one example page.
# Asserted on the variant set rather than on value_on_page(), because the earlier
# version passed this by luck: it generated the degenerate needles "00"/"000"
# (from 3/1e3 formatted as "0.00") and simply did not find them on the sample
# page. Those needles would have matched any page carrying a run of zeros, and
# would have counted 3 as verifiable.
check(number_variants(3) == set(),
      "A: a single-digit value yields NO needles (excluded, not guessed)",
      f"got {number_variants(3)}")
check(number_variants(1) == set(), "A: value 1 yields no needles")
check(all(len(v) >= 2 and not v.startswith("00") for v in number_variants(4500000)),
      "A: no degenerate all-zero needles survive",
      f"got {number_variants(4500000)}")

ga = GroundingEvaluator().run()
check(ga["measured"] is True, "A ran on the real corpus")
check(ga["denominator"] > 4000, f'A compared a real population ({ga["denominator"]})')
check(0.0 < ga["score"] < 1.0, "A is neither 0 nor a perfect 1",
      f'got {ga["score"]} — a perfect score here means the matcher is matching anything')
check(ga["numerator"] + ga["mismatches"] == ga["denominator"],
      "A accounts for every compared node")
check(len(ga.get("mismatch_examples") or []) > 0,
      "A shows examples of what failed, so a human can audit the residual")
# Skipped nodes must be visible: a metric that silently drops what it cannot
# check reports a denominator it chose for itself.
check("skipped" in ga and sum(ga["skipped"].values()) > 0,
      "A reports what it could not compare and why")


# ---------------------------------------------------------------------------
# Group 9 — the .docx export carries the SAME content as the Markdown.
#
# The export is a second presentation of numbers a reader may act on, so the
# risk it has to be held against is silent divergence: a converter that drops a
# table, swallows a row, or leaks raw "**" into the prose produces a document
# that looks authoritative and disagrees with its own source. The assertions
# therefore compare against the Markdown rather than checking the file merely
# opens.
# ---------------------------------------------------------------------------
print("\n[9] .docx export matches the Markdown")
# python-docx is an optional, lazily-imported dependency (repo convention), so a
# bare clone may legitimately not have it -> skip. The CONVERTER missing is not
# the same thing and must fail loudly, so its import stays outside the guard:
# folding both into one try/except is what made this group skip silently while
# export_docx.py did not exist at all.
try:
    import docx  # noqa: F401

    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    print("  SKIP — python-docx not installed (pip install python-docx)")

if _HAS_DOCX:
    from export_docx import export  # noqa: E402

    md_path = REPO_ROOT / "evalu" / "evaluation_report.md"
    if not md_path.exists():
        print("  (report not generated yet — skipping)")
    else:
        out = REPO_ROOT / "evalu" / "evaluation_report.docx"
        export(md_path, out)
        check(out.exists() and out.stat().st_size > 5000,
              "export writes a non-trivial .docx")

        d = docx.Document(str(out))
        para_text = "\n".join(p.text for p in d.paragraphs)
        table_cells = [c.text for t in d.tables for r in t.rows for c in r.cells]
        all_text = para_text + "\n" + "\n".join(table_cells)

        md_lines = md_path.read_text(encoding="utf-8").splitlines()

        # Every Markdown table must survive as a real Word table, with every
        # data row. A converter that renders tables as flat text loses the
        # comparison the tables exist to make.
        md_data_rows = sum(1 for ln in md_lines
                           if ln.startswith("|") and not re.match(r"^\|[\s:-]+\|", ln))
        docx_rows = sum(len(t.rows) for t in d.tables)
        check(docx_rows == md_data_rows,
              "every Markdown table row became a Word table row",
              f"docx={docx_rows}, markdown={md_data_rows}")
        check(len(d.tables) >= 8, f"tables survived as tables ({len(d.tables)})")

        # Markdown syntax must be rendered, not printed. A stray "**" tells the
        # reader the document was machine-dumped and undermines every number in it.
        for junk, what in [("**", "bold markers"), ("|---", "table separators"),
                           ("###", "heading hashes")]:
            check(junk not in all_text, f"no raw {what} leaked into the .docx")

        # Headline numbers must actually be present — the cheapest guard against
        # an export that succeeds while emitting an empty shell. Thousands
        # separators are stripped first: the report renders 1001 as "1,001", and
        # matching the raw digits would fail on a document that is perfectly
        # correct.
        digits_text = all_text.replace(",", "")
        for needle in (str(flat["evidence_asymmetry_abstention"]["numerator"]),
                       str(flat["indicator_alignment_coverage"]["numerator"]),
                       str(an["by_role"]["contradicts"]["violations"])):
            check(needle in digits_text, f"headline value {needle} present in .docx")

        # Heading structure is what makes a 250-line report navigable in Word.
        heads = [p for p in d.paragraphs if (p.style.name or "").startswith("Heading")]
        check(len(heads) >= 10, f"headings preserved as Word headings ({len(heads)})")


# ---------------------------------------------------------------------------
# Group 5 — the rendered report may not contain a number nothing measured.
# ---------------------------------------------------------------------------
print("\n[5] rendered report carries only measured numbers")
report_md = REPO_ROOT / "evalu" / "evaluation_report.md"
if report_md.exists():
    text = report_md.read_text(encoding="utf-8")
    for needle, what in [("0.864", "fabricated Context Precision"),
                         ("0.912", "fabricated Context Recall"),
                         ("85/100", "fabricated alignment coverage")]:
        check(needle not in text, f"report has no {what} ({needle})")
    check("measured=False" in text or "KHÔNG ĐO ĐƯỢC" in text or "not measured" in text.lower(),
          "report names its unmeasured metrics instead of hiding them")
else:
    print("  (report not generated yet — skipping)")


print(f"\n{'=' * 60}")
if failures:
    print(f"FAILED — {len(failures)}/{checks} checks failed:")
    for f in failures:
        print(f"  - {f}")
    sys.exit(1)
print(f"OK — {checks}/{checks} checks passed")
