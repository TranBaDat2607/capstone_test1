#!/usr/bin/env python3
"""
export_docx.py — render the evaluation report to Word (.docx).

Converts ``evalu/evaluation_report.md`` rather than re-rendering from the JSON
payload, and that choice is the point. A second renderer reading the same data
would be free to drift from the first, and the two documents would then disagree
about numbers a reader might act on — with no way to tell which one is current.
Converting the Markdown keeps exactly one source of truth for content; this file
only changes how it looks.

The parser therefore only has to handle the constructs `run_evaluation.py`
actually emits — headings, pipe tables, bullets, blockquotes, and inline
`**bold**` / `` `code` ``. It is not a general Markdown implementation and does
not pretend to be; anything it does not recognise is emitted as plain text so
content is never silently dropped.

``test/test_evalu_metrics.py`` group [9] compares the output back against the
Markdown: every table row must survive as a Word table row, and no raw syntax
may leak into the prose.

python-docx is imported lazily and is deliberately NOT in requirements.txt, in
line with the repo's convention for optional tools.

Run:  python evalu/export_docx.py [input.md] [output.docx]
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import List, Optional

EVALU_DIR = Path(__file__).resolve().parent
DEFAULT_MD = EVALU_DIR / "evaluation_report.md"
DEFAULT_DOCX = EVALU_DIR / "evaluation_report.docx"

# `**bold**`, `*italic*` and `` `code` `` — the only inline markup the report
# generator produces. Captured in one pass so the run splitting stays ordered.
_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
_TABLE_SEP_RE = re.compile(r"^\|[\s:\-|]+\|$")


def _add_runs(paragraph, text: str) -> None:
    """Write `text` into `paragraph`, converting inline markup to real runs."""
    for piece in _INLINE_RE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            paragraph.add_run(piece[1:-1]).italic = True
        else:
            paragraph.add_run(piece)


def _split_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _flush_table(doc, rows: List[List[str]]) -> None:
    """Emit collected pipe-table rows as a real Word table.

    Ragged rows are padded rather than rejected: losing a row to a malformed
    line would be a silent content loss, which is the one failure mode this
    converter must not have.
    """
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(width):
            cell = cells[j]
            cell.paragraphs[0].text = ""
            _add_runs(cell.paragraphs[0], row[j] if j < len(row) else "")
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def export(md_path: Path = DEFAULT_MD, docx_path: Path = DEFAULT_DOCX) -> Path:
    """Convert the Markdown report to .docx. Returns the written path."""
    from docx import Document  # lazy: optional dependency
    from docx.shared import Pt

    md_path, docx_path = Path(md_path), Path(docx_path)
    if not md_path.exists():
        raise FileNotFoundError(
            f"{md_path} not found — run `python evalu/run_evaluation.py` first.")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    pending_table: List[List[str]] = []

    def flush() -> None:
        nonlocal pending_table
        _flush_table(doc, pending_table)
        pending_table = []

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()

        # --- table rows accumulate until a non-table line ends the block ---
        if line.startswith("|"):
            if not _TABLE_SEP_RE.match(line):
                pending_table.append(_split_row(line))
            continue
        flush()

        if not line.strip():
            continue

        if line.startswith("---"):
            # Horizontal rules are section breaks in the source; a blank spacer
            # reads better in Word than a literal line of dashes.
            doc.add_paragraph()
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            # Built empty and filled through _add_runs: headings in this report
            # carry inline markup too (a headline figure is written as
            # "### **0.9778** (97.8%) — ..."), and add_heading(text) would drop
            # the text in verbatim, printing the asterisks to the reader.
            heading = doc.add_heading("", 0 if level == 1 else min(level - 1, 4))
            _add_runs(heading, text)
            continue

        if line.startswith(">"):
            para = doc.add_paragraph(style="Intense Quote")
            _add_runs(para, line.lstrip("> ").strip())
            continue

        if line.lstrip().startswith(("- ", "* ")):
            para = doc.add_paragraph(style="List Bullet")
            _add_runs(para, line.lstrip()[2:])
            continue

        if re.match(r"^\d+\.\s", line.lstrip()):
            para = doc.add_paragraph(style="List Number")
            _add_runs(para, re.sub(r"^\d+\.\s", "", line.lstrip()))
            continue

        para = doc.add_paragraph()
        _add_runs(para, line)

    flush()
    doc.save(str(docx_path))
    return docx_path


def main(argv: Optional[List[str]] = None) -> int:
    args = list(argv if argv is not None else sys.argv[1:])
    md = Path(args[0]) if args else DEFAULT_MD
    out = Path(args[1]) if len(args) > 1 else DEFAULT_DOCX
    try:
        written = export(md, out)
    except ImportError:
        print("python-docx chua duoc cai. Chay:  pip install python-docx")
        return 1
    except FileNotFoundError as exc:
        print(str(exc))
        return 1
    print(f"Da xuat: {written}  ({written.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
