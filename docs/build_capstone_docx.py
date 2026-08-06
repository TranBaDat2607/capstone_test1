"""Render the writable portion of the capstone report to .docx.

WHY THIS SCRIPT EXISTS (and is not just a hand-made Word file)

`docs/CAPSTONE_REPORT_DRAFT.md` is a working draft that mixes three kinds of
content: prose that is backed by an artifact in this repo ([REAL]), prose that
cannot be written until a stage is run ([TODO-RUN]), and prose only the team can
write ([WRITE]). Handing over a Word file that silently blends the three is how a
placeholder ends up in a submitted thesis.

So this script writes ONLY what is currently writable, in finished academic
English, and renders every remaining gap as a visible red marker that is
impossible to miss on the page. Regenerating is a one-liner, so the document
never drifts from the numbers in the repo:

    python docs/build_capstone_docx.py

Every figure quoted below was read out of a real artifact and re-verified on
2026-08-03:
  * config/gri_catalog.json                     -> 145 codes, 59/51/35, 20 TT96
  * graph_output/resolved/resolved_graph.json   -> 10,425 nodes / 14,402 edges
  * graph_output/quality/quality_report_*.json  -> the Q1-Q8 ablation tables
  * graph_output/crosscheck/aaa_crosscheck_stats.json
  * gri/benchmark_results/parser_benchmark.json

NO CITATION IS INVENTED. Reference slots are emitted as [CITE: <what is needed>]
markers for the team to fill; a fabricated bibliography is worse than an empty one.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Capstone_Report_v0.1.docx"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
RED = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x59, 0x59, 0x59)

_table_no = [0]
_figure_no = [0]


# --------------------------------------------------------------------------
# document scaffolding
# --------------------------------------------------------------------------
def new_document() -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    # East-Asian font binding, otherwise Word substitutes for Vietnamese glyphs
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, colour in (
        ("Heading 1", 16, RGBColor(0, 0, 0)),
        ("Heading 2", 14, RGBColor(0, 0, 0)),
        ("Heading 3", 12, RGBColor(0, 0, 0)),
        ("Heading 4", 12, RGBColor(0, 0, 0)),
    ):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = colour
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    _add_page_numbers(doc.sections[0])
    return doc


def _add_page_numbers(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for instr in ("begin", "PAGE", "end"):
        el = OxmlElement("w:fldChar" if instr != "PAGE" else "w:instrText")
        if instr == "PAGE":
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), instr)
        run._r.append(el)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)


# --------------------------------------------------------------------------
# content helpers
# --------------------------------------------------------------------------
def h(doc, level: int, text: str):
    return doc.add_heading(text, level=level)


def p(doc, text: str = "", *, italic=False, align=None, space_after=None):
    """Paragraph with lightweight **bold** / *italic* / `code` inline markup."""
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    _emit_runs(para, text, base_italic=italic)
    return para


def _emit_runs(para, text: str, *, base_italic=False, base_bold=False):
    import re

    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[CITE:[^\]]*\])")
    pos = 0
    for m in token.finditer(text):
        if m.start() > pos:
            _run(para, text[pos:m.start()], base_italic, base_bold)
        chunk = m.group(0)
        if chunk.startswith("**"):
            # recurse: bold spans may contain `code` or *italic* inside them
            _emit_runs(para, chunk[2:-2], base_italic=base_italic, base_bold=True)
        elif chunk.startswith("`"):
            r = _run(para, chunk[1:-1], base_italic, base_bold)
            r.font.name = MONO_FONT
            r.font.size = Pt(10.5)
        elif chunk.startswith("[CITE:"):
            r = _run(para, chunk, True, False)
            r.font.color.rgb = RED
        elif chunk.startswith("*"):
            _emit_runs(para, chunk[1:-1], base_italic=True, base_bold=base_bold)
        else:
            _run(para, chunk, base_italic, base_bold)
        pos = m.end()
    if pos < len(text):
        _run(para, text[pos:], base_italic, base_bold)


def _run(para, text, italic, bold):
    r = para.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.name = BODY_FONT
    r.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    return r


def bullet(doc, text: str, *, numbered=False, level=0):
    style = "List Number" if numbered else "List Bullet"
    para = doc.add_paragraph(style=style)
    para.paragraph_format.left_indent = Cm(0.8 + 0.6 * level)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(3)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _emit_runs(para, text)
    return para


def todo(doc, kind: str, text: str):
    """A visible, unmissable gap marker. Red, boxed, prefixed."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    _shade(para, "FFF2F2")
    r = para.add_run(f"[{kind}] ")
    r.bold = True
    r.font.color.rgb = RED
    r.font.name = BODY_FONT
    r2 = para.add_run(text)
    r2.italic = True
    r2.font.color.rgb = RED
    r2.font.name = BODY_FONT
    r2.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    return para


def _shade(para, hex_fill: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def equation(doc, text: str, number: str = ""):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    r = para.add_run(text)
    r.italic = True
    r.font.name = "Cambria Math"
    r.font.size = Pt(12)
    if number:
        tab = para.add_run("\t\t(" + number + ")")
        tab.font.name = BODY_FONT
    return para


def figure_slot(doc, caption: str, source_hint: str):
    _figure_no[0] += 1
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box.paragraph_format.space_before = Pt(8)
    _shade(box, "F2F2F2")
    r = box.add_run(f"[ FIGURE {_figure_no[0]} PLACEHOLDER — {source_hint} ]")
    r.bold = True
    r.font.color.rgb = GREY
    r.font.name = BODY_FONT
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cr = cap.add_run(f"Figure {_figure_no[0]}. {caption}")
    cr.italic = True
    cr.font.size = Pt(11)
    cr.font.name = BODY_FONT
    return _figure_no[0]


def table(doc, caption: str, headers, rows, *, widths=None, right_align_from=1):
    """Caption ABOVE the table (thesis convention), Table Grid, bold header row."""
    _table_no[0] += 1
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(4)
    cr = cap.add_run(f"Table {_table_no[0]}. {caption}")
    cr.bold = True
    cr.font.size = Pt(11)
    cr.font.name = BODY_FONT

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, txt in enumerate(headers):
        _cell(hdr[i], txt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E8E8E8")
    for row in rows:
        cells = t.add_row().cells
        for i, txt in enumerate(row):
            al = WD_ALIGN_PARAGRAPH.RIGHT if i >= right_align_from else WD_ALIGN_PARAGRAPH.LEFT
            _cell(cells[i], str(txt), align=al)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return _table_no[0]


def _cell(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(2)
    _emit_runs(para, text, base_bold=bold)
    for r in para.runs:
        r.font.size = Pt(10.5)
    if fill:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ==========================================================================
# THE REPORT
# ==========================================================================
def build() -> Document:
    doc = new_document()

    # ---------------------------------------------------------------- cover
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Design and Implementation of a Temporal Knowledge-Graph System "
                  "for Evidence-Grounded Greenwashing Analysis of Vietnamese Listed Companies")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = BODY_FONT

    for _ in range(2):
        doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Capstone Project — AIP491\nFPT University, Hanoi")
    sr.font.size = Pt(14)
    sr.font.name = BODY_FONT

    doc.add_paragraph()
    todo(doc, "WRITE", "Group number, full name and student ID of every member, "
                       "supervisor name and title, submission month/year.")

    page_break(doc)

    # ------------------------------------------------------- how to read v0.1
    h(doc, 1, "How to read this draft (remove before submission)")
    p(doc, "This document contains **only the sections that can be written today**. Every "
           "sentence in it describes behaviour that exists in the codebase, and every number "
           "was read out of an artifact in the repository on 3 August 2026 rather than "
           "estimated. Nothing has been rounded up and nothing has been invented.")
    p(doc, "Three kinds of gap remain, and each is marked in red where it occurs so that no "
           "placeholder can survive to submission by accident:")
    bullet(doc, "**[WRITE]** — content only the team can supply (names, opinions, stylistic choices).")
    bullet(doc, "**[TODO-RUN]** — a number that requires running a stage that has not been run yet. "
                "The exact command is given at each marker.")
    bullet(doc, "**[CITE: …]** — a reference slot. **No bibliography entry has been generated**, "
                "because a fabricated citation is a more serious defect in a thesis than a missing "
                "one. The bracket says what kind of source is needed; the team supplies the source.")
    p(doc, "A consolidated list of every outstanding item appears in **Annex Z** at the end of the "
           "document, ordered by how much it would strengthen the defence.")
    p(doc, "Regenerate this file at any time with `python docs/build_capstone_docx.py`; the "
           "numbers are hard-coded from verified artifacts, so the document cannot silently drift "
           "away from the repository.")

    page_break(doc)

    # ------------------------------------------------------------- abstract
    h(doc, 1, "Abstract")
    p(doc, "Environmental, Social and Governance (ESG) disclosure has become a central input to "
           "investment, regulatory and reputational judgement, yet it suffers from a structural "
           "asymmetry: a company both authors and benefits from its own sustainability narrative. "
           "This creates a persistent risk of *greenwashing* — reporting that overstates "
           "environmental or social performance relative to actual conduct. Existing automated ESG "
           "systems largely read the report alone, and therefore inherit exactly the bias they "
           "should be detecting.")
    p(doc, "This thesis presents a temporal knowledge-graph (KG) system that makes greenwashing "
           "analysis falsifiable by construction. The system ingests two deliberately independent "
           "evidence channels into a single graph: the **claim side**, drawn from annual and "
           "sustainability reports published by the company, and the **conduct side**, drawn from "
           "independent Vietnamese news media. Both channels are extracted into one "
           "temporally-versioned schema of 28 node classes and approximately 50 directed edge "
           "labels, under a design in which time lives on edges and event nodes while entity "
           "identity remains timeless, and in which every node retains sentence-level provenance "
           "back to the sentence that produced it.")
    p(doc, "The system comprises five modules: (i) a **dual-standard indicator metadata layer** "
           "that normalises both the Vietnamese regulatory vocabulary (Circular 96/2020/TT-BTC, "
           "Decision 2171/QĐ-BTC, QCVN 09, the SSC–IFC guide) and 145 GRI disclosure codes into a "
           "machine-readable axis with a confirmed crosswalk between them; (ii) a **claim-side "
           "report pipeline** combining layout-aware PDF linearisation, Vietnamese-aware sentence "
           "segmentation and a fine-tuned ViDeBERTa multi-label ESG classifier; (iii) a "
           "**conduct-side news pipeline** built on the same schema but a distinct extraction "
           "prompt; (iv) a **temporal KG construction stage** using structured-output LLM triple "
           "extraction followed by offline schema validation, deterministic entity resolution, "
           "provenance stamping and indicator-axis materialisation; and (v) a **claim–conduct "
           "cross-check stage** that retrieves conduct-side candidates for each sustainability "
           "claim and produces an LLM-adjudicated, provenance-carrying advisory dossier.")
    p(doc, "Because no ground-truth greenwashing labels exist for Vietnamese listed companies, the "
           "system deliberately emits **evidence and an advisory assessment, never a greenwashing "
           "score or verdict**, and is evaluated under a label-free protocol: an eight-attribute "
           "graph quality instrument (Q1–Q8) measured before and after each design change. On a "
           "corpus of 13 annual reports and 30 independent news articles for a listed "
           "plastics-and-construction-materials issuer, the resulting graph contains 10,425 nodes "
           "and 14,402 edges. Three controlled ablations show that the temporal-integrity redesign "
           "reduced schema consistency violations from 1,098 to 1; that materialising the "
           "standard-indicator axis raised the share of masked-answerable queries from 26.3% to "
           "34.8% while cutting graph leaves from 82.2% to 75.8%; and that layout-aware document "
           "linearisation eliminates all 78 cross-standard mis-attributions and all title "
           "truncation incurred by positional text extraction on the GRI corpus. Cross-checking "
           "1,093 extracted claims against the conduct pool produced 3,461 adjudicated "
           "claim–evidence pairs, yielding 70 apparently-supported and 22 apparently-contradicted "
           "claims, with 1,001 claims explicitly returned as *unverified — insufficient evidence* "
           "rather than silently assumed clean.")
    p(doc, "")
    p(doc, "**Keywords:** Greenwashing; ESG disclosure; Temporal Knowledge Graph; GRI Standards; "
           "Circular 96/2020/TT-BTC; Vietnamese NLP; Evidence-grounded reasoning; Large Language "
           "Models.")

    page_break(doc)

    # ---------------------------------------------------------------- TOC
    h(doc, 1, "Table of Contents")
    todo(doc, "WRITE", "In Word: References ▸ Table of Contents ▸ Automatic Table. The heading "
                       "styles in this document are already the built-in Heading 1/2/3, so the "
                       "table will generate correctly with no manual editing. Do the same for "
                       "List of Figures and List of Tables (Insert ▸ Table of Figures, "
                       "caption label 'Figure' / 'Table').")

    page_break(doc)

    # =============================================================== 1
    h(doc, 1, "1. Introduction")

    h(doc, 2, "1.1 Motivation")
    todo(doc, "CITE", "This subsection is written in full, but every citation slot below is "
                      "empty. Insert 1–2 references per paragraph to match the density expected "
                      "of a related-work-grade introduction.")

    p(doc, "Over the past decade, environmental, social and governance (ESG) disclosure has moved "
           "from a peripheral corporate-communications practice to a primary input in investment "
           "allocation, credit assessment and regulatory supervision [CITE: ESG disclosure "
           "literature review or ESG-financial performance meta-analysis]. Asset managers screen "
           "portfolios against sustainability criteria, exchanges impose disclosure obligations as "
           "a condition of listing, and reputational consequences increasingly follow from what a "
           "company does or does not report. The information contained in a sustainability report "
           "is therefore no longer merely narrative; it is priced.")

    p(doc, "In Vietnam, this shift has taken a specific regulatory form that a system trained only "
           "on international frameworks cannot read. Circular 96/2020/TT-BTC on disclosure in the "
           "securities market requires listed companies to report environmental and social impact "
           "in their annual reports; Decision 2171/QĐ-BTC provides the national guidance on "
           "greenhouse-gas reporting; QCVN 09 sets the technical standard for energy-efficient "
           "buildings; and the State Securities Commission, jointly with the IFC, publishes an ESG "
           "reporting guide aimed specifically at Vietnamese issuers [CITE: the four regulatory "
           "instruments, cited as primary sources]. A system that understands only the GRI "
           "Standards will not recognise the vocabulary that Vietnamese companies actually use, and "
           "will therefore fail to locate the disclosures that regulation obliges them to make.")

    p(doc, "This thesis focuses on the construction, building-materials and real-estate sector, and "
           "the choice is methodological rather than one of convenience. The sector is materially "
           "exposed on the environmental dimension — energy consumption, emissions, water, waste, "
           "and recycled-material content are all first-order to its operations — and it is "
           "directly governed by QCVN 09. Consequently its ESG claims are, in principle, "
           "*checkable*: they resolve into physical quantities that leave traces outside the "
           "company's own report. This contrasts with service sectors, where the majority of "
           "disclosure is qualitative and there is correspondingly little to check a claim "
           "against. Selecting a sector whose claims are falsifiable is a precondition for the "
           "problem being studied at all.")

    p(doc, "The core difficulty is an asymmetry between who speaks and who benefits. A company "
           "authors its own sustainability narrative and is also the party that gains from that "
           "narrative being believed. Greenwashing — the practice of presenting environmental or "
           "social performance more favourably than conduct warrants — is the predictable "
           "consequence [CITE: greenwashing definition and measurement]. The methodological "
           "implication is easy to state and easy to overlook: any automated system whose only "
           "input is the corporate report, however sophisticated its language model, inherits "
           "precisely the bias it was built to detect. It can assess whether a report is "
           "internally coherent, well-structured or comprehensive; it cannot assess whether the "
           "report is true.")

    p(doc, "Retrieval-augmented generation (RAG), the dominant architecture for question answering "
           "over corporate documents [CITE: Lewis et al., Retrieval-Augmented Generation], does not "
           "close this gap, because greenwashing is not a retrieval problem. Greenwashing is a "
           "question about a *discrepancy between two sources at two different points in time*. "
           "Detecting it therefore requires three properties that flat retrieval does not provide: "
           "(a) an evidence channel that is independent of the party being assessed; (b) a "
           "representation that carries time explicitly, so that a commitment made in 2012 can be "
           "compared with conduct observed in 2024; and (c) provenance at sentence granularity, so "
           "that a human reader can open the source page and overturn the system's conclusion. "
           "Taken together, these three requirements point directly at a temporal knowledge graph "
           "rather than at a vector index.")

    p(doc, "This thesis therefore proposes a two-channel temporal knowledge graph with a "
           "claim–conduct cross-check layer on top of it. Report-derived claims and news-derived "
           "conduct evidence are extracted into a single schema and a single identity space, but "
           "remain distinguishable at query time; each sustainability claim is then matched against "
           "the conduct evidence available for it, and adjudicated into a provenance-carrying "
           "advisory dossier. The system deliberately stops short of emitting a greenwashing score. "
           "The justification for that decision is developed in §5.1 and defended in §6.2, and it is "
           "the single most consequential design choice in the work.")

    h(doc, 2, "1.2 Objectives")
    p(doc, "The work pursues four objectives:")
    bullet(doc, "To design a **temporal ESG knowledge-graph schema** for Vietnamese listed "
                "companies in which entity identity is timeless, observations and events carry "
                "validity intervals, and every node retains sentence-level provenance to its source "
                "document and page.", numbered=True)
    bullet(doc, "To build a **two-channel ingestion pipeline** that populates this graph from two "
                "structurally independent sources — company-authored reports (the claim side) and "
                "independent Vietnamese news media (the conduct side) — under one schema and one "
                "identity space.", numbered=True)
    bullet(doc, "To align the graph to a **dual-standard indicator axis** covering both the "
                "Vietnamese regulatory ESG vocabulary and the GRI Standards, with an explicit and "
                "confirmed-only crosswalk between them.", numbered=True)
    bullet(doc, "To implement and evaluate an **evidence-grounded claim–conduct cross-check** that "
                "produces auditable, provenance-carrying advisory assessments, and to evaluate the "
                "system as a whole under a **label-free protocol** appropriate to a domain in which "
                "no ground-truth greenwashing labels exist.", numbered=True)

    h(doc, 2, "1.3 Contributions")
    bullet(doc, "A temporal ESG knowledge-graph schema of 28 node classes and approximately 50 edge "
                "labels, governed by eight explicit design principles (P1–P8) and a "
                "machine-checkable invariant set. The P1 rule that entity identity must be timeless, "
                "and the T1/T2/T3 tier partition, are enforced by an offline linter rather than by "
                "convention, so a violation fails a test instead of surviving review.", numbered=True)
    bullet(doc, "An ESG knowledge graph for Vietnamese listed companies that ingests an "
                "*independent conduct channel* alongside the report channel, and keeps the two "
                "distinguishable at query time through a `source_type` stamp carried on every node "
                "and every edge.", numbered=True)
    bullet(doc, "A dual-standard indicator axis unifying Circular 96/2020/TT-BTC and the GRI "
                "Standards through a confirmed-only crosswalk, materialised as graph structure "
                "(`measuredUnder`, `alignsWithIndicator`, `equivalentTo`) rather than held in an "
                "offline lookup table.", numbered=True)
    bullet(doc, "A label-free evaluation instrument (Q1–Q8) that makes design changes measurable in "
                "a domain with no ground truth, demonstrated on three controlled before/after "
                "ablations.", numbered=True)
    bullet(doc, "A documented and defended negative design decision: the system emits evidence and "
                "an advisory assessment but **never a greenwashing score**, because no label exists "
                "against which such a score could be validated.", numbered=True)

    page_break(doc)

    # =============================================================== 2
    h(doc, 1, "2. Related Work")
    todo(doc, "CITE", "The analytical structure of this chapter — in particular the gap statement "
                      "that closes each subsection — is written out in full and can be kept as is. "
                      "The references are NOT written. Each [CITE: …] marker states what kind of "
                      "source belongs there. Target roughly 30–40 references overall to match the "
                      "density of comparable capstone reports.")

    h(doc, 2, "2.1 Automated ESG Assessment from Sustainability Reports")
    p(doc, "Early automated approaches to ESG assessment applied text mining to GRI-aligned "
           "sustainability reports, correlating surface textual properties — length, tone, "
           "boilerplate ratio, readability — with third-party ESG ratings [CITE: text-mining ESG "
           "reports; ESG rating prediction from disclosure text]. More recent work replaces "
           "hand-designed features with large language models, either scoring reports directly or "
           "answering analyst questions over them [CITE: LLM-based ESG assessment framework]. "
           "Across this line of work the unit of output is the *firm*: a score, a rating, or a "
           "summary attributed to the company as a whole. The reader is given no mechanism to "
           "interrogate which specific statement in the report produced the number, and therefore "
           "no mechanism to disagree with it. **This is the first gap: assessment is performed at "
           "firm level, when the object that can actually be verified or falsified is an individual "
           "claim.**")

    h(doc, 2, "2.2 Greenwashing Detection")
    p(doc, "Work aimed specifically at greenwashing falls into three families, each limited by the "
           "evidence it is willing to admit. *Disclosure-only* approaches measure the gap between "
           "how much a firm says and how it is rated by a third party [CITE: greenwashing measured "
           "as disclosure–performance gap]; the limitation is circularity, since the third-party "
           "rating is itself largely derived from disclosure [CITE: ESG rating divergence / "
           "construction]. *Claim-versus-outcome* approaches compare stated commitments against "
           "independent emissions inventories or audit data [CITE: claim vs. verified outcome "
           "study]; this is methodologically the strongest family, but the required independent "
           "data is essentially unavailable in the Vietnamese market. *Linguistic-cue* approaches "
           "train classifiers on vagueness, hedging and unsubstantiated superlatives [CITE: "
           "linguistic markers of greenwashing]; these detect a rhetorical style, not a factual "
           "discrepancy, and a well-written false claim evades them entirely.")
    p(doc, "All three families are constrained by the same underlying condition: **no ground-truth "
           "greenwashing labels exist**. A supervised detector cannot be trained honestly, and a "
           "reported accuracy figure cannot be validated. This observation is the reason the present "
           "work positions its output as *advisory evidence* rather than *classification*, and the "
           "reason its evaluation protocol (§5.1) is label-free by construction. **This is the "
           "second gap, and the one this thesis addresses most directly.**")
    todo(doc, "WRITE", "Recommended addition here, worth roughly half a page: the list in "
                       "docs/EVALUATION_WITHOUT_LABELS.md §8 of the evaluation metrics this group "
                       "tried and discarded, with the reason each failed. Reporting a dead end "
                       "is unusual in a capstone and reads as methodological maturity to an "
                       "examining board.")

    h(doc, 2, "2.3 Knowledge Graphs and Temporal Representation for Corporate Disclosure")
    p(doc, "Knowledge graphs represent entities and typed relations explicitly, which supports "
           "multi-hop reasoning that flat retrieval cannot express [CITE: Hogan et al., Knowledge "
           "Graphs, ACM Computing Surveys 2021]. A recent line of work constructs such graphs "
           "directly from unstructured text by prompting large language models to emit triples "
           "[CITE: Trajanoska et al.; Carta et al., LLM-based KG construction], and this is the "
           "approach adopted here. Separately, the temporal-database literature distinguishes "
           "*valid time*, when a fact holds in the world, from *transaction time*, when the system "
           "recorded it [CITE: bi-temporal data model], and the temporal-KG literature extends "
           "entity representations with validity intervals and version chains [CITE: temporal "
           "knowledge graph representation].")
    p(doc, "These two lines have not been combined for corporate disclosure. Published ESG "
           "knowledge graphs are predominantly atemporal, and an atemporal graph is structurally "
           "incapable of expressing the sentence *the company committed to X in 2021 and acted "
           "contrary to X in 2024* — which is the precise shape of the phenomenon under study. "
           "**This is the third gap: the time dimension is missing from exactly the representation "
           "that needs it most.**")

    h(doc, 2, "2.4 Document Parsing and Vietnamese NLP for ESG Documents")
    p(doc, "Converting a published sustainability standard or annual report into machine-readable "
           "text is not a solved preprocessing step. Positional extractors iterate over text runs "
           "in coordinate order and recover characters without structure [CITE: PyMuPDF]. A newer "
           "generation of neural document-parsing systems — Nougat, Marker/Surya, Docling and "
           "olmOCR — replaces geometric heuristics with learned layout understanding, emitting "
           "Markdown that preserves heading hierarchy, list structure and table cell boundaries "
           "[CITE: Nougat; Marker/Surya; Docling; olmOCR]. §5.4 of this thesis quantifies what that "
           "difference is worth on the GRI corpus specifically, rather than assuming it.")
    p(doc, "On the language side, Vietnamese requires dedicated tooling: `underthesea` provides "
           "word segmentation and sentence splitting adapted to Vietnamese orthography [CITE: "
           "underthesea], while PhoBERT and ViDeBERTa provide pretrained representations [CITE: "
           "PhoBERT; ViDeBERTa]. Diacritics, compound words and the absence of whitespace-delimited "
           "word boundaries mean that an English pipeline applied unchanged degrades badly rather "
           "than gracefully. **This is the fourth gap: no existing work combines layout-aware "
           "parsing, Vietnamese-specific language models and a temporal knowledge graph for ESG.**")

    h(doc, 2, "2.5 Positioning of This Work")
    p(doc, "This thesis sits at the intersection of §2.2 and §2.3, and uses §2.4 as its technical "
           "substrate. It departs from §2.1 in its unit of analysis: the object assessed is **an "
           "individual sustainability claim**, not a firm. It departs from §2.2 in its evidence "
           "base, admitting an independent conduct channel rather than reasoning from disclosure "
           "alone, and in its output, which is an auditable dossier rather than a label. It departs "
           "from §2.3 in applying temporal representation to a task whose central question is "
           "explicitly diachronic. The cost of these choices is that the system cannot report a "
           "detection accuracy; §5.1 argues that in this domain such a figure could not have been "
           "honest in any case.")

    page_break(doc)

    # =============================================================== 3
    h(doc, 1, "3. Methodology")

    h(doc, 2, "3.1 Overview of the System")
    fig_arch = figure_slot(doc, "System architecture: five modules and the two evidence channels.",
                           "adapt from docs/PIPELINE_DIAGRAMS.md fig. 1")
    p(doc, f"As shown in Figure {fig_arch}, the system consists of five modules. The **Indicator "
           "Metadata Module** (§3.2) normalises Vietnamese ESG regulation and the GRI Standards "
           "into a single machine-readable indicator axis. The **Claim-Side Report Processing "
           "Module** (§3.3) converts company-published PDF reports into ESG-labelled, page-anchored "
           "sentences. The **Conduct-Side News Module** (§3.4) builds a structurally independent "
           "evidence channel from Vietnamese news media using the same schema. The **Temporal "
           "Knowledge Graph Module** (§3.5) fuses both channels into one temporally-versioned, "
           "entity-resolved graph. Finally, the **Claim–Conduct Cross-Check Module** (§3.6) "
           "retrieves conduct evidence for each sustainability claim and produces an auditable "
           "advisory dossier.")
    p(doc, "One constraint runs through every module and is treated as a design requirement rather "
           "than a feature: **sentence-level traceability is preserved at every stage**. The triple "
           "`(source_pdf, page, sentence_index)` is attached when a sentence is first extracted and "
           "survives classification, triple extraction, validation, entity resolution and graph "
           "loading. Consequently every statement the system finally makes can be traced back to a "
           "specific sentence on a specific page of a specific document, and a reader who disagrees "
           "with the system can open that page and say so. In a system that makes assertions about "
           "the conduct of real listed companies, this is a safeguard before it is a convenience "
           "(§6.2).")

    # ---- 3.2
    h(doc, 2, "3.2 Indicator Metadata Module")
    p(doc, "The purpose of this module is to turn two heterogeneous bodies of standards — "
           "Vietnamese regulatory texts and the GRI Standards — into **one machine-readable "
           "indicator axis** to which every KPI observation and every sustainability claim is "
           "subsequently anchored.")

    h(doc, 3, "3.2.1 Source corpora")
    p(doc, "Two corpora are processed in parallel through the same shape of pipeline.")
    p(doc, "**(a) Vietnamese regulatory corpus.** Circular 96/2020/TT-BTC on securities-market "
           "disclosure (whose Appendix IV mandates reporting of environmental and social impact), "
           "Decision 2171/QĐ-BTC, QCVN 09:2013/BXD and the SSC–IFC ESG reporting guide. A six-stage "
           "provenance pipeline downloads each instrument, extracts the relevant provisions "
           "**verbatim**, and emits a controlled vocabulary of **35 KPI definitions**. Each "
           "definition carries a `source` block naming the instrument, the article or clause and "
           "the page, so that every indicator can be traced to the line of law that defines it.")
    p(doc, "**(b) GRI corpus.** Forty-two GRI Standards documents — the universal standards GRI 1, "
           "2 and 3, the sector standards GRI 11–14, and the topic-specific 200/300/400 series — "
           "totalling approximately 45 MB of PDF. The source PDFs are version-controlled alongside "
           "the code so that the catalogue can be rebuilt from exactly the documents that produced "
           "it.")

    h(doc, 3, "3.2.2 Layout-aware document linearisation")
    p(doc, "GRI standards are published as visually complex PDFs: two-column layouts, requirement "
           "boxes, sidebars, footnotes and disclosure tables interleaved on the same page. "
           "Classical text extractors that iterate over positional text runs recover characters but "
           "not *structure* — a requirement box merges into body prose, a disclosure table "
           "collapses into an unordered token stream, and the hierarchical relationship between a "
           "disclosure code and its guidance is lost. Because every downstream stage keys on the "
           "disclosure unit, structural loss at this stage propagates irrecoverably.")
    p(doc, "Each standard is therefore linearised with **Marker** [CITE: Paruchuri et al., "
           "marker/Surya], a layout-aware document-conversion pipeline that composes a "
           "layout-detection model, an order-detection model, a table-recognition model and an OCR "
           "model over a text-first extraction backbone, emitting Markdown that preserves heading "
           "hierarchy, list structure and table cell boundaries. Marker belongs to the same "
           "generation of neural document-parsing systems as Nougat, Docling and olmOCR [CITE: as "
           "in §2.4], all of which replace geometric heuristics with learned layout understanding; "
           "Marker was selected for its explicit table reconstruction and for its Markdown target, "
           "which retains precisely the heading hierarchy that the segmenter in §3.2.3 consumes. "
           "Conversion runs through the hosted inference endpoint under a three-worker concurrency "
           "bound with exponential backoff on HTTP 429, and every conversion result is "
           "**content-addressed and cached to disk**, so the corpus is converted exactly once and "
           "every subsequent rebuild of the catalogue is byte-reproducible from the cache without "
           "re-invoking the model.")
    p(doc, "Formally, let S be a GRI standard PDF of P pages. Linearisation produces an ordered "
           "Markdown token sequence")
    equation(doc, "M(S)  =  m₁ ⊕ m₂ ⊕ ⋯ ⊕ m_P", "3.1")
    p(doc, "in which each m_p retains block type (heading level, paragraph, list item, table) and "
           "reading order, and ⊕ denotes order-preserving concatenation. §5.4 reports a controlled "
           "measurement of what this choice is worth relative to positional extraction.")
    todo(doc, "WRITE", "Equations are rendered as plain text here because python-docx cannot emit "
                       "native Word equation objects. In Word, select each equation line and use "
                       "Insert ▸ Equation to convert it; the symbols are already correct.")

    h(doc, 3, "3.2.3 Disclosure segmentation")
    p(doc, "The linearised Markdown is segmented into **disclosure units** by a deterministic "
           "parser operating over the heading hierarchy and the GRI disclosure-code pattern. Because "
           "Markdown preserves heading depth, segmentation is a structural operation rather than a "
           "similarity-based guess: a disclosure boundary is a heading that introduces a disclosure "
           "code, and the unit extends to the next heading of the same or higher level. This yields")
    equation(doc, "D(S)  =  { d₁, d₂, … , d_n }", "3.2")
    p(doc, "where each dᵢ pairs one disclosure code with its requirements and guidance.")

    h(doc, 3, "3.2.4 The ownership rule: a non-obvious correctness constraint")
    p(doc, "GRI standards *re-list disclosures that belong to other standards*. The sector "
           "standards GRI 11–14 and the 2024/25 rewrites GRI 101–103 each reproduce disclosures "
           "owned by topic-specific standards, so a naive first-file-wins attribution assigns "
           "ownership to whichever document happened to be read first. Each disclosure is instead "
           "attributed to the standard whose identifier is a prefix of the disclosure code, a rule "
           "denoted standard_of(·).")
    p(doc, "In this corpus the rule is not cosmetic. Applying it corrected the attribution of **80 "
           "of 136 entries** and repaired **31 mangled titles**. The constraint is pinned by a "
           "regression test rather than left to reviewer vigilance, because it is exactly the class "
           "of defect that produces a plausible-looking catalogue containing systematically wrong "
           "ownership.")

    h(doc, 3, "3.2.5 Catalogue schema and provenance")
    p(doc, "The module emits a catalogue of **145 GRI disclosure codes**, each a record")
    equation(doc, "Cᵢ = ⟨ code, title_en, title_vi, pillar, requirement_type, units, "
                  "tt96_equivalent, versions, provenance ⟩", "3.3")
    p(doc, "with the pillar distribution shown in Table 1 below, and **20 codes carrying a "
           "confirmed Circular-96 equivalent**. Provenance is a `(source_pdf, page, sha256)` "
           "triple: the SHA-256 digest of the source PDF is stored per record, so a claim traced "
           "back to a GRI requirement can be verified against the exact document revision that "
           "produced it. The `versions[]` field records each standard's version year, effective "
           "date and status, which is what allows the graph to state that a 2016 disclosure was "
           "superseded by a 2020 revision — a temporal fact rather than a static label.")
    table(doc, "Pillar distribution of the GRI catalogue (n = 145 disclosure codes).",
          ["Pillar", "Disclosure codes", "Share"],
          [["Governance", "59", "40.7%"],
           ["Environmental", "51", "35.2%"],
           ["Social", "35", "24.1%"],
           ["**Total**", "**145**", "**100%**"]],
          widths=[6.0, 4.0, 3.0])

    h(doc, 3, "3.2.6 The dual-standard crosswalk")
    p(doc, "A hand-curated crosswalk maps Vietnamese regulatory indicators onto GRI codes. Only "
           "rows explicitly marked *confirmed* are materialised as `equivalentTo` edges in the "
           "graph; draft rows are retained for review but excluded unless explicitly enabled by a "
           "command-line flag. This is a deliberate **precision-over-recall** decision: an "
           "incorrect equivalence silently attributes evidence to the wrong regulatory requirement, "
           "which is a materially worse failure than an absent edge, because it is invisible to the "
           "reader.")

    # ---- 3.3
    h(doc, 2, "3.3 Claim-Side Report Processing Module")
    figure_slot(doc, "Claim-side report processing: PDF to ESG-labelled, page-anchored sentences.",
                "adapt from docs/PIPELINE_DIAGRAMS.md fig. 3")
    p(doc, "This module converts published annual reports into a set of ESG-labelled sentences, "
           "each anchored to the page it came from.")

    h(doc, 3, "3.3.1 Corpus acquisition")
    p(doc, "A threaded downloader retrieves annual reports for the company list held in a "
           "spreadsheet of sector issuers. The downloader is **resumable** — an interrupted run "
           "does not re-fetch completed files — and includes an archive-extraction step for the "
           "`.rar` and `.7z` containers in which Vietnamese issuers frequently publish. The raw "
           "corpus currently holds **1,416 PDF files**.")

    h(doc, 3, "3.3.2 Text extraction with page anchoring")
    p(doc, "This channel uses positional text extraction (PyMuPDF), in deliberate contrast to the "
           "layout-aware linearisation used for the standards corpus in §3.2.2. The two channels "
           "have different requirements. What is needed here is an **exact page number for every "
           "sentence**, not the reconstruction of a table; Vietnamese annual reports are "
           "predominantly text-native rather than scanned; and per-page neural conversion of a "
           "1,416-document corpus is not economically viable, whereas the 42-document standards "
           "corpus is converted exactly once and cached forever. The measurement in §5.4 shows that "
           "positional extraction loses *structure*, not *characters* — which is a fatal loss when "
           "the unit of interest is a disclosure heading, and an acceptable one when the unit of "
           "interest is a sentence with a page number. The extractor normalises to Unicode NFC to "
           "preserve Vietnamese diacritics and attaches the page number to every block.")

    h(doc, 3, "3.3.3 Vietnamese-aware sentence segmentation")
    p(doc, "Sentences are segmented with `underthesea` rather than by punctuation rules. Vietnamese "
           "financial reporting is dense with decimal separators written as `1.234,5`, with "
           "abbreviations such as `TNHH`, `CP` and `TP.`, and with full stops inside proper nouns; "
           "a naive full-stop rule fragments these. The output is one JSONL record per sentence, "
           "carrying the traceability triple `(source_pdf, page, sentence_index)` that is preserved "
           "through to the final node in the graph.")

    h(doc, 3, "3.3.4 Multi-label ESG classification with ViDeBERTa")
    p(doc, "Each sentence is classified into the label set {E, S, G, Neutral} by a fine-tuned "
           "**ViDeBERTa-v3** model [CITE: ViDeBERTa]. The task is deliberately formulated as "
           "*multi-label* rather than multi-class, because a single sentence can legitimately carry "
           "more than one dimension — *“The Board of Directors approved the emission-reduction "
           "plan”* is simultaneously a governance statement and an environmental one, and "
           "forcing a single label would discard half of that.")
    todo(doc, "TODO-RUN", "This is the largest remaining gap in the report and an examining board "
                          "will certainly ask about it. Required: model configuration (base "
                          "checkpoint, max sequence length), training setup (epochs, learning rate, "
                          "batch size, optimiser), the train/validation split and its size, and a "
                          "per-label precision / recall / F1 table plus micro- and macro-averages. "
                          "Source: notebooks/kaggle_esg_classify.ipynb (GPU). Until this table "
                          "exists, the classification stage is the one component of the pipeline "
                          "with no quantitative evidence behind it.")
    p(doc, "Classified sentences are written to the labelled corpus and then trimmed into "
           "graph-ready records carrying only the fields the extraction stage consumes.")

    # ---- 3.4
    h(doc, 2, "3.4 Conduct-Side News Module")
    figure_slot(doc, "Conduct-side news ingestion: the independent evidence channel.",
                "adapt from docs/PIPELINE_DIAGRAMS.md fig. 2")

    h(doc, 3, "3.4.1 Design rationale")
    p(doc, "This module is what distinguishes the system from ESG pipelines that read only the "
           "report. If the evidence used to challenge a claim originates from the same document "
           "that made the claim, the system cannot detect greenwashing; it can only verify that a "
           "document written by an interested party is internally consistent. The conduct channel "
           "must therefore be **structurally independent**: a different author, a different "
           "publication incentive and a different editorial process.")

    h(doc, 3, "3.4.2 Pipeline")
    p(doc, "The crawler runs the chain companies → query generation → Google News RSS, Bing and "
           "DuckDuckGo → fetch (disk-cached and rate-limited) → content extraction with "
           "`trafilatura` → normalisation into **the same sentence schema used by the report "
           "channel**. Sharing the schema is intentional: it allows both channels to enter one "
           "identity space while remaining distinguishable through the `source_type` stamp. News "
           "sentences pass through **the same ViDeBERTa classifier** as report sentences, then "
           "through a preprocessing stage that normalises publication dates and removes "
           "boilerplate.")

    h(doc, 3, "3.4.3 The date_uncertain contract")
    p(doc, "A small design decision carries disproportionate weight for honesty. When an article "
           "does not state an explicit date or period for the fact it reports, the system falls "
           "back to the article's publication date as a proxy — but is **required** to mark that "
           "fact `date_uncertain = true`. The flag travels with the evidence all the way into the "
           "final dossier, where it surfaces as an explicit caveat to the reader. The system never "
           "silently assumes a year.")

    h(doc, 3, "3.4.4 Self-verification guard")
    p(doc, "Articles served from domains owned by the company under assessment are not admissible "
           "as independent evidence. The cross-check stage removes `verifiedBy` edges arising from "
           "own-domain sources and records them separately under "
           "`flagged_non_independent_support` — meaning they remain visible to the reader, but do "
           "not count as verification. Without this guard, a company that publishes its own press "
           "releases widely enough could verify its own claims, which would invert the purpose of "
           "the conduct channel.")

    # ---- 3.5
    h(doc, 2, "3.5 Temporal Knowledge Graph Module")

    h(doc, 3, "3.5.1 Schema and the T1/T2/T3 tier model")
    figure_slot(doc, "Temporal KG schema: the T1/T2/T3 tier partition and where time is stored.",
                "adapt from docs/PIPELINE_DIAGRAMS.md fig. 8")
    p(doc, "A single schema file is the source of truth for the graph: approximately **28 node "
           "classes** and **50 directed edge labels**. Classes are partitioned into three tiers, "
           "and each class belongs to exactly one tier — a property enforced by test rather than by "
           "documentation.")
    table(doc, "The T1/T2/T3 tier model and the placement of time.",
          ["Tier", "Meaning", "Example classes", "Where time lives"],
          [["T1", "Durable entities",
            "Organization, Facility, Person, Standard, Location",
            "Not on the node. Identity is timeless; history is held in `temporal_versions` and "
            "`supersedes` edges."],
           ["T2", "Observations and events",
            "KPIObservation, Emission, Waste, Controversy, Penalty, MediaReport",
            "On the node itself (`valid_from` / `valid_to`)."],
           ["T3", "Statements and norms",
            "SustainabilityClaim, Goal, Initiative, StandardIndicator",
            "On the node and on the connecting edges."]],
          widths=[1.6, 3.4, 4.8, 5.2], right_align_from=99)

    h(doc, 3, "3.5.2 The eight design principles")
    p(doc, "Eight principles govern the temporal design. Three carry most of the engineering "
           "consequence and are stated here in full; the complete set appears in Appendix A.")
    bullet(doc, "**P1 — Entity identity is timeless.** No time-valued field may appear in the "
                "identity keys of a T1 class. Violating P1 splits a single company into one entity "
                "per year and destroys every cross-year inference the system exists to perform. The "
                "rule is **automatically linted** by the quality stage, not left to convention. "
                "§5.2 shows what enforcing it was worth.")
    bullet(doc, "**P2 — In the resolved graph, time lives on edges and on T2/T3 nodes**, never on a "
                "T1 entity node.")
    bullet(doc, "**P4 — Dates are canonical ISO.** Every temporal value is normalised to "
                "`YYYY[-MM[-DD]]`, and a version chain containing an open version must have exactly "
                "one member flagged as current.")

    h(doc, 3, "3.5.3 LLM triple extraction")
    figure_slot(doc, "Triple extraction, offline validation and entity resolution.",
                "adapt from docs/PIPELINE_DIAGRAMS.md fig. 5–6")
    p(doc, "For every page containing at least one sentence classified as ESG-relevant, the system "
           "assembles a prompt from the page text, the KPI observations extracted for that page and "
           "the graph schema, and requests **structured output**: typed JSON containing nodes and "
           "edges with temporal metadata attached. Two prompt templates are used, one per channel — "
           "a claim-side template producing SustainabilityClaim, Goal, Initiative and reported "
           "KPIObservation nodes, and a conduct-side template producing Controversy, MediaReport, "
           "Penalty and observed KPIObservation nodes. Every node and edge is stamped with its "
           "`source_type` at this point, and the stamp is never recomputed later.")
    p(doc, "**Language guard.** Both templates require Vietnamese output for the `name`, `title` "
           "and `description` fields. The reason is not stylistic: if the model translates a proper "
           "noun into English on one page and leaves it in Vietnamese on another, entity resolution "
           "will later split one real-world entity into two. The templates are pinned "
           "byte-for-byte by a test, because a prompt that is reworded harmlessly still runs "
           "successfully while changing every extraction it produces.")

    h(doc, 3, "3.5.4 Offline validation and repair")
    p(doc, "Validation runs in three phases, and the ordering is the substantive design decision:")
    bullet(doc, "**Phase 1 (offline).** Reverse the direction of edges emitted backwards and "
                "validate against the schema. An edge label may be legal for several distinct "
                "(source class, target class) pairs; the validator accepts any matching pair and "
                "auto-swaps a reversed one.", numbered=True)
    bullet(doc, "**Phase 1.5 (offline, P4).** Normalise all dates to ISO form, warn where "
                "`valid_from > valid_to`, and default the `date_uncertain` flag on news-side T2 "
                "nodes.", numbered=True)
    bullet(doc, "**Phase 2 (LLM).** Only triples that are *still* invalid after both offline phases "
                "are batched to a language model for repair. This ordering keeps LLM cost "
                "proportional to the number of **errors** rather than to the volume of **data**.",
           numbered=True)
    p(doc, "**Value-preservation guard.** The repair model is permitted to correct the *shape* of a "
           "triple — its class, predicate and temporal fields — but is forbidden from translating, "
           "reformatting, inventing or dropping a property *value*. This is enforced in code rather "
           "than by prompt instruction: after the model responds, property values are restored from "
           "the original. Without the guard, a model prompted in English will helpfully “correct” "
           "a Vietnamese proper noun, and the resulting split entity will not surface until entity "
           "resolution several stages later.")

    h(doc, 3, "3.5.5 Entity resolution")
    figure_slot(doc, "Entity resolution stages A–D.", "adapt from docs/ENTITY_RESOLUTION.md")
    p(doc, "Entity resolution proceeds in four stages:")
    bullet(doc, "**Stage A — deterministic merge** on identity keys, augmented by two *frozen* "
                "anchors: an **issuer anchor** that collapses every name variant of the reporting "
                "company itself, and a **standards anchor** that collapses the four or more "
                "attested spellings of “GRI” and both the Vietnamese and English forms of "
                "Circular 96 onto one canonical node each.")
    bullet(doc, "**Stage B — blocking** on a Vietnamese-normalised signature combined with cosine "
                "similarity over embeddings.")
    bullet(doc, "**Stage C — LLM adjudication** of the residual ambiguous pairs, under an explicit "
                "budget cap.")
    bullet(doc, "**Stage D — consolidation** by disjoint-set union, preserving temporal history "
                "throughout.")
    p(doc, "**The results reported in this thesis were produced with Stages B and C disabled**, "
           "because the API project behind the embedding and adjudication key is billing-blocked. "
           "Only Stage A and the deterministic part of Stage B were active. This is a real "
           "limitation rather than a presentational detail — the residual-duplicate figure in §5.2 "
           "would be lower with the full pipeline enabled — and it is restated in §6.1.")
    p(doc, "The resolved graph contains **10,425 nodes and 14,402 edges**, distributed by class as "
           "shown in Table 3.")
    table(doc, "Node distribution of the resolved graph by class (n = 10,425).",
          ["Class", "Nodes", "Class", "Nodes"],
          [["KPIObservation", "4,906", "Location", "248"],
           ["SustainabilityClaim", "1,217", "Regulation", "220"],
           ["Goal", "722", "Product", "215"],
           ["Initiative", "495", "Standard", "212"],
           ["Organization", "438", "Person", "196"],
           ["Investment", "282", "ClaimKeyword", "141"],
           ["Facility", "277", "Community", "110"],
           ["Project", "255", "MediaReport", "91"]],
          widths=[5.0, 2.4, 5.0, 2.4], right_align_from=1)
    p(doc, "By channel, **10,100 nodes derive from reports, 208 from news, and 117 carry no "
           "channel stamp**. That ratio is the quantitative basis for the thin-conduct-channel "
           "limitation discussed in §5.5 and §6.1, and it should be read alongside every "
           "cross-check result in this thesis.")

    h(doc, 3, "3.5.6 Provenance patching and the indicator axis")
    figure_slot(doc, "Materialisation of the standard-indicator axis.",
                "adapt from docs/STANDARD_INDICATOR_AXIS.md")
    p(doc, "**Provenance patching** stamps `source_doc` and `source_page` — and, for news-derived "
           "nodes, the article title, URL and domain — onto claim and evidence nodes, using a "
           "four-tier precedence: direct parse of the source identifier, exact index lookup, "
           "recomputation of the stable identifier, and finally a page-token match. The stage "
           "operates under a hard invariant: **it never reorders nodes**, because the graph loader "
           "keys Neo4j nodes by array index and the dossiers reference nodes positionally.")
    p(doc, "**Indicator-axis materialisation** appends approximately 35 StandardIndicator nodes and "
           "four kinds of edge: `partOf` (indicator to its parent document), `measuredUnder` "
           "(KPIObservation or Emission to indicator, read from a previously canonicalised "
           "indicator identifier and never guessed at this stage), `equivalentTo` (Circular 96 to "
           "GRI, confirmed crosswalk rows only), and a keyword tier of `alignsWithIndicator` "
           "(Claim, Goal or Initiative to indicator, longest matching phrase winning). The stage is "
           "**append-only** and asserts that property of its own output, so that dossiers computed "
           "against an earlier revision of the graph remain valid.")
    p(doc, "**The self-reported-zero rule.** A Penalty node whose amount is zero represents a "
           "company stating that it *was never fined*. Such a node is flagged `self_reported_zero` "
           "and generates **no conduct edge**. This is the most dangerous shape in the entire "
           "problem: a self-authored assertion being counted as independent evidence of good "
           "conduct. Handling it correctly is a one-line rule whose absence would have quietly "
           "inflated the supported-claim count.")

    h(doc, 3, "3.5.7 Graph materialisation in Neo4j")
    p(doc, "The resolved node and edge collections are loaded into Neo4j as a property graph. Nodes "
           "are keyed by array index, since entities are already resolved and must not be "
           "de-duplicated a second time. Edges are merged on a temporal edge key, so that the same "
           "relation asserted for several different years remains several distinct edges rather "
           "than collapsing into one. Version chains become `supersedes`-linked node sequences for "
           "the classes where superseding is legally meaningful, and are stored as a JSON property "
           "otherwise.")

    # ---- 3.6
    h(doc, 2, "3.6 Claim–Conduct Cross-Check Module")
    figure_slot(doc, "Claim–conduct cross-check: retrieval, adjudication and dossier generation.",
                "adapt from docs/PIPELINE_DIAGRAMS.md fig. 7")

    h(doc, 3, "3.6.1 Formulation")
    p(doc, "For each sustainability claim c, the system retrieves a candidate set of conduct-side "
           "evidence E(c). For each pair (c, e) an adjudicating language model returns one of three "
           "labels,")
    equation(doc, "y(c, e) ∈ { supports,  contradicts,  irrelevant }", "3.4")
    p(doc, "together with a confidence value and a **natural-language rationale**. The rationale is "
           "mandatory rather than optional: it is what a human analyst reads in order to decide "
           "whether to accept the system's judgement, and an assessment that cannot be argued with "
           "cannot be audited.")

    h(doc, 3, "3.6.2 Retrieval")
    p(doc, "Candidates are retrieved within an **asymmetric temporal window**: one year before the "
           "claim, fifty years after it, with the top eight candidates retained. The asymmetry is "
           "deliberate. A claim made in 2012 can be contradicted by conduct observed in 2020, but "
           "essentially cannot be contradicted by conduct observed in 2005; a symmetric window "
           "would spend half of its retrieval budget on evidence that is logically incapable of "
           "bearing on the claim. §5.7 (E1) records the cost of this choice — the window is wide "
           "enough that it should be paired with a distance-decay weighting, which it currently is "
           "not.")

    h(doc, 3, "3.6.3 Adjudication and aggregation")
    p(doc, "LLM adjudication is **mandatory and has no deterministic fallback**. If no provider is "
           "available the stage aborts at start-up rather than silently degrading to a heuristic "
           "that would produce numbers of a different and undocumented kind. The provider used for "
           "the reported results is OpenAI `gpt-4o-mini`; support for the alternative provider was "
           "removed outright from this stage after its API project became permanently unavailable.")
    p(doc, "Pair-level verdicts are aggregated into three advisory assessments: "
           "`appears_supported`, `appears_contradicted` and `unverified_insufficient_evidence`. "
           "Where a dossier contains both supporting and contradicting evidence, **contradiction "
           "takes precedence over support**. This asymmetry is intentional and conservative: the "
           "cost of surfacing a mixed case for human attention is small, whereas the cost of "
           "reporting a claim as supported while contradicting evidence sits in the same dossier is "
           "the exact failure the system exists to prevent.")

    h(doc, 3, "3.6.4 The advisory dossier")
    p(doc, "Each claim yields a record containing its identifier, text and year; the assessment; an "
           "explicit `assessment_is_advisory: true` flag; lists of supporting and contradicting "
           "evidence, where each item carries the node index, class, text, confidence, rationale, "
           "`date_uncertain` flag and independence flag; the separately-recorded "
           "non-independent support; and a list of caveats.")
    p(doc, "One caveat appears in **every** dossier without exception:", italic=False)
    q = doc.add_paragraph(style="Intense Quote")
    q.add_run("No ground-truth greenwashing label exists; this is an advisory opinion.").italic = True

    h(doc, 3, "3.6.5 Presentation layer")
    p(doc, "Two surfaces consume the dossiers. The **claim ledger** is a per-company table ordered "
           "signal-first — contradicted, then supported, then unverified — rendered exclusively "
           "from Neo4j so that what an analyst reads is what the database actually holds. The **ESG "
           "Evidence View** is a three-column web interface aligned to the Circular 96 and GRI "
           "axes, served by a standard-library HTTP server reading live from Neo4j. It displays "
           "only claims that carry an `alignsWithIndicator` edge, which means each card's E/S/G "
           "column is read from the linked indicator's declared pillar rather than inferred from "
           "the claim text.")
    figure_slot(doc, "The ESG Evidence View interface (three-column TT96/GRI layout).",
                "screenshot: run python api/main.py, open http://localhost:8000")

    page_break(doc)

    # =============================================================== 4
    h(doc, 1, "4. Dataset")
    p(doc, "Table 4 summarises the components of the dataset. Every figure is measured from the "
           "artifact named, not estimated.")
    table(doc, "Dataset components.",
          ["Component", "Source", "Size", "Role"],
          [["GRI standards corpus", "GRI (official)", "42 PDFs → 145 disclosure codes",
            "Indicator axis (international)"],
           ["VN regulatory corpus", "TT96/2020, QĐ 2171, QCVN 09, SSC–IFC", "35 KPI definitions",
            "Indicator axis (national)"],
           ["Crosswalk", "Hand-curated, confirmed only", "20 GRI codes with a TT96 equivalent",
            "Dual-standard linking"],
           ["Raw report corpus", "Sector-wide annual reports", "1,416 PDFs downloaded",
            "Ingestion pool"],
           ["Claim-side corpus (analysed)", "AAA annual reports 2011–2025", "13 documents",
            "Claim channel"],
           ["Conduct-side corpus", "Vietnamese news media", "30 articles, 22 distinct domains",
            "Conduct channel"],
           ["Resolved graph", "Output of §3.5", "10,425 nodes / 14,402 edges",
            "Analysis substrate"],
           ["Claim set", "Extracted SustainabilityClaim nodes", "1,093 claims cross-checked",
            "Evaluation unit"]],
          widths=[4.2, 4.2, 4.4, 3.6], right_align_from=99)

    p(doc, "**Depth versus breadth.** The deeply-analysed corpus is concentrated on a **single "
           "issuer**, with a fifteen-year time series, while the raw corpus already contains 1,416 "
           "sector-wide PDFs. This is a conscious trade-off rather than an incomplete run: temporal "
           "depth is a *precondition* for detecting a discrepancy between a claim in one year and "
           "conduct in another, and temporal depth is more expensive to obtain than breadth. The "
           "consequence for the strength of the claims made in this thesis is stated plainly in "
           "§6.1.")
    todo(doc, "TODO-RUN", "Running two or three additional issuers end-to-end before the defence "
                          "would convert this chapter from 'case study' to 'multi-issuer "
                          "evaluation', which is a materially stronger claim. The pipeline already "
                          "supports multiple issuers; no code change is required.")

    p(doc, "**Independence of the news channel.** The 30 conduct-side articles come from 22 "
           "distinct domains, including vietstock.vn, tinnhanhchungkhoan.vn, vneconomy.vn and "
           "thoibaotaichinhvietnam.vn. Domains owned by the issuer itself are marked "
           "**non-independent** and excluded from the verification path in accordance with §3.4.4, "
           "while remaining visible to the reader.")

    page_break(doc)

    # =============================================================== 5
    h(doc, 1, "5. Experiments and Results")

    h(doc, 2, "5.1 Evaluation Design under Absence of Ground Truth")
    p(doc, "**(a) Why there is no ground truth.** No labelled set of greenwashing instances exists "
           "for Vietnamese listed companies. Third-party ESG ratings cannot substitute for such "
           "labels, because those ratings are themselves derived largely from the same corporate "
           "disclosure the system is attempting to test; using them as ground truth would be "
           "circular, and a high agreement score against them would measure conformity to "
           "disclosure rather than fidelity to conduct. The system therefore emits no score, and "
           "the evaluation is displaced onto properties that can be measured without labels.")
    p(doc, "**(b) The instrument: eight graph-quality attributes.** The quality stage measures "
           "eight attributes of the resolved graph entirely offline — no language model and no "
           "database are involved, so the instrument is free to run, deterministic and repeatable. "
           "It is run with a label before and after every design change, and the difference between "
           "the two runs is the experimental result. Table 5 defines the attributes.")
    table(doc, "The Q1–Q8 label-free graph-quality instrument.",
          ["#", "Attribute", "What it measures"],
          [["Q1", "Accuracy", "Non-NFC names; names corrupted during text extraction."],
           ["Q2", "Consistency", "Illegal edges; non-ISO dates; `valid_from > valid_to`; broken "
                                 "current-version chains; versions split by formatting; missing "
                                 "`date_uncertain`; T1 classes carrying time in their identity keys."],
           ["Q3", "Conciseness", "Surplus duplicate T1 entity nodes."],
           ["Q4", "Completeness", "Coverage of the conduct channel (Controversy, Penalty, "
                                  "MediaReport, news-derived KPI)."],
           ["Q5", "Timeliness", "Share of edges and T2 nodes carrying `valid_from`; share of "
                                "news-derived T2 nodes carrying `date_uncertain`."],
           ["Q6", "Provenance", "Share of nodes carrying `source_type`; share of KPI observations "
                                "whose source identifier parses."],
           ["Q7", "Traversability", "Median degree; share of leaf nodes; share of masked-answerable "
                                    "queries; share of claim→conduct paths reachable structurally; "
                                    "share of T2 nodes with degree ≥ 2."],
           ["Q8", "Independence", "Distribution of conduct evidence across channels."]],
          widths=[1.2, 3.0, 10.8], right_align_from=99)
    p(doc, "**(c) Designs not yet executed.** Three further label-free designs have been specified "
           "but not run, and are reported here as future work with a completed design rather than "
           "as results: *metamorphic relations* (paraphrasing a claim must not change its verdict), "
           "a *negative control with permutation testing* (pairing a company's claims against "
           "another company's conduct pool must collapse the contradiction rate towards chance), "
           "and *Krippendorff's α* for agreement across repeated adjudication runs.")
    todo(doc, "TODO-RUN", "The negative control is by far the highest value-per-hour item left in "
                          "the entire evaluation. It is cheap, it requires no new code beyond a "
                          "swapped conduct pool, and it directly answers the sharpest question an "
                          "examiner can ask — 'how do you know the system is not simply labelling "
                          "at random?' A collapse towards zero demonstrates genuine specificity; a "
                          "failure to collapse is itself a finding worth reporting honestly.")

    # 5.2
    h(doc, 2, "5.2 Ablation A — Temporal Integrity Redesign")
    p(doc, "**Setup.** Identical corpus, identical language model; the only variable is the offline "
           "temporal-handling logic — ISO date normalisation, version-chain repair, "
           "`date_uncertain` defaulting, and the removal of time-valued fields from the identity "
           "keys of T1 classes (P1). Both configurations were measured with the same instrument.")
    table(doc, "Effect of the temporal-integrity redesign.",
          ["Metric", "Baseline", "After redesign", "Δ"],
          [["Graph size (nodes / edges)", "10,573 / 13,008", "10,362 / 13,047", "−211 / +39"],
           ["**Q2 total consistency violations**", "**1,098**", "**1**", "**−99.9%**"],
           ["  — broken current-version chains", "660", "0", "−660"],
           ["  — versions split by formatting", "312", "0", "−312"],
           ["  — missing `date_uncertain`", "124", "0", "−124"],
           ["  — T1 classes with time in identity", "2", "0", "−2"],
           ["Q3 surplus duplicate T1 nodes", "271", "60", "−78%"],
           ["Q5 T2 nodes with `valid_from`", "0.0%", "87.7%", "+87.7 pp"],
           ["Q5 news T2 with `date_uncertain`", "0.0%", "100.0%", "+100 pp"],
           ["Q7 leaf nodes", "83.2%", "82.2%", "−1.0 pp"],
           ["Q7 masked-answerable queries", "25.1%", "26.3%", "+1.2 pp"]],
          widths=[7.0, 3.2, 3.4, 2.8])
    p(doc, "**Interpretation.** Three observations are worth drawing out. First, the graph "
           "simultaneously *lost* nodes and *gained* edges (−211 / +39). That combination is the "
           "signature of correct consolidation rather than of data loss: 211 duplicate nodes were "
           "merged, and edges previously scattered across those duplicates now converge on a single "
           "entity. Second, the 78% reduction in surplus duplicate T1 nodes is the direct "
           "consequence of P1 — before the redesign, the same organisation was being split into a "
           "separate entity for each year in which it was mentioned. Third, the figure of 87.7% for "
           "T2 nodes carrying a validity date is deliberately not 100%: the remainder are "
           "observations for which the source document genuinely states no date. The system leaves "
           "the field empty rather than inventing a value, and the figure is reported unrounded for "
           "that reason.")

    # 5.3
    h(doc, 2, "5.3 Ablation B — Standard-Indicator Axis Materialisation")
    p(doc, "**Setup.** The same resolved graph, with the indicator-axis materialisation stage "
           "(§3.5.6) enabled and disabled.")
    table(doc, "Effect of materialising the standard-indicator axis.",
          ["Metric", "Before", "After", "Δ"],
          [["Nodes / edges", "10,362 / 13,047", "10,393 / 13,790", "+31 / +743"],
           ["Q7 leaf nodes", "82.2%", "**75.8%**", "**−6.4 pp**"],
           ["Q7 masked-answerable queries", "26.3%", "**34.8%**", "**+8.5 pp**"],
           ["Q7 T2 nodes with degree ≥ 2", "10.1%", "**19.9%**", "**+9.8 pp**"],
           ["Q7 median degree", "1.0", "1.0", "—"],
           ["Q2 consistency violations", "1", "1", "0"]],
          widths=[7.0, 3.2, 3.4, 2.8])
    p(doc, "**Interpretation.** Thirty-one additional nodes produced 743 additional edges, meaning "
           "each indicator node connects on average about 24 observations or claims. The share of "
           "T2 nodes with degree of at least two nearly doubled: before this stage, most KPI "
           "observations were dangling leaves that could not participate in any reasoning path at "
           "all. This is the mechanism that converts a *bag of observations* into a **traversable "
           "graph**, and it is why the masked-answerable share rises by 8.5 percentage points from "
           "an operation that adds only 0.3% more nodes. That Q2 violations remain at 1 confirms "
           "the stage is genuinely append-only as designed and breaks no existing invariant.")
    table(doc, "Structural anchoring per T2 class after materialisation.",
          ["Class", "Nodes", "Degree ≥ 2"],
          [["Emission", "24", "100.0%"],
           ["Project", "255", "53.3%"],
           ["Investment", "282", "50.4%"],
           ["KPIObservation", "4,906", "16.9%"],
           ["Initiative", "495", "14.7%"],
           ["Waste", "15", "13.3%"],
           ["MediaReport", "91", "9.9%"],
           ["Controversy", "2", "0.0%"],
           ["Penalty", "4", "0.0%"],
           ["ThirdPartyVerification", "24", "0.0%"]],
          widths=[6.0, 3.0, 3.4])
    p(doc, "Three classes sit at 0% anchoring — Controversy, Penalty and ThirdPartyVerification. "
           "This is a known limitation and its cause is the thinness of the conduct channel "
           "documented in §3.5.5, not a defect in the anchoring algorithm: with two Controversy "
           "nodes and four Penalty nodes in the entire graph, there is very little for the "
           "algorithm to anchor.")

    # 5.4 — NEW parser ablation
    h(doc, 2, "5.4 Ablation C — Document Linearisation for the Indicator Axis")
    p(doc, "**Motivation.** §3.2.2 selects a layout-aware parser over positional text extraction. "
           "That is a design assertion, and an assertion of that kind should be measured rather "
           "than declared, particularly since comparable work uses positional extraction for the "
           "same corpus. This ablation measures it.")
    p(doc, "**Setup.** Three arms differ **only** in how the PDF is converted to text: positional "
           "raw extraction; positional extraction with blocks re-sorted into reading order by "
           "coordinate — the method used by comparable work on this corpus; and layout-aware "
           "linearisation as described in §3.2.2. A **single disclosure segmenter, copied verbatim "
           "from the production parser, is applied identically to all three arms**, so the "
           "comparison isolates the parser and cannot be won by tuning a detector. Ground truth is "
           "the disclosure list published by GRI itself in its official content index, which is "
           "derived from no parser in this experiment and therefore cannot favour any arm. Of the "
           "42 standards, 36 appear in the content index and are scored; the remaining six are "
           "excluded rather than scored against an empty denominator. The measurement is fully "
           "offline and reproducible.")
    table(doc, "Parser ablation on 36 GRI standards, scored against the official GRI content index.",
          ["Parser", "Recall", "Precision", "F1", "Title fidelity", "False pos."],
          [["Positional, raw text", "**100.0%**", "62.6%", "77.0%", "81.8%", "82"],
           ["Positional, ordered blocks", "**100.0%**", "62.6%", "77.0%", "81.8%", "82"],
           ["Layout-aware (Marker/Surya)", "98.5%", "**100.0%**", "**99.3%**", "**100.0%**", "**0**"]],
          widths=[5.6, 2.2, 2.4, 1.8, 2.6, 1.8])
    p(doc, "**Positional extraction wins on recall, and that is reported rather than suppressed.** "
           "It recovers every disclosure code in the corpus. The difference between the arms lies "
           "entirely in *precision* and *title fidelity*, and both failures are structural.")
    p(doc, "Positional extraction produces 82 spurious disclosure units. The next table separates "
           "these by kind, because a single total would conceal an important distinction — and because four "
           "of them are an artefact of the ground truth rather than a parser error.")
    table(doc, "False positives by kind.",
          ["Parser", "Total FP", "Ground-truth gap (own standard)", "Real contamination (cross-standard)"],
          [["Positional, raw text", "82", "4", "78"],
           ["Positional, ordered blocks", "82", "4", "78"],
           ["Layout-aware (Marker/Surya)", "0", "0", "0"]],
          widths=[5.0, 2.4, 4.4, 4.6])
    p(doc, "A code whose prefix matches the standard being parsed — for example `306-1` found "
           "inside GRI 306 — is almost certainly a **ground-truth gap**: the 2021 content index "
           "retired disclosures that the 2016 PDF still contains, so the parser is penalised for "
           "extracting genuine content. Four of the 82 are of this kind. The remaining **78 are "
           "genuine cross-standard contamination**: a code such as `2-23` harvested out of a "
           "cross-reference inside GRI 101 and then attributed to GRI 101. This is precisely the "
           "mis-attribution that the ownership rule of §3.2.4 exists to prevent, arriving one stage "
           "earlier and by a different route.")
    p(doc, "Title fidelity fails in a single consistent pattern: truncation at a line break. For "
           "GRI 305-7, positional extraction yields *“Nitrogen oxides (NOx), sulfur oxides "
           "(SOx), and other significant”*, silently dropping *“air emissions”*. That "
           "string becomes the display name of a StandardIndicator node, so the defect is visible "
           "to the end user of the evidence interface and not merely internal.")
    p(doc, "**Is the gap structural or character-level?** A permissive control answers this "
           "directly. Mention recall counts a disclosure code appearing *anywhere* in the extracted "
           "characters, with no structural requirement at all.")
    table(doc, "Control: character-level recall versus structural recall.",
          ["Parser", "Mention recall (characters)", "Unit recall (structure)", "Markdown headings", "Table rows"],
          [["Positional, raw text", "100.0%", "100.0%", "0", "0"],
           ["Positional, ordered blocks", "100.0%", "100.0%", "0", "0"],
           ["Layout-aware (Marker/Surya)", "100.0%", "98.5%", "2,322", "1,418"]],
          widths=[4.6, 3.4, 3.2, 2.6, 2.2])
    p(doc, "Every arm recovers 100% of disclosure codes at character level. Positional extraction "
           "therefore loses no *text* — it loses *structure*, recovering zero headings and zero "
           "table rows against 2,322 and 1,418 respectively. Without heading structure there is no "
           "signal distinguishing a real disclosure heading from a sentence that merely mentions a "
           "disclosure code, which is the direct cause of the 78 contaminating false positives.")
    p(doc, "**A detector limitation, separated from the parser.** The two disclosures missed by the "
           "layout-aware arm are not parser failures. The production segmenter cannot match a "
           "disclosure code wrapped in Markdown emphasis, and only a parser rich enough to emit "
           "emphasis can ever be penalised by that. Re-running with a relaxed pattern applied "
           "identically to all three arms separates the two concerns.")
    table(doc, "Detector ablation: emphasis-tolerant segmentation applied to all arms.",
          ["Parser", "Recall (production)", "Recall (tolerant)", "Precision (tolerant)", "Title fidelity (tolerant)"],
          [["Positional, raw text", "100.0%", "100.0%", "62.6%", "81.8%"],
           ["Positional, ordered blocks", "100.0%", "100.0%", "62.6%", "81.8%"],
           ["Layout-aware (Marker/Surya)", "98.5%", "**100.0%**", "**100.0%**", "**100.0%**"]],
          widths=[4.6, 3.0, 2.8, 3.0, 3.2])
    p(doc, "**Conclusion.** Positional text extraction recovers every disclosure code but cannot "
           "distinguish a disclosure heading from a cross-reference to one, yielding 78 spurious "
           "disclosure units and truncating 18.2% of titles at line breaks. Layout-aware "
           "linearisation resolves both, at 100% precision and 100% title fidelity, and at no "
           "recall cost once the segmenter tolerates Markdown emphasis.")
    todo(doc, "WRITE", "This ablation also surfaced a genuine defect in the production segmenter — "
                       "it silently skips disclosures whose code is emphasised — which was left "
                       "unfixed because repairing it changes a version-controlled artifact "
                       "(config/gri_catalog.json). Decide before submission whether to fix it. "
                       "Either choice is defensible; reporting the defect and the decision is "
                       "stronger than reporting neither.")

    # 5.5
    h(doc, 2, "5.5 Claim–Conduct Cross-Check Results")
    table(doc, "Claim–conduct cross-check results for the analysed issuer.",
          ["Quantity", "Value"],
          [["Sustainability claims extracted", "1,093"],
           ["Conduct pool", "124 (16 MediaReport + 108 KPIObservation)"],
           ["Claims with at least one candidate", "748 (68.4%)"],
           ["Candidate claim–evidence pairs", "3,461"],
           ["Average candidates per claim", "3.17"],
           ["LLM adjudications performed", "3,461 (0 failures)"],
           ["Linking edges written", "152"],
           ["  → `appears_supported`", "70"],
           ["  → `appears_contradicted`", "22"],
           ["  → `unverified_insufficient_evidence`", "1,001"]],
          widths=[9.0, 7.0])
    p(doc, "**On the figure of 1,001.** This number should not be read as a failure rate, and "
           "presenting it as one would misrepresent the system. It is designed behaviour. "
           "Ninety-two per cent of claims are returned as not verifiable because the independent "
           "conduct corpus available for a mid-capitalisation Vietnamese issuer genuinely is thin — "
           "30 articles covering fifteen years of reporting. A system that returned a verdict for "
           "all 1,093 claims from that evidence base would be manufacturing confidence it does not "
           "possess. The honest output is the one the system actually produces, and it carries its "
           "own caveat:")
    q2 = doc.add_paragraph(style="Intense Quote")
    q2.add_run("Thin independent conduct — absence of contradiction is NOT exoneration.").italic = True

    # 5.6
    h(doc, 2, "5.6 Qualitative Case Studies")
    p(doc, "Three dossiers, one per assessment class, drawn unmodified from the system's output.")

    h(doc, 3, "Case 1 — Supported (social / labour)")
    p(doc, "**Claim** (annual report 2012): *“The company implements monthly, quarterly and "
           "annual bonus systems to motivate employees.”* → `appears_supported`, confidence "
           "0.90. **Evidence**: a KPIObservation recording an employee stock-bonus payout of 114.5 "
           "billion VND against 2024 profits. **Comment**: the system connected a policy commitment "
           "made in 2012 to an observed disbursement in 2024 — precisely the cross-year inference "
           "that the temporal representation exists to support, and one that an atemporal graph "
           "could not have expressed. The dossier automatically carries a `date_uncertain` caveat "
           "on the evidence.")

    h(doc, 3, "Case 2 — Contradicted (governance / capital)")
    p(doc, "**Claim** (annual report 2012): *“Actively sought investment sources in order to "
           "use capital from shareholders and investors effectively.”* → "
           "`appears_contradicted`. **Contradicting evidence**: total assets changed by −11.5% as "
           "of 30 June 2025, confidence 0.90. **Supporting evidence also present**: Factory No. 6, "
           "500.6 billion VND, 2016. Because both directions are present, the dossier carries the "
           "caveat *“Evidence is mixed”*, and under the precedence rule of §3.6.3 the "
           "contradiction determines the assessment. **Comment**: this case is also the clearest "
           "instance of error class E1 below — a 2012 claim adjudicated against 2025 evidence with "
           "no discounting for the thirteen-year gap.")

    h(doc, 3, "Case 3 — Unverified")
    p(doc, "**Claim** (annual report 2011): a statement concerning social and charitable activities "
           "and gift-giving. → `unverified_insufficient_evidence`, with no retrieved candidates at "
           "all. **Comment**: this is the archetype of the 1,001 unverified claims — a qualitative "
           "commitment with no measurable KPI attached and no independent reporting. It marks the "
           "system's genuine capability boundary, and it is more informative about that boundary "
           "than either of the two preceding cases.")

    # 5.7
    h(doc, 2, "5.7 Error Analysis")
    p(doc, "Four error classes are identifiable from the current output. Their mechanisms are "
           "established; their relative frequencies are not yet quantified.")
    p(doc, "**E1 — Temporal-distance mismatch.** The adjudicator pairs a 2012 claim with 2025 "
           "evidence (Case 2) without discounting for temporal distance. The retrieval window of "
           "fifty years forward is too wide to be used without a decay function. *Proposed "
           "remedy:* apply a distance-decay weight over the year gap during aggregation.")
    p(doc, "**E2 — Generic-claim over-matching.** Highly generic claims — *“the company strives "
           "for sustainable development”* — match almost any KPI observation, generating both "
           "spurious support and spurious contradiction. *Proposed remedy:* score claim specificity "
           "before retrieval and route claims below a threshold directly to `unverified` rather "
           "than adjudicating them.")
    p(doc, "**E3 — Sparse conduct channel.** The dominant cause of the 1,001 unverified claims. "
           "This is a data-availability limitation and not an algorithmic error. *Proposed remedy:* "
           "extend the conduct channel with administrative-penalty records, State Securities "
           "Commission and stock-exchange announcements, and environmental monitoring data.")
    p(doc, "**E4 — Caveat fatigue from `date_uncertain` propagation.** Most news-side evidence must "
           "use the publication date as a proxy, so the uncertain-date caveat appears on a large "
           "share of dossiers and progressively loses its warning value through familiarity. "
           "*Proposed remedy:* extract sentence-level temporal expressions with named-entity "
           "recognition instead of falling back to the article-level publication date.")
    todo(doc, "TODO-RUN", "Manually reviewing roughly 20 dossiers and estimating the share "
                          "attributable to each of E1–E4 would turn this section from a taxonomy "
                          "into a measurement, and would let it be presented as a distribution "
                          "table. This is a few hours of reading and requires no code.")

    page_break(doc)

    # =============================================================== 6
    h(doc, 1, "6. Discussion, Limitations and Ethics")

    h(doc, 2, "6.1 Limitations")
    bullet(doc, "**The conduct channel is thin.** Only 208 of 10,425 nodes originate from news. "
                "The direct consequence is that 92% of claims cannot be verified. This is stated "
                "plainly rather than reframed, because the conduct channel is the component on "
                "which the entire contribution rests, and its current size bounds every conclusion "
                "in §5.5.", numbered=True)
    bullet(doc, "**One issuer is analysed in depth.** The results constitute a case study with "
                "genuine temporal depth, not evidence generalisable to the sector.", numbered=True)
    bullet(doc, "**There is no ground truth**, so no precision or recall figure for *greenwashing "
                "detection* exists or can exist here. What is reported is graph-quality attributes "
                "and internal consistency. A reader looking for a detection accuracy will not find "
                "one, by design.", numbered=True)
    bullet(doc, "**Entity resolution ran in a degraded configuration.** Stages B and C — embedding "
                "blocking and LLM adjudication — are implemented and covered by tests, but were "
                "disabled for the reported results because of an API billing block. The residual "
                "duplicate count of 60 in §5.2 would be lower had they been enabled, so the figure "
                "should be read as an upper bound rather than as the system's best achievable "
                "result.", numbered=True)
    bullet(doc, "**Dependence on a commercial language model.** Adjudication runs on `gpt-4o-mini`; "
                "results are not bit-for-bit reproducible across model versions, and a future "
                "deprecation would require re-running the paid stage.", numbered=True)
    bullet(doc, "**No systematic human evaluation.** No panel of independent annotators has "
                "re-scored the verdicts, so no inter-annotator agreement figure is available and "
                "the adjudicator's judgements are unaudited.", numbered=True)

    h(doc, 2, "6.2 Ethical Considerations")
    p(doc, "This system makes statements about the conduct of **real, publicly listed companies**, "
           "which places obligations on its design that a purely technical evaluation would not "
           "capture.")
    bullet(doc, "**Advisory, never a verdict.** Every output carries an explicit advisory flag and "
                "the caveat that no ground-truth label exists. The system does not emit a "
                "greenwashing score, and this is a deliberate refusal rather than an unimplemented "
                "feature.")
    bullet(doc, "**Absence of contradiction is not exoneration.** This is hard-coded as a caveat on "
                "every statistical summary the system produces, because it is the most likely "
                "misreading of a result dominated by 1,001 unverified claims.")
    bullet(doc, "**Provenance is a safeguard, not a feature.** Because every statement resolves to "
                "a file, a page and a sentence, an incorrect allegation can always be refuted by a "
                "reader opening the cited page. A system that made unfalsifiable accusations about "
                "named companies would be indefensible regardless of its accuracy.")
    bullet(doc, "**Self-verification is blocked.** Content published by the company itself is never "
                "counted as independent evidence in its own favour.")
    bullet(doc, "**Risk of misuse.** If a reader ignores the caveats and reads "
                "`appears_contradicted` as a finding of fact, the system could cause unwarranted "
                "reputational harm. For that reason the claim ledger presents its coverage caveat "
                "in the header, above the signal-first ordering, rather than in a footnote where it "
                "would be read last or not at all.")

    page_break(doc)

    # =============================================================== 7
    h(doc, 1, "7. Conclusion")
    p(doc, "This thesis addressed the detection of greenwashing in Vietnamese listed companies in "
           "the construction, building-materials and real-estate sector — a setting in which the "
           "party authoring the sustainability narrative is also the party benefiting from it, and "
           "in which systems that read only the corporate report inherit the very bias they are "
           "meant to expose. The approach taken was a two-channel temporal knowledge graph: "
           "report-derived claims and independently-sourced news conduct extracted into a single "
           "schema and a single identity space, kept distinguishable at query time, and joined by a "
           "cross-check layer that adjudicates each claim against the conduct evidence available "
           "for it. The system was realised as five modules spanning dual-standard indicator "
           "metadata, claim-side report processing, conduct-side news ingestion, temporal graph "
           "construction, and claim–conduct cross-checking.")
    p(doc, "The resulting graph contains 10,425 nodes and 14,402 edges built from 13 annual reports "
           "and 30 independent news articles. Three controlled ablations, measured with a "
           "label-free eight-attribute instrument, quantify the design decisions behind it: the "
           "temporal-integrity redesign reduced schema consistency violations from 1,098 to 1 and "
           "surplus duplicate entities by 78%; materialising the standard-indicator axis raised the "
           "share of masked-answerable queries from 26.3% to 34.8% and nearly doubled the share of "
           "observation nodes participating in more than one relation, from an operation adding "
           "only 31 nodes; and layout-aware document linearisation eliminated all 78 cross-standard "
           "mis-attributions and all title truncation that positional extraction incurs on the GRI "
           "corpus. Cross-checking 1,093 extracted claims produced 3,461 adjudicated pairs with no "
           "failures, yielding 70 apparently-supported and 22 apparently-contradicted claims — and "
           "1,001 claims explicitly returned as unverified rather than silently assumed clean.")
    p(doc, "That last figure indicates the direction of future work more clearly than any other "
           "result in this thesis. The binding constraint is not the graph, the schema or the "
           "adjudicator, but the volume of independent conduct evidence: extending the conduct "
           "channel with administrative-penalty records, regulator and exchange announcements, and "
           "environmental monitoring data would convert a large share of those unverified claims "
           "into decidable ones. Beyond that, scaling the analysis from one issuer to the sector, "
           "executing the specified negative-control and metamorphic evaluations, and adding a "
           "path-based reasoning layer over the graph are the natural next steps. The design "
           "principle that should survive all of them is the one adopted here at the outset: in a "
           "domain with no ground truth, a system that shows its evidence and declines to score is "
           "more useful, and more defensible, than one that supplies a number nobody can check.")

    page_break(doc)

    # =============================================================== refs
    h(doc, 1, "References")
    todo(doc, "CITE", "This section is intentionally EMPTY. No reference has been generated, "
                      "because an invented citation is a more serious defect in a thesis than a "
                      "missing one — and fabricated entries are trivially detected by an examiner "
                      "who looks one up. Use the groups below as a collection checklist; each "
                      "[CITE: …] marker in the body names the group it needs. Target 30–40 entries.")
    for group, items in [
        ("ESG and greenwashing", "ESG disclosure literature review; ESG–financial performance "
                                 "meta-analysis; greenwashing definition and measurement; ESG "
                                 "rating divergence."),
        ("Vietnamese regulation (primary sources)", "Circular 96/2020/TT-BTC; Decision "
                                                    "2171/QĐ-BTC; QCVN 09:2013/BXD; SSC–IFC ESG "
                                                    "reporting guide."),
        ("Standards", "GRI Universal Standards 2021; GRI topic-specific standards; GRI content "
                      "index template 2021 (used as ground truth in §5.4)."),
        ("Document parsing", "PyMuPDF; Marker/Surya; olmOCR; Nougat; Docling."),
        ("Vietnamese NLP", "underthesea; PhoBERT; ViDeBERTa."),
        ("Knowledge graphs and temporal representation", "Hogan et al., Knowledge Graphs (ACM "
                                                         "CSUR 2021); bi-temporal data models; "
                                                         "LLM-based KG construction (Trajanoska "
                                                         "et al.; Carta et al.); Neo4j."),
        ("RAG and LLMs", "Lewis et al., Retrieval-Augmented Generation; structured output and "
                         "constrained decoding; LLM-as-a-judge; hallucination in high-stakes "
                         "domains."),
        ("Evaluation methodology", "Metamorphic testing; permutation testing; Krippendorff's α; "
                                   "FEVER."),
    ]:
        bullet(doc, f"**{group}** — {items}")

    page_break(doc)

    # =============================================================== appendix
    h(doc, 1, "Appendix")
    p(doc, "The following appendices are specified but not yet rendered into this document. Each "
           "can be produced mechanically from the artifact named.")
    bullet(doc, "**Appendix A — Graph schema.** The 28 node classes, approximately 50 edge labels "
                "and the T1/T2/T3 tier assignment, together with the full statement of principles "
                "P1–P8. Source: `config/schema.json`.")
    bullet(doc, "**Appendix B — Extraction prompt templates.** Claim-side and conduct-side "
                "templates including the language constraint of §3.5.3.")
    bullet(doc, "**Appendix C — Adjudication prompt and output schema** of the cross-check stage.")
    bullet(doc, "**Appendix D — Indicator metadata schema.** The GRI catalogue record and the KPI "
                "definition record, including the provenance and SHA-256 blocks.")
    bullet(doc, "**Appendix E — Three complete advisory dossiers** corresponding to Cases 1–3 in "
                "§5.6, reproduced verbatim from the system output.")
    bullet(doc, "**Appendix F — Neo4j constraints and analyst queries.**")
    bullet(doc, "**Appendix G — Parser ablation reproduction package.** Benchmark script, test "
                "suite, per-document extractions for all three arms, and side-by-side comparison "
                "tables for all 36 standards. Source: `gri/benchmark_results/`.")

    page_break(doc)

    # =============================================================== annex Z
    h(doc, 1, "Annex Z — Outstanding items (remove this annex before submission)")
    p(doc, "Every gap in this document, collected in one place and ordered by how much closing it "
           "would strengthen the defence.")
    table(doc, "Outstanding items, highest impact first.",
          ["#", "Item", "Where", "Kind", "Effort"],
          [["1", "ViDeBERTa evaluation table: architecture, hyper-parameters, split sizes, "
                 "per-label P/R/F1", "§3.3.4", "TODO-RUN", "GPU run"],
           ["2", "Negative control: AAA claims against another issuer's conduct pool", "§5.1",
            "TODO-RUN", "Low"],
           ["3", "References — the entire bibliography", "All", "CITE", "Medium"],
           ["4", "§1.1 and §2 citation slots (marked inline in red)", "§1.1, §2", "CITE", "Medium"],
           ["5", "E1–E4 frequency estimates from ~20 reviewed dossiers", "§5.7", "TODO-RUN", "Low"],
           ["6", "Two or three additional issuers end-to-end", "§4", "TODO-RUN", "High"],
           ["7", "Figures 1–10: draw or screenshot into the placeholders", "Throughout", "WRITE",
            "Medium"],
           ["8", "Convert plain-text equations to Word equation objects", "§3.2, §3.6", "WRITE",
            "Low"],
           ["9", "Decide whether to fix the emphasis defect in the GRI segmenter", "§5.4", "WRITE",
            "Low"],
           ["10", "Title page: names, student IDs, supervisor, date", "Cover", "WRITE", "Trivial"],
           ["11", "Discarded-metrics discussion from EVALUATION_WITHOUT_LABELS.md §8", "§2.2",
            "WRITE", "Low"],
           ["12", "Generate Table of Contents, List of Figures, List of Tables in Word", "Front",
            "WRITE", "Trivial"]],
          widths=[0.9, 6.6, 2.6, 2.2, 1.8], right_align_from=99)

    p(doc, "**Sections deliberately not drafted here.** §1.1 and §2 are written in full prose but "
           "carry no references; they are complete as argument and incomplete as scholarship. The "
           "bibliography is empty by choice, for the reason stated in the References section. No "
           "number anywhere in this document was estimated, rounded up or carried over from a "
           "previous draft without being re-read from the artifact that produces it.")

    return doc


def main() -> int:
    doc = build()
    doc.save(OUT)
    size_kb = OUT.stat().st_size / 1024
    print(f"wrote {OUT}  ({size_kb:.0f} KB)")
    print(f"  tables: {_table_no[0]}   figure placeholders: {_figure_no[0]}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
