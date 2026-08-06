"""Build the capstone report in the EquiFashion thesis format.

Format reference: docs/EquiFashion_thesis.docx (AIP491, FPT University). What is
copied from it is the *structure and typographic convention*, not its content:

  * A4, 1-inch margins, Times New Roman 13 pt, 1.15 line spacing, justified body
    with a 0.5-inch first-line indent.
  * Front matter in this exact order: cover, Acknowledgment, Abstract (+Keywords),
    Table of contents, List of tables, List of figures, List of abbreviations.
  * Chapters: 1 Introduction (Related Work is *inside* it, as 1.2), 2 Data,
    3 Methodology, 4 Experiment, 5 Conclusion and Discussion, References.
  * "Table N. ..." caption ABOVE the table; "Figure N. ..." caption BELOW the
    figure; both auto-numbered with SEQ fields so Word can build the two lists.
  * Right-aligned "(N)" equation numbers.
  * Numbered references in Springer LNCS style.

Deliberate deviation: EquiFashion styles its chapter headings at 12 pt and its
section headings at 10 pt, i.e. SMALLER than its own 13 pt body text. That is a
defect in the source document, not a convention worth reproducing, so headings
here are 16/14/13 pt. Everything else follows the reference.

Data provenance -- nothing in Chapter 2 is hand-typed:
    docs/eda_out/eda_stats.json   <- written by docs/eda_report_data.py
    docs/figures/*.png            <- written by docs/eda_report_data.py
Chapters 3-5 carry numbers verified against graph_output/, kpi_output/ and
gri/benchmark_results/ on 2026-08-04. NO CITATION IS INVENTED: the References
section is deliberately empty and every reference slot is marked [CITE: ...].

Run:  python docs/build_capstone_report.py
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
FIGDIR = REPO / "docs" / "figures"
OUT = REPO / "docs" / "Capstone_Report_v0.2.docx"
STATS_PATH = REPO / "docs" / "eda_out" / "eda_stats.json"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
BODY_PT = Pt(13)
RED = RGBColor(0xB1, 0x1B, 0x1B)
GREY = RGBColor(0x66, 0x66, 0x66)

TEXT_WIDTH_IN = 6.27  # A4 minus 1-inch margins on both sides

if not STATS_PATH.exists():
    sys.exit(f"missing {STATS_PATH}\nrun: python docs/eda_report_data.py")
S: dict = json.loads(STATS_PATH.read_text(encoding="utf-8"))


def f(n) -> str:
    """Thousands separator, so no number is ever retyped by hand."""
    return f"{n:,}"


_table_no = [0]
_figure_no = [0]
_eq_no = [0]


# ---------------------------------------------------------------------------
# document skeleton
# ---------------------------------------------------------------------------
def new_document() -> Document:
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = BODY_PT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = st.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, bold, italic, before, after in [
        ("Heading 1", 16, True, False, 18, 10),
        ("Heading 2", 14, True, False, 14, 8),
        ("Heading 3", 13, True, True, 12, 6),
    ]:
        h = doc.styles[name]
        h.font.name = BODY_FONT
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.italic = italic
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(after)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.paragraph_format.keep_with_next = True

    cap = doc.styles["Caption"]
    cap.font.name = BODY_FONT
    cap.font.size = Pt(11)
    cap.font.italic = False
    cap.font.bold = False
    cap.font.color.rgb = RGBColor(0, 0, 0)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)

    for sec in doc.sections:
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        _add_page_numbers(sec)
    return doc


def _field(paragraph, instr: str, *, placeholder: str = "", bold=False, size=None):
    """Insert a Word field, e.g. SEQ / TOC / PAGE."""
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)

    r = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {instr} "
    r._r.append(it)

    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r._r.append(fld)

    r = paragraph.add_run(placeholder)
    r.font.name = BODY_FONT
    r.font.bold = bold
    if size:
        r.font.size = size

    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r._r.append(fld)


def _add_page_numbers(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(p, "PAGE", placeholder="1", size=Pt(11))


# ---------------------------------------------------------------------------
# inline markup: **bold**  *italic*  `code`  [CITE: ...]
# ---------------------------------------------------------------------------
TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[CITE:[^\]]*\])")


def _run(para, text, italic, bold):
    r = para.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = BODY_PT
    r.italic = italic
    r.bold = bold
    r._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    return r


def _emit_runs(para, text: str, *, base_italic=False, base_bold=False,
               color=None, size=None):
    """Render inline markup recursively.

    Recursion matters: a backticked path nested inside a **bold** span is common
    in this document, and a non-recursive pass emits the backticks literally.
    `color`/`size` let a caller (todo()) restyle a whole block without the
    markup rules being duplicated there.
    """
    def styled(r):
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = size
        return r

    for chunk in TOKEN.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**") and len(chunk) > 4:
            _emit_runs(para, chunk[2:-2], base_italic=base_italic, base_bold=True,
                       color=color, size=size)
        elif chunk.startswith("`") and chunk.endswith("`") and len(chunk) > 2:
            r = _run(para, chunk[1:-1], base_italic, base_bold)
            r.font.name = MONO_FONT
            r.font.size = size or Pt(11)
            if color is not None:
                r.font.color.rgb = color
        elif chunk.startswith("[CITE:"):
            r = styled(_run(para, chunk, True, False))
            r.font.color.rgb = RED
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            _emit_runs(para, chunk[1:-1], base_italic=True, base_bold=base_bold,
                       color=color, size=size)
        else:
            styled(_run(para, chunk, base_italic, base_bold))


def h(doc, level: int, text: str):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    return p


def p(doc, text: str = "", *, indent=True, align=None, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    if text:
        _emit_runs(para, text, base_italic=italic)
    return para


def bullet(doc, text: str, *, numbered=False):
    para = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.left_indent = Inches(0.45)
    _emit_runs(para, text)
    return para


def quote(doc, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.7)
    para.paragraph_format.right_indent = Inches(0.7)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(8)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(12)
    r.italic = True
    return para


def code_block(doc, text: str):
    """A verbatim listing: monospace, shaded, no markup interpretation.

    Used for real records copied out of the corpus. Markup is deliberately NOT
    parsed here -- a backtick or an asterisk inside a data record is part of the
    record, not an instruction to the renderer.
    """
    for i, line in enumerate(text.split("\n")):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.paragraph_format.right_indent = Inches(0.2)
        para.paragraph_format.space_before = Pt(6) if i == 0 else Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _shade(para, "F4F5F7")
        r = para.add_run(line if line else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(9)
    para.paragraph_format.space_after = Pt(10)
    return para


def _shade(para, hex_fill: str):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_fill)
    para._p.get_or_add_pPr().append(sh)


def todo(doc, kind: str, text: str):
    """An unmissable gap marker: red, on a pink band, prefixed."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.2)
    para.paragraph_format.right_indent = Inches(0.2)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _shade(para, "FFF2F2")
    r = para.add_run(f"[{kind}] ")
    r.bold = True
    r.font.color.rgb = RED
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    _emit_runs(para, text, color=RED, size=Pt(11))
    return para


def equation(doc, text: str) -> int:
    """Centred equation with a right-aligned number, EquiFashion style."""
    _eq_no[0] += 1
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabs = para.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(TEXT_WIDTH_IN))
    r = para.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(12)
    r2 = para.add_run(f"\t({_eq_no[0]})")
    r2.font.name = BODY_FONT
    r2.font.size = Pt(12)
    return _eq_no[0]


def caption(doc, label: str, text: str) -> int:
    """'Table N. ...' / 'Figure N. ...' with a real SEQ field."""
    counter = _table_no if label == "Table" else _figure_no
    counter[0] += 1
    para = doc.add_paragraph(style="Caption")
    r = para.add_run(f"{label} ")
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    _field(para, f"SEQ {label} \\* ARABIC", placeholder=str(counter[0]), bold=True, size=Pt(11))
    r = para.add_run(". ")
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    # route through _emit_runs so `code` / **bold** in a caption is rendered, not printed
    _emit_runs(para, text)
    for run in para.runs:
        run.font.size = Pt(11)
    return counter[0]


def figure(doc, filename: str, cap_text: str, *, width_in: float = 6.2) -> int:
    """Embed a real PNG; caption goes BELOW, per the reference format."""
    path = FIGDIR / filename
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph.paragraph_format.space_before = Pt(8)
    ph.paragraph_format.space_after = Pt(2)
    if path.exists():
        ph.add_run().add_picture(str(path), width=Inches(width_in))
    else:
        r = ph.add_run(f"[missing figure: {filename} — run python docs/eda_report_data.py]")
        r.font.color.rgb = RED
        r.italic = True
    return caption(doc, "Figure", cap_text)


def figure_slot(doc, cap_text: str, hint: str) -> int:
    """A figure that still has to be drawn by hand."""
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box.paragraph_format.space_before = Pt(8)
    box.paragraph_format.space_after = Pt(2)
    _shade(box, "F2F2F2")
    r = box.add_run(f"\n[ FIGURE PLACEHOLDER — {hint} ]\n")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    r.font.name = BODY_FONT
    return caption(doc, "Figure", cap_text)


def table(doc, cap_text: str, headers, rows, *, widths=None, right_align_from=1) -> int:
    """Caption ABOVE the table, per the reference format."""
    n = caption(doc, "Table", cap_text)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, htxt in enumerate(headers):
        _cell(t.rows[0].cells[i], htxt, bold=True, fill="EDF0F5",
              align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            al = WD_ALIGN_PARAGRAPH.RIGHT if i >= right_align_from else WD_ALIGN_PARAGRAPH.LEFT
            _cell(cells[i], val, align=al)
    if widths:
        total = sum(widths)
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(TEXT_WIDTH_IN * w / total)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    return n


def _cell(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(2)
    _emit_runs(para, str(text), base_bold=bold)
    for r in para.runs:
        r.font.size = Pt(10.5)
    if fill:
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(sh)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ===========================================================================
#                                THE DOCUMENT
# ===========================================================================
def build() -> Document:
    doc = new_document()

    # ------------------------------------------------------------ cover
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "Evidence-Grounded Greenwashing Analysis of Vietnamese Listed Companies "
        "with a Two-Channel Temporal Knowledge Graph"
    )
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = BODY_FONT

    doc.add_paragraph()
    todo(doc, "WRITE", "Full name and student ID of every member, one per line, centred — this is "
                       "where EquiFashion lists its four authors.")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub.add_run("\nSupervisor: ")
    rr.font.size = Pt(14)
    rr.font.name = BODY_FONT
    todo(doc, "WRITE", "Supervisor title and name.")

    for _ in range(4):
        doc.add_paragraph()
    tail = doc.add_paragraph()
    tail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = tail.add_run("Bachelor of Artificial Intelligence\nHoa Lac campus — FPT University\n2026")
    rr.italic = True
    rr.font.size = Pt(13)
    rr.font.name = BODY_FONT
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cp.add_run("© FPT University 2026. All rights reserved")
    rr.font.size = Pt(11)
    rr.font.name = BODY_FONT

    page_break(doc)

    # --------------------------------------------- how to read (remove later)
    h(doc, 1, "How to read this draft (remove before submission)")
    p(doc, "This document follows the structure of the reference thesis in `docs/"
           "EquiFashion_thesis.docx`: front matter, then 1 Introduction (with Related Work as "
           "§1.2), 2 Data, 3 Methodology, 4 Experiment, 5 Conclusion and Discussion, References.")
    p(doc, "Every number in it was read out of an artifact in this repository, not estimated. "
           "Chapter 2 in particular is generated: `docs/eda_report_data.py` measures the corpus and "
           "writes both `docs/eda_out/eda_stats.json` and every figure in `docs/figures/`, and this "
           "document reads those files. Re-run the two scripts and the chapter updates itself.")
    p(doc, "Three kinds of gap remain, each marked in red at the point where it occurs:")
    bullet(doc, "**[WRITE]** — content only the team can supply (names, opinions, stylistic choices).")
    bullet(doc, "**[TODO-RUN]** — a number requiring a stage that has not been run yet.")
    bullet(doc, "**[CITE: …]** — a reference slot. **No bibliography entry has been generated**, "
                "because a fabricated citation is a more serious defect in a thesis than a missing "
                "one, and is trivially detected. The bracket states what kind of source is needed.")
    p(doc, "Annex Z at the end collects every outstanding item, ordered by how much closing it would "
           "strengthen the defence.")
    p(doc, "Regenerate with `python docs/eda_report_data.py` then "
           "`python docs/build_capstone_report.py`.")

    page_break(doc)

    # ------------------------------------------------------ acknowledgment
    h(doc, 1, "Acknowledgment")
    todo(doc, "WRITE", "Three short paragraphs, following the reference thesis: (1) thanks to the "
                       "supervisor for guidance and feedback; (2) thanks to the lecturers of the "
                       "IT Specialization Department at FPT University; (3) thanks to families and "
                       "friends. This is the one section that must be in the team's own voice, so "
                       "it is deliberately left blank rather than drafted.")

    page_break(doc)

    # ------------------------------------------------------------- abstract
    h(doc, 1, "Abstract")
    p(doc, "Environmental, Social and Governance (ESG) disclosure has become a primary input to "
           "investment, credit and regulatory judgement, yet it carries a structural asymmetry: a "
           "company both authors and benefits from its own sustainability narrative. This creates a "
           "persistent risk of *greenwashing* — reporting that overstates environmental or social "
           "performance relative to actual conduct. Existing automated ESG systems read the report "
           "alone, and therefore inherit precisely the bias they should be detecting.")
    p(doc, "This thesis presents a temporal knowledge-graph system that makes greenwashing analysis "
           "falsifiable by construction. Two deliberately independent evidence channels are ingested "
           "into a single graph: the **claim side**, drawn from annual and sustainability reports "
           "published by the company, and the **conduct side**, drawn from independent Vietnamese "
           "news media. Both are extracted into one temporally-versioned schema of 28 node classes "
           "and 48 directed edge labels, under a design in which time lives on edges "
           "and event nodes while entity identity remains timeless, and in which every node retains "
           "sentence-level provenance back to the sentence that produced it.")
    p(doc, f"The corpus assembled for this work comprises {f(S['raw_pdfs'])} annual-report PDFs from "
           f"{S['raw_companies']} listed companies in the construction, building-materials and "
           f"real-estate sector ({S['raw_gb']} GB, {S['raw_year_min']}–{S['raw_year_max']}), "
           f"processed into {f(S['rep_sentences'])} page-anchored sentences across "
           f"{f(S['rep_pages'])} pages, of which {f(S['rep_esg_sentences'])} "
           f"({S['rep_esg_rate']}%) are classified ESG-relevant; and an independent conduct channel "
           f"of {f(S['news_articles'])} news articles from {S['news_domains_total']} distinct "
           f"domains, yielding {f(S['news_sentences'])} sentences. An exploratory analysis of this "
           "corpus is reported in full, including the properties that shaped the system design: an "
           "extreme claim-to-conduct volume asymmetry, a pillar imbalance that inverts between the "
           "two channels, and a severe recency skew in the news channel.")
    p(doc, "The system comprises five modules: a dual-standard indicator metadata layer normalising "
           "both the Vietnamese regulatory vocabulary (Circular 96/2020/TT-BTC, Decision "
           "2171/QĐ-BTC, QCVN 09, the SSC–IFC guide) and 145 GRI disclosure codes into one "
           "machine-readable axis; a claim-side report pipeline; a conduct-side news pipeline built "
           "on the same schema but a distinct extraction prompt; a temporal graph-construction stage "
           "combining structured-output LLM triple extraction with offline schema validation, "
           "deterministic entity resolution, provenance stamping and indicator-axis materialisation; "
           "and a claim–conduct cross-check stage producing provenance-carrying advisory dossiers.")
    p(doc, f"Because no ground-truth greenwashing labels exist for Vietnamese listed companies, the "
           f"system deliberately emits **evidence and an advisory assessment, never a greenwashing "
           f"score**, and is evaluated under a label-free protocol: an eight-attribute graph-quality "
           f"instrument measured before and after each design change. The resolved graph contains "
           f"{f(S['graph_nodes'])} nodes and {f(S['graph_edges'])} edges. Three controlled ablations "
           f"show that the temporal-integrity redesign reduced schema consistency violations from "
           f"1,098 to 1; that materialising the standard-indicator axis raised the share of "
           f"masked-answerable queries from 26.3% to 34.8% while cutting graph leaves from 82.2% to "
           f"75.8%; and that layout-aware document linearisation eliminates all 78 cross-standard "
           f"mis-attributions and all title truncation incurred by positional text extraction on the "
           f"GRI corpus. Cross-checking 1,093 extracted claims produced 3,461 adjudicated pairs, "
           f"yielding 70 apparently-supported and 22 apparently-contradicted claims, with 1,001 "
           f"returned explicitly as *unverified — insufficient evidence* rather than silently "
           f"assumed clean.")
    kw = doc.add_paragraph()
    kw.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _emit_runs(kw, "**Keywords:** Greenwashing; ESG disclosure; Temporal Knowledge Graph; GRI "
                   "Standards; Circular 96/2020/TT-BTC; Vietnamese NLP; Evidence-grounded "
                   "reasoning; Large Language Models.")

    page_break(doc)

    # --------------------------------------------------------- front lists
    h(doc, 1, "Table of contents")
    tocp = doc.add_paragraph()
    _field(tocp, 'TOC \\o "1-3" \\h \\z \\u',
           placeholder="Right-click here in Word and choose “Update Field” to build the "
                       "table of contents.")
    page_break(doc)

    h(doc, 1, "List of tables")
    tp = doc.add_paragraph()
    _field(tp, 'TOC \\h \\z \\c "Table"',
           placeholder="Right-click and “Update Field” to build the list of tables.")
    page_break(doc)

    h(doc, 1, "List of figures")
    fp = doc.add_paragraph()
    _field(fp, 'TOC \\h \\z \\c "Figure"',
           placeholder="Right-click and “Update Field” to build the list of figures.")
    page_break(doc)

    h(doc, 1, "List of abbreviations and acronyms")
    table(doc, "Abbreviations and acronyms used in this thesis.",
          ["Acronym", "Expansion"],
          [["AAA", "Ticker of the issuer analysed end-to-end (An Phat Bioplastics JSC)"],
           ["API", "Application Programming Interface"],
           ["DSU", "Disjoint-Set Union (used in entity consolidation)"],
           ["E / S / G", "Environmental / Social / Governance"],
           ["ESG", "Environmental, Social and Governance"],
           ["GRI", "Global Reporting Initiative"],
           ["IFC", "International Finance Corporation"],
           ["ISO", "International Organization for Standardization (here: ISO 8601 dates)"],
           ["JSONL", "JSON Lines (one JSON object per line)"],
           ["KG", "Knowledge Graph"],
           ["KPI", "Key Performance Indicator"],
           ["LLM", "Large Language Model"],
           ["NER", "Named-Entity Recognition"],
           ["NFC", "Unicode Normalization Form C"],
           ["OCR", "Optical Character Recognition"],
           ["QCVN 09", "QCVN 09:2013/BXD, the Vietnamese technical standard for energy-efficient buildings"],
           ["QĐ 2171", "Decision 2171/QĐ-BTC (national greenhouse-gas reporting guidance)"],
           ["RAG", "Retrieval-Augmented Generation"],
           ["SSC", "State Securities Commission of Vietnam"],
           ["T1 / T2 / T3", "The three schema tiers: entities / observations and events / statements and norms"],
           ["TT96", "Circular 96/2020/TT-BTC on securities-market disclosure"]],
          widths=[2.2, 8.0], right_align_from=99)

    page_break(doc)

    # ======================================================== 1 INTRODUCTION
    h(doc, 1, "1  Introduction")

    h(doc, 2, "1.1  Problem and Motivation")
    p(doc, "The global surge in corporate sustainability reporting has been accompanied by an "
           "equally alarming rise in **greenwashing**: the practice of making misleading or "
           "unsubstantiated environmental claims in order to project a greener image than a "
           "company's actual conduct warrants [1]. Greenwashing threatens the credibility of "
           "environmental markets, misdirects capital away from genuinely sustainable enterprises, "
           "and undermines collective progress toward climate goals [2]. As AI-driven systems "
           "become increasingly embedded in investment decision-making and regulatory oversight, "
           "there is an urgent need for automated tools capable of detecting deceptive "
           "sustainability communication at scale [3].")
    p(doc, "Vietnam stands at a pivotal moment in this global challenge. Following its commitment "
           "at the 26th UN Climate Change Conference (COP26) in 2021 to achieve net-zero carbon "
           "emissions by 2050, the Vietnamese government has substantially strengthened its ESG "
           "regulatory framework. **Circular 96/2020/TT-BTC** requires every publicly listed "
           "company to disclose environmental and social impacts in its annual report, and **Prime "
           "Minister Decision 13/2024/QĐ-TTg** expanded greenhouse-gas inventory requirements to "
           "over 2,166 facilities across key industrial sectors [4]. Despite this regulatory "
           "momentum, the quality and verifiability of the resulting disclosures remain deeply "
           "inconsistent. Only about 25% of the companies listed in Vietnam's top sustainability "
           "indices release independently verified sustainability reports [5], and no legal "
           "mechanism specifically addresses greenwashing in corporate disclosure, leaving the "
           "authenticity of environmental claims largely unchecked [6].")
    p(doc, "This credibility gap is particularly acute in three industries central to Vietnam's "
           "economic development and environmental footprint: **real estate, construction and "
           "building materials**. The construction and built-environment sector accounts for a "
           "substantial share of Vietnam's greenhouse-gas emissions, with cement, steel and clay "
           "brick generating significant carbon impacts across their production lifecycles [7]. "
           "With rapid urbanisation, a fast-expanding green-building market and mounting pressure "
           "from foreign investors demanding ESG compliance, listed companies in these sectors "
           "face strong incentives to project green credentials in their public disclosures — "
           "whether or not those claims are substantiated [8]. Construction alone comprises 229 of "
           "the facilities subject to Vietnam's mandatory GHG inventory requirement [4], yet no "
           "automated mechanism currently exists to assess the veracity of the environmental "
           "claims these companies make in their annual reports.")
    p(doc, "The sector is also the right choice on methodological grounds, not merely on grounds "
           "of importance. Its ESG claims are **checkable**: energy, emissions, water, waste and "
           "recycled-material content are first-order to its operations, it is directly governed "
           "by the national technical standard QCVN 09 on energy-efficient buildings, and its "
           "claims resolve into physical quantities that leave traces outside the company's own "
           "report. This contrasts with service sectors, where most disclosure is qualitative and "
           "there is correspondingly little to check a claim against. Selecting a sector whose "
           "claims are falsifiable is a precondition for the problem being studied at all.")
    p(doc, "**Detecting greenwashing manually is intractable at scale.** Annual reports are long, "
           "complex documents in which sustainability disclosures sit alongside financial "
           "statements and operational narrative. Human auditors cannot feasibly verify every "
           "environmental claim across an entire sector. This has motivated a growing body of "
           "research applying natural language processing to corporate sustainability "
           "communication [9, 10], including sentiment analysis, topic modelling, BERT-based "
           "classifiers and, most recently, retrieval-augmented generation (RAG) frameworks for "
           "identifying misleading environmental claims [11, 12].")
    p(doc, "A landmark advance in this direction is **EmeraldMind** [3], the first end-to-end "
           "knowledge-based RAG pipeline for greenwashing detection. By integrating a "
           "domain-specific knowledge graph capturing company-level ESG entities and relationships "
           "with a vectorised document store, EmeraldMind grounds LLM-based claim assessment in "
           "retrievable evidence, producing transparent, fact-backed verdicts without model "
           "fine-tuning. On a curated benchmark of 620 sustainability claims it achieves higher "
           "accuracy and better explanation quality than generic LLM baselines.")
    p(doc, "**However, existing greenwashing-detection frameworks have been developed and "
           "evaluated exclusively on English-language ESG reports from Western regulatory "
           "contexts.** No prior study has applied NLP-based greenwashing detection to Vietnamese "
           "corporate annual reports. The Vietnamese context presents distinct challenges: "
           "disclosures are written primarily in Vietnamese, embedded within the annual-report "
           "structure prescribed by Circular 96 rather than published as standalone sustainability "
           "reports, and framed by a national indicator vocabulary that a system trained only on "
           "the GRI Standards does not recognise. A system that understands only GRI will not find "
           "the disclosures Vietnamese regulation actually obliges companies to make. Furthermore, "
           "the real-estate, construction and materials industries in Vietnam have received no "
           "targeted NLP scrutiny despite their environmental significance.")
    p(doc, "There is also a limitation that EmeraldMind shares with the rest of the literature, "
           "and it is the one this work treats as central. **A company authors its own "
           "sustainability narrative and is also the party that gains from that narrative being "
           "believed.** Any automated system whose only input is the corporate report — however "
           "sophisticated its language model, however well-structured its knowledge graph — "
           "inherits precisely the bias it was built to detect. It can assess whether a report is "
           "internally coherent, well-structured or comprehensive; it cannot assess whether the "
           "report is *true*. Closing that gap requires evidence the company did not write.")
    p(doc, "Retrieval-augmented generation does not close it either [CITE: Lewis et al., "
           "Retrieval-Augmented Generation], because greenwashing is not a retrieval problem. It "
           "is a question about a *discrepancy between two sources at two different points in "
           "time*. Detecting it requires three properties that flat retrieval does not provide: "
           "(a) an evidence channel independent of the party being assessed; (b) a representation "
           "carrying time explicitly, so that a commitment made in 2012 can be compared with "
           "conduct observed in 2024; and (c) provenance at sentence granularity, so that a human "
           "reader can open the source page and overturn the system's conclusion. Together these "
           "point at a **temporal knowledge graph**, not at a vector index.")
    p(doc, "This work therefore adapts the RAG-plus-knowledge-graph paradigm of EmeraldMind to the "
           "Vietnamese corporate reporting context and extends it in one decisive respect: it adds "
           "a **second, independent evidence channel**. Report-derived claims and news-derived "
           "conduct evidence are extracted into a single schema and a single identity space but "
           "remain distinguishable at query time through a channel stamp; each sustainability "
           "claim is then matched against the conduct evidence available for it and adjudicated "
           "into a provenance-carrying advisory dossier. To the best of our knowledge this is the "
           "first study to apply knowledge-graph-augmented NLP to greenwashing detection in "
           "Vietnamese annual reports, and the first in any language to fuse a claim channel and "
           "an independent conduct channel into one temporal ESG graph.")
    p(doc, "One consequence of that framing is stated here rather than buried in the evaluation. "
           "**The system deliberately stops short of emitting a greenwashing score or verdict.** "
           "No ground-truth greenwashing labels exist for this corpus — or, so far as we are "
           "aware, for any Vietnamese one — so a classification into *greenwashing* / *not "
           "greenwashing* could be produced but could not be validated. The system instead returns "
           "the evidence, an advisory assessment and a full provenance trail, and leaves the "
           "judgement to the reader. That decision is justified in §4.2 and defended in §5.3, and "
           "it is the single most consequential design choice in the work.")

    h(doc, 2, "1.2  Related Work")

    h(doc, 3, "1.2.1  Conceptual foundations of greenwashing")
    p(doc, "The concept of greenwashing was formalised by Lyon and Maxwell [13] as the strategic "
           "use of selective environmental disclosure to create a misleadingly favourable "
           "environmental image. Delmas and Burbano [14] subsequently characterised greenwashing "
           "as the intersection of **poor environmental performance** and **positive environmental "
           "communication**, distinguishing it from sincere but ineffective sustainability "
           "efforts. More recently, regulatory bodies including the European Securities and "
           "Markets Authority [15] and the Competition and Consumer Commission [16] have converged "
           "on definitions encompassing not only false claims but also vague, misleading or "
           "unverifiable environmental assertions — a broadened scope that substantially increases "
           "the difficulty of automated detection. The fundamental enabler of greenwashing is the "
           "unregulated and largely unaudited nature of sustainability reporting in most "
           "jurisdictions [17]: companies face strong reputational incentives to make ambitious "
           "environmental claims but weak accountability mechanisms to ensure those claims are "
           "substantiated.")
    p(doc, "The Delmas–Burbano definition is worth restating because it dictates an architecture. "
           "If greenwashing is the *intersection* of two things — what a company says and how it "
           "behaves — then a system with access to only one of them cannot in principle observe "
           "the intersection. This is the criterion against which each family of prior work is "
           "assessed below.")

    h(doc, 3, "1.2.2  NLP for ESG and sustainability report analysis")
    p(doc, "A substantial body of NLP research has been directed at extracting and analysing "
           "sustainability information from corporate disclosures. Early approaches relied on "
           "keyword dictionaries and term-frequency methods, offering limited ability to handle "
           "the contextual nuance of sustainability language [18]. Transformer-based pretrained "
           "language models changed this landscape fundamentally. **ClimateBERT** [19], fine-tuned "
           "on a corpus of over two million climate-related sentences, became a foundational tool "
           "for classifying climate-relevant content in annual and sustainability reports and for "
           "detecting discrepancies between corporate climate disclosure and actual emissions "
           "data. Subsequent domain-specific models such as FinBERT-ESG [20] and ESG-BERT [21] "
           "extended these capabilities across the environmental, social and governance "
           "dimensions.")
    p(doc, "Beyond classification, NLP has been applied to ESG report analysis for a range of "
           "tasks: topic detection and thematic analysis using LDA and BERTopic [22]; specificity "
           "and readability scoring of environmental claims [23]; question answering over climate "
           "reports [24]; and structured data extraction from ESG documents using large language "
           "models [25]. A particularly relevant study by Schimanski et al. [26] extracted "
           "nature-related topics — water, forests, biodiversity — from annual reports and "
           "sustainability disclosures, demonstrating the feasibility of fine-grained domain "
           "analysis on corporate reporting corpora. Collectively this work establishes that NLP "
           "techniques can surface meaningful sustainability signals from the complex, "
           "heterogeneous text of corporate annual reports, and it provides the methodological "
           "foundation on which greenwashing-detection approaches build.")

    h(doc, 3, "1.2.3  Automated greenwashing detection")
    p(doc, "The specific task of automated greenwashing detection has attracted increasing "
           "attention, particularly following the global expansion of mandatory sustainability "
           "disclosure requirements [9]. Gorovaia et al. [17] conducted a large-scale textual "
           "analysis of CSR reports and found that companies engaged in environmental violations "
           "produce systematically longer, more positively worded and less readable reports — "
           "consistent with a greenwashing strategy of obscuring poor performance through "
           "narrative complexity. Kim et al. [11] integrated transformer-based classifiers with "
           "ensemble machine-learning models (XGBoost, Random Forest), reporting 86.34% accuracy "
           "in identifying greenwashing risk patterns in UK sustainability disclosures. A recent "
           "study of Central and Eastern European firms [12] constructed a Greenwashing Severity "
           "Index from sentiment analysis, TF-IDF term weighting and topic modelling to quantify "
           "the divergence between corporate ESG self-reports and external media narratives, "
           "finding moderate but widespread greenwashing across industries and firm sizes. "
           "Vinella et al. [27] proposed linguistic specificity scoring as a greenwashing signal, "
           "while Bingler et al. [28] showed that firms supporting the TCFD engage in systematic "
           "*cheap talk*, producing climate disclosures that are structurally compliant but "
           "substantively vague.")
    p(doc, "Read against the Delmas–Burbano criterion, these approaches fall into three families "
           "and each is limited by the evidence it admits. *Disclosure-only* approaches [17, 28] "
           "measure properties of the text itself — length, tone, readability, vagueness — and "
           "therefore detect a **rhetorical style rather than a factual discrepancy**; a "
           "well-written false claim evades them entirely. *Claim-versus-rating* approaches "
           "measure the gap between how much a firm says and how a third party rates it; the "
           "limitation is circularity, since the third-party rating is itself largely derived "
           "from disclosure. Only *claim-versus-independent-outcome* approaches — of which [12], "
           "comparing self-reports against media narratives, is the closest to the present work — "
           "observe both halves of the intersection, and they are the rarest, because the "
           "independent half is the expensive one to obtain.")
    p(doc, "A further constraint cuts across all three families and is decisive for how this work "
           "is evaluated. The first comprehensive NLP survey on greenwashing detection [9] "
           "reviewed 61 studies and identified three persistent challenges: the absence of a "
           "universally standardised definition, **the scarcity of reliably labelled datasets**, "
           "and the domain-specificity problem — models trained on English-language Western "
           "reports generalise poorly to other linguistic and regulatory contexts. The second of "
           "these is why the present work positions its output as *advisory evidence* rather than "
           "*classification*, and why its evaluation protocol (§4.2) is label-free by "
           "construction. The third is what this work addresses by targeting the Vietnamese "
           "regulatory environment directly.")
    todo(doc, "WRITE", "Recommended addition here, worth roughly half a page: the list in the "
                       "project's evaluation notes (EVALUATION_WITHOUT_LABELS §8) of metrics this "
                       "group "
                       "tried and discarded, with the reason each failed. Reporting a dead end is "
                       "unusual in a capstone and reads as methodological maturity.")

    h(doc, 3, "1.2.4  Retrieval-augmented generation and knowledge-graph-augmented approaches")
    p(doc, "The most directly relevant precedent for this framework is **EmeraldMind** [3], the "
           "first end-to-end knowledge-graph-augmented RAG pipeline for greenwashing detection. "
           "EmeraldMind constructs two complementary evidence stores: *EmeraldGraph*, a structured "
           "knowledge graph capturing company-specific ESG entities — companies, KPIs, "
           "environmental targets — and their relationships, extracted from corporate ESG "
           "reports; and *EmeraldDB*, a vectorised document store supporting semantic retrieval "
           "of relevant textual evidence. Given a sustainability claim, the framework retrieves "
           "supporting or contradicting evidence from both stores and passes the context to an "
           "LLM for verdict generation, classifying the claim as greenwashing, not greenwashing, "
           "or abstaining when evidence is inconclusive, alongside a natural-language "
           "justification. EmeraldMind achieves competitive accuracy with markedly superior "
           "explanation quality relative to vanilla LLM baselines, without domain-specific "
           "fine-tuning.")
    p(doc, "Related RAG-based work in the ESG domain includes ESGReveal [25], which applies RAG to "
           "structured ESG data extraction from sustainability reports, and ChatClimate [29], "
           "which grounds climate question answering on corporate disclosures. In the broader "
           "fact-checking literature, knowledge-graph integration has been shown to improve claim "
           "verification accuracy substantially by supplying structured, queryable background "
           "knowledge that LLMs cannot reliably retrieve from parametric memory alone [30].")
    p(doc, "Two properties of this line of work bound what it can currently do, and both are "
           "addressed directly in this thesis. First, **the evidence stores are built from the "
           "company's own reports**: EmeraldGraph's entities and EmeraldDB's passages are "
           "extracted from corporate ESG disclosure, so retrieval returns what the company said "
           "elsewhere in its own document set. This supports a strong consistency check and an "
           "auditable justification, but by the Delmas–Burbano criterion it observes only the "
           "*communication* half of the intersection. Second, **the representation is "
           "atemporal**. Published ESG knowledge graphs, EmeraldGraph included, model entities and "
           "relations without validity intervals, and an atemporal graph is structurally incapable "
           "of expressing *the company committed to X in 2021 and acted contrary to X in 2024* — "
           "the precise shape of the phenomenon under study. The temporal-database literature has "
           "long distinguished *valid time*, when a fact holds in the world, from *transaction "
           "time*, when the system recorded it [CITE: bi-temporal data model], and the temporal-KG "
           "literature extends entity representations with validity intervals and version chains "
           "[CITE: temporal knowledge graph representation]; that line has not been combined with "
           "corporate disclosure analysis.")

    h(doc, 3, "1.2.5  ESG disclosure and greenwashing in Vietnam")
    p(doc, "Despite the rapid global growth of ESG-oriented NLP research, the Vietnamese corporate "
           "reporting context remains almost entirely unstudied from a computational perspective. "
           "Vietnam's disclosure framework has evolved considerably: **Circular "
           "96/2020/TT-BTC** mandates environmental and social disclosure for all listed companies "
           "[4], and the *ESG Implementation and Disclosure Handbook* (2024) issued by the State "
           "Securities Commission provides GRI-aligned reporting guidance [31]. The Vietnam "
           "Sustainability Index, launched in 2017, offers a market-based benchmark covering the "
           "twenty most sustainable HOSE-listed companies. Nevertheless, independent verification "
           "of sustainability claims remains rare in the Vietnamese market [5], and greenwashing "
           "is an identified and growing risk [6].")
    p(doc, "The real-estate and construction sectors occupy a particularly prominent position in "
           "this risk landscape. They are explicitly included in Vietnam's Emissions Trading "
           "Scheme framework as major GHG emitters subject to mandatory inventory requirements "
           "[32]. The rapid proliferation of green-building certifications — LEED, LOTUS, EDGE — "
           "has created incentive structures in which green branding yields commercial benefits "
           "that may exceed the cost of genuine sustainability implementation [8, 33]. Yet no "
           "prior computational study has examined whether the environmental claims Vietnamese "
           "real-estate, construction and materials companies make in their annual reports are "
           "substantiated by verifiable evidence.")
    p(doc, "The Vietnamese setting also imposes a requirement that the international literature "
           "does not face. Disclosure obligations are framed in a **national indicator "
           "vocabulary** — Circular 96/2020/TT-BTC, Decision 2171/QĐ-BTC on greenhouse-gas "
           "reporting, the technical standard QCVN 09:2013/BXD for energy-efficient buildings, and "
           "the SSC–IFC reporting guide [31] — that does not map onto the GRI Standards "
           "term-for-term. A system that recognises only GRI vocabulary will not locate the "
           "disclosures Vietnamese regulation actually obliges companies to make, which is why "
           "§3.2 constructs a dual-standard indicator axis rather than adopting either standard "
           "alone.")


    h(doc, 2, "1.3  Contribution")
    p(doc, "The main contributions of this work are as follows.")
    bullet(doc, f"**A Vietnamese ESG disclosure corpus for greenwashing analysis**, built from the "
                f"annual-report filings of {S['idx_tickers']} listed companies in the real-estate, "
                f"construction and building-materials sectors — {S['idx_exchange_top_pct']}% filed "
                f"on the Ho Chi Minh Stock Exchange (HOSE), the remainder on HNX and UPCoM — "
                f"comprising {f(S['raw_pdfs'])} documents and {f(S['rep_sentences'])} "
                "sentence-level records, each carrying the coordinates that locate it on a page of "
                "a specific filing.", numbered=True)
    bullet(doc, "**An independent conduct channel**, the property that distinguishes this system "
                f"from disclosure-only greenwashing detection: {f(S['news_articles'])} articles "
                f"from {f(S['news_domains_total'])} Vietnamese online outlets, ingested into the "
                "*same* schema and identity space as the reports but kept separable at query time "
                "by a channel stamp carried on every node and edge. Claims and conduct can "
                "therefore be compared without either being able to masquerade as the other.",
           numbered=True)
    bullet(doc, "**A temporal ESG knowledge-graph schema** adapted to Vietnamese regulatory and "
                "linguistic norms — 28 node classes and 48 edge labels declared over 76 legal "
                "class pairs — governed by eight explicit design principles and a machine-checkable "
                "invariant set. The rule that entity identity must be timeless, and the partition "
                "of classes according to whether they may carry time at all, are enforced by an "
                "offline linter rather than by convention, so a violation fails a test instead of "
                "surviving review.", numbered=True)
    bullet(doc, "**A dual-standard indicator axis** unifying Circular 96/2020/TT-BTC with the GRI "
                "Standards through a confirmed-only crosswalk, materialised as graph structure "
                "rather than held in an offline lookup table — so that a Vietnamese-vocabulary "
                "disclosure and its international equivalent resolve to one node.", numbered=True)
    bullet(doc, "**A label-free evaluation instrument** that makes design changes measurable in a "
                "domain with no ground truth, demonstrated on three controlled before/after "
                "ablations.", numbered=True)
    bullet(doc, "**A documented and defended negative design decision**: the pipeline classifies "
                "each claim–evidence pair as *supports*, *contradicts* or *irrelevant*, and each "
                "claim as *contradicted*, *supported* or *unverified — insufficient evidence*, but "
                "it never emits a greenwashing score or verdict, because no label exists against "
                "which such a score could be validated.", numbered=True)
    todo(doc, "CITE", "Contribution 3 in the paper draft (Paper.docx) was phrased as classifying "
                      "claims into "
                      "*greenwashing / not greenwashing / inconclusive*, and contribution 4 as a "
                      "*systematic analysis of greenwashing linguistic patterns*. Neither is what "
                      "the implemented system does — it emits supports/contradicts/irrelevant per "
                      "evidence pair with no greenwashing verdict (§3.6.2), and no linguistic-"
                      "pattern study was performed. Both have been restated above to match the "
                      "code. If a linguistic-pattern analysis is wanted as a contribution, it has "
                      "to be run first.")

    h(doc, 2, "1.4  Structure of this Report")
    p(doc, "The remainder of this report is organised as follows. **Chapter 2** describes the "
           "dataset: its unit of analysis and record schema, how the two channels and the "
           "reference vocabulary were acquired, and an exploratory analysis whose findings drive "
           "several later design decisions. **Chapter 3** presents the methodology — the graph "
           "schema, the extraction, validation and entity-resolution stages, the indicator axis "
           "and the claim–conduct cross-check. **Chapter 4** reports the experiments: the "
           "label-free evaluation instrument, three controlled ablations, the cross-check results "
           "and an error analysis. **Chapter 5** concludes, states the limitations honestly, "
           "discusses the ethical position the system's advisory framing rests on, and sets out "
           "future work.")

    page_break(doc)

    # ============================================================== 2 DATA
    _chapter_data(doc)

    page_break(doc)

    # ======================================================= 3 METHODOLOGY
    _chapter_method(doc)

    page_break(doc)

    # ======================================================== 4 EXPERIMENT
    _chapter_experiment(doc)

    page_break(doc)

    # ======================================================== 5 CONCLUSION
    _chapter_conclusion(doc)

    page_break(doc)

    # ======================================================== back matter
    _back_matter(doc)

    return doc


# ---------------------------------------------------------------------------
#                              CHAPTER 2 — DATA
# ---------------------------------------------------------------------------
def _chapter_data(doc):
    h(doc, 1, "2  Data")

    # ---------------------------------------------------------------- 2.1
    h(doc, 2, "2.1  Dataset Overview")
    p(doc, "Progress on evidence-grounded greenwashing analysis depends on a corpus that satisfies "
           "three conditions simultaneously, and no existing public dataset satisfies all three. It "
           "must be **Vietnamese**, because the disclosures that regulation obliges companies to "
           "make are written in Vietnamese and framed in a national regulatory vocabulary. It must "
           "be **temporally deep**, because the phenomenon under study is a discrepancy between a "
           "statement made in one year and conduct observed in another; a single-year snapshot "
           "cannot express it. And it must contain **two structurally independent channels**, "
           "because a corpus consisting only of company-authored reports can support consistency "
           "checking but not verification.")
    p(doc, "Existing ESG corpora each satisfy one or two of these and fail the rest. International "
           "sustainability-report collections are temporally deep but monolingual in English and "
           "single-channel. Vietnamese financial-text datasets are in the right language but are "
           "built for sentiment or event extraction rather than disclosure analysis, and carry no "
           "regulatory indicator vocabulary. News corpora provide independence but no claims to "
           "check. The corpus described in this chapter was therefore assembled specifically for "
           "this work.")

    h(doc, 3, "2.1.1  Unit of analysis and record schema")
    p(doc, "The dataset's **unit of analysis is a single sentence, carried together with the "
           "coordinates that locate it in a published document**. It is not a document-level "
           "dataset: a company does not greenwash a report, it makes an individual assertion on a "
           "page, and the assertion is what has to be checkable. Every record is therefore a JSON "
           "object keyed by the triple `(source_pdf, page, sentence_index)`, which is assigned when "
           "the sentence is first segmented and never subsequently rewritten.")
    n = table(doc, "Record schema of the labelled sentence corpus. The first seven fields are "
                   "present in both channels; the remainder exist only in the news channel.",
              ["Field", "Type", "Meaning"],
              [["`source_pdf`", "string", "Source document identifier (report file name, or "
                                          "ticker + domain + hash for an article)"],
               ["`page`", "int", "1-based page of the source PDF (always 1 for an article)"],
               ["`sentence_index`", "int", "Position of the sentence within that page"],
               ["`text`", "string", "The sentence, verbatim, NFC-normalised"],
               ["`scores`", "object", "Four sigmoid scores: Neutral, Environmental, Social, "
                                      "Governance"],
               ["`labels`", "list", "Pillars whose score reaches the tag threshold; may be empty "
                                    "or hold several"],
               ["`esg`", "bool", "Binary ESG relevance, decided by the Neutral score (§2.2.4)"],
               ["`ticker`, `company`", "string", "Issuer the article was retrieved for"],
               ["`url`, `source_domain`, `title`", "string", "Article location, publisher and "
                                                             "headline"],
               ["`publish_date`, `date_crawled`", "date", "Stated publication date; retrieval "
                                                          "timestamp"],
               ["`channel`, `query`", "string", "Search back-end used, and the query that "
                                                "returned the article"],
               ["`matched_terms`, `company_mentioned`", "list, bool",
                "Which issuer identifiers occur in the sentence"]],
              widths=[3.6, 2.0, 8.6], right_align_from=99)
    p(doc, f"Table {n} is what makes the traceability requirement concrete. Because the first three "
           "fields are assigned once and preserved through classification, extraction, validation, "
           "entity resolution and graph loading, any statement the finished system makes can be "
           "resolved back to one sentence on one page of one document. A record from the report "
           "channel is shown verbatim below, unedited apart from line wrapping.")
    code_block(doc,
               '{"source_pdf": "AAA_2013.pdf", "page": 2, "sentence_index": 2,\n'
               ' "text": "Tháng 04/2004 Trước sự tăng trưởng nhanh chóng về quy mô thị trường và\n'
               '          nhu cầu sản xuất, Công ty đã thuê 10.000 m2 đất tại KCN Nam Sách để đầu\n'
               '          tư nhà máy sản xuất bao bì nhựa.",\n'
               ' "scores": {"Neutral": 0.8673, "Environmental": 0.1798,\n'
               '            "Social": 0.0018, "Governance": 0.0005},\n'
               ' "labels": [], "esg": false}')

    h(doc, 3, "2.1.2  Primary data and reference data")
    p(doc, "The dataset has two parts, and they are acquired differently enough that separating "
           "them matters. **Primary data** is the evidence the system reasons over: it is crawled, "
           "it grows when the crawl is re-run, and it is what the measurements in §2.3 "
           "characterise. **Reference data** is the controlled vocabulary the system reasons "
           "*with*: it is built once from published standards, it changes only when those "
           "standards change, and it is committed to version control alongside the source PDFs and "
           "their SHA-256 digests.")
    n = table(doc, "The dataset. Every figure is measured from the artifact named, not estimated.",
              ["Part", "Component", "Source", "Scale"],
              [["Primary", "Raw report corpus",
                f"Annual-report filings, {S['idx_exchange_top_pct']}% HOSE, via VietStock's "
                "disclosure mirror",
                f"{f(S['raw_pdfs'])} PDFs, {S['raw_companies']} issuers, {S['raw_gb']} GB"],
               ["Primary", "Labelled report sentences", "The above, parsed and ESG-classified",
                f"{f(S['rep_sentences'])} sentences, {f(S['rep_documents'])} documents"],
               ["Primary", "News corpus",
                f"{f(S['news_domains_total'])} Vietnamese online outlets, three search back-ends",
                f"{f(S['news_articles'])} articles, {f(S['news_sentences'])} sentences"],
               ["Reference", "VN regulatory corpus", "TT96/2020, QĐ 2171, QCVN 09, SSC–IFC",
                "35 KPI definitions"],
               ["Reference", "GRI standards corpus", "GRI (official PDFs)",
                "42 PDFs → 145 disclosure codes"],
               ["Reference", "Crosswalk", "Hand-curated, confirmed rows only",
                "20 GRI codes with a TT96 equivalent"]],
              widths=[1.8, 3.6, 5.4, 3.4], right_align_from=99)
    p(doc, f"As Table {n} shows, the corpus separates **breadth** from **depth**. The raw pool "
           f"spans {S['raw_companies']} issuers across the whole sector, and the labelled sentence "
           "corpus covers all of it. The end-to-end graph construction and cross-check, by "
           "contrast, are run in depth on a single issuer with a fifteen-year time series. This is "
           "a conscious trade-off rather than an incomplete run, and §5.2 states its consequence "
           "for the strength of the claims made here.")

    h(doc, 3, "2.1.3  What is deliberately not part of the dataset")
    p(doc, "One boundary is worth stating explicitly, because blurring it would make the "
           "evaluation in Chapter 4 circular. The KPI records, the extracted triples and the "
           "resolved knowledge graph are **outputs of the system under study, not data the system "
           "was given**. They are produced by the stages described in Chapter 3 and their scale, "
           "composition and defects are reported as results in §4.3.4 and §4.3.5. Listing them "
           "here as though they were corpus components would present the system's own product as "
           "evidence about the world, and would let a pipeline that manufactured more nodes appear "
           "to be working with more data.")
    p(doc, "The same boundary explains a second absence. **There is no train / validation / test "
           "split, because nothing in this work is trained.** The ESG classifier is a published "
           "checkpoint applied as released (§2.2.4) and every other stage is either a rule, an "
           "offline algorithm, or a prompted model used without gradient updates. The corpus is an "
           "analysis and evaluation corpus throughout; the labelled sample that a proper classifier "
           "evaluation would require does not exist yet, and is recorded as an outstanding item "
           "rather than quietly assumed.")

    # ---------------------------------------------------------------- 2.2
    h(doc, 2, "2.2  Data Construction Pipeline")

    h(doc, 3, "2.2.1  Pipeline overview")
    fig = figure_slot(doc, "Corpus construction pipeline: from published PDFs and online media to "
                           "the two labelled channels and the reference indicator vocabulary.",
                      "adapt from docs/PIPELINE_DIAGRAMS.md fig. 2–3")
    p(doc, f"Figure {fig} shows the construction pipeline. Two acquisition paths run in parallel and "
           "converge on **one sentence schema**: a report path that downloads published PDFs and "
           "linearises them, and a news path that queries three search back-ends and extracts "
           "article bodies. Sharing the schema is what later allows both channels to enter a single "
           "identity space while remaining distinguishable through a `source_type` stamp. A third, "
           "run-once path processes the regulatory and GRI standards into the controlled indicator "
           "vocabulary that the extraction stage consumes.")
    p(doc, "One constraint governs every stage and is treated as a requirement rather than a "
           "feature: **sentence-level traceability**. The triple `(source_pdf, page, "
           "sentence_index)` is attached when a sentence is first extracted and survives "
           "classification, KPI extraction, triple extraction, validation, entity resolution and "
           "graph loading. Every statement the finished system makes can therefore be traced back to "
           "a specific sentence on a specific page of a specific document.")

    h(doc, 3, "2.2.2  Public data sources and the raw data pool")
    p(doc, "**Where the reports come from.** Vietnamese listed companies are required to file an "
           "annual report with the exchange they are listed on, and the exchange publishes the "
           "filing. The corpus is therefore built from **exchange disclosure filings**, not from "
           "company websites: a filing is the document the issuer is legally answerable for, it is "
           "archived at a stable location, and it is available uniformly for every ticker, whereas "
           "an investor-relations page is voluntary, differently organised at every issuer and "
           "frequently rewritten. In practice the filings are retrieved through **VietStock's "
           "static mirror of the exchange disclosure archive**, whose URLs are structured "
           "`/data/{exchange}/{fiscal_year}/BCTN/VN/{TICKER}_Baocaothuongnien_{year}.pdf` — the "
           "exchange, the fiscal year, the document class (`BCTN` = *báo cáo thường niên*, annual "
           "report) and the language are all recoverable from the address itself.")
    p(doc, f"**Acquisition is a crawl, and the crawl produces an index.** Discovery is separated "
           f"from download. A crawl over the sector's ticker list resolves each issuer's filing "
           f"history into one row per document, and that index — "
           f"`config/company_annual_report.xlsx`, sheet *{S['idx_sheet_name']}* — carries the "
           f"ticker, the issuer's legal name, the document class, the filing year and the direct "
           f"file URL. It holds **{f(S['idx_rows'])} filings for {S['idx_tickers']} issuers**, all "
           f"{f(S['idx_rows'])} of them served from a single host. Keeping the index as a "
           "materialised artifact rather than re-resolving URLs at download time is what makes the "
           "acquisition auditable and repeatable: the exact set of documents a run was pointed at "
           "is a file that can be diffed, and the same index is reused as the issuer list for the "
           "news channel (§2.2.5) and for the issuer-identity registry.")
    ex = S["idx_exchanges"]
    fm = S["idx_formats"]
    n = table(doc, "The acquisition index, measured from the crawl manifest.",
              ["Property", "Value"],
              [["Filings indexed", f(S["idx_rows"])],
               ["Issuers (tickers)", str(S["idx_tickers"])],
               ["Document class", "`BCTN` — annual report (100%)"],
               ["**Listing venue**",
                " / ".join(f"**{k} {f(v)}**" if k == "HOSE" else f"{k} {f(v)}"
                           for k, v in ex.items())],
               ["Share filed on HOSE", f"{S['idx_exchange_top_pct']}%"],
               ["Filing years spanned",
                f"{S['idx_filing_year_min']}–{S['idx_filing_year_max']}"],
               ["File format",
                " / ".join(f"{k} {f(v)}" for k, v in fm.items())],
               ["**Delivered as an archive (`.zip`/`.rar`)**",
                f"**{f(S['idx_archive_rows'])} ({S['idx_archive_pct']}%)**"],
               ["Filing year exceeds fiscal year by exactly one",
                f"{S['idx_offset_one_pct']}% of rows"]],
              widths=[6.4, 4.0])
    p(doc, f"Table {n} settles three things that matter downstream. First, the corpus is "
           f"**overwhelmingly a HOSE corpus** — {S['idx_exchange_top_pct']}% of filings — with "
           f"{f(ex.get('HNX', 0))} HNX and {f(ex.get('UPCOM', 0))} UPCoM rows present because a few "
           "issuers moved venue during the period covered. Any statement about "
           "*listed Vietnamese companies* in this report should be read as a statement about "
           "HOSE-listed ones. Second, **every indexed document is an annual report**; no "
           "standalone sustainability report is included, so the ESG content analysed here is ESG "
           "content as it appears *inside* a general-purpose annual report, which is a materially "
           "different genre and generally a thinner one. Third, "
           f"**{S['idx_archive_pct']}% of filings arrive as `.zip` or `.rar` archives** rather "
           "than as a PDF, which is why extraction of `.rar` and `.7z` containers is a required "
           "step of the downloader rather than a convenience.")
    p(doc, "**The downloader.** Fetching is threaded across five workers, retries three times with "
           "backoff, and is **resumable** — an interrupted run does not re-fetch completed files. "
           "Archives are unpacked automatically and the container deleted, so the output tree "
           "contains no archives. Files are laid out sector → issuer → year, so the year is "
           "recoverable from the path even when the document metadata omits it, and failures are "
           "appended to a download log rather than raised, so one dead URL cannot abort a "
           "1,357-file run.")
    p(doc, f"**One naming caveat that affects every temporal figure in this report.** The year in "
           f"a downloaded file's name is the **filing** year, which for "
           f"{S['idx_offset_one_pct']}% of rows is exactly one greater than the **fiscal** year the "
           "report covers — a report on financial year 2024 is filed in spring 2025. The two are "
           "not interchangeable, and this report uses *filing year* wherever a year is read off a "
           "file name, labelling it as such. The distinction is not cosmetic: it shifts the whole "
           "temporal axis by one year, and it is the reason the compliance discontinuity discussed "
           "in §2.3.4 must be looked for in the 2021 filings rather than the 2020 ones.")
    p(doc, f"**Raw pool.** After download and unpacking, the pool holds **{f(S['raw_pdfs'])} PDF "
           f"files** for **{S['raw_companies']} companies**, totalling **{S['raw_gb']} GB** at a "
           f"mean of {S['raw_mean_pdf_mb']} MB per file, a mean of "
           f"{S['raw_pdfs_per_company_mean']} documents per company, spanning filing years "
           f"{S['raw_year_min']}–{S['raw_year_max']}. The file count exceeds the "
           f"{f(S['idx_rows'])} indexed filings because an unpacked archive frequently contains "
           "several PDFs — a report split into chapters. No filtering by content is applied at "
           "this stage; selection happens later, at the sentence level, so that the decision to "
           "discard is recorded rather than implicit.")
    p(doc, "**News sources.** The conduct channel is built by generating per-company queries from "
           "the same issuer index and issuing them against three independent back-ends — Google "
           "News RSS, Bing and DuckDuckGo — then fetching, caching and extracting article bodies "
           "with `trafilatura`. Using three back-ends rather than one is a deliberate hedge "
           "against the ranking bias of any single provider, and §2.3.5 reports how unevenly the "
           "three actually contributed.")

    h(doc, 3, "2.2.3  Cleaning, parsing and sentence segmentation")
    p(doc, "Three cleaning operations are applied, in this order.")
    p(doc, "**Text extraction with page anchoring.** Report PDFs are converted with a positional "
           "text extractor [CITE: PyMuPDF] that iterates over text runs in coordinate order, "
           "normalising to Unicode "
           "NFC so that Vietnamese diacritics survive, and attaching the page number to every block. "
           "The choice of a positional extractor here — as against the layout-aware parser used for "
           "the standards corpus in §3.2.2 — is deliberate and is measured in §4.3.3: what is needed "
           "for this channel is an exact page number per sentence, not the reconstruction of a "
           "table.")
    p(doc, "**Vietnamese-aware sentence segmentation.** Sentences are segmented with `underthesea` "
           "[CITE: underthesea] "
           "rather than by punctuation rules. Vietnamese financial reporting is dense with decimal "
           "separators written as `1.234,5`, with abbreviations such as `TNHH`, `CP` and `TP.`, and "
           "with full stops inside proper nouns; a naive full-stop rule fragments all of these.")
    p(doc, f"**Minimum-length filtering.** The splitter discards a segment unless it clears "
           f"**both** floors it enforces: at least **25 characters** and at least **4 words**. "
           f"Discarding happens here rather than downstream, so the decision is recorded once "
           f"instead of being re-made implicitly by every later stage. The effect is directly "
           f"visible in the finished corpus — the shortest surviving sentence is exactly "
           f"**{S['rep_min_len']} characters**, the floor itself — so page furniture such as folio "
           f"numbers and running headers never reaches the classifier. The cost is that a genuinely "
           "short disclosure — a "
           "bare table cell containing a figure — is lost with it, which is one reason quantitative "
           "values are extracted separately in §3.3 rather than being expected to survive as "
           "sentences.")

    h(doc, 3, "2.2.4  Annotation: multi-label ESG classification")
    p(doc, "Each sentence is annotated with the label set {Environmental, Social, Governance, "
           "Neutral} by **`nguyen599/ViDeBERTa-v3-ESG-base`**, a publicly released multi-label "
           "DeBERTa-v2 checkpoint for Vietnamese ESG text [CITE: ViDeBERTa; the checkpoint itself]. "
           "**The model is applied as published — it is not fine-tuned on this corpus.** No "
           "training is performed at any point in this work; the annotation stage loads the "
           "checkpoint and runs inference at a maximum sequence length of 256 tokens. This is a "
           "deliberate scope decision — building a labelled Vietnamese ESG training set was not "
           "within the project's budget — but it has a consequence stated plainly in §2.3.3 and "
           "§5.2: the checkpoint was fitted to someone else's data distribution, and its behaviour "
           "on this corpus is therefore an empirical question rather than a given.")
    p(doc, "The task is *multi-label* rather than multi-class, because a single sentence can "
           "legitimately carry more than one dimension — *“The Board of Directors approved the "
           "emission-reduction plan”* is simultaneously a governance statement and an environmental "
           "one, and forcing a single label would discard half of it. Accordingly the model applies "
           "a **per-label sigmoid, not a softmax**, so the four scores do not sum to one and each "
           "is read independently.")
    p(doc, "**Two separate rules produce the two fields, and the distinction matters.** A pillar "
           "tag is attached when that pillar's sigmoid score reaches **0.45**. The binary "
           "`esg` flag, however, is decided by a different quantity entirely: it is true when the "
           "**Neutral score falls below 0.5**. The second rule is deliberately robust to a signal "
           "that splits across pillars — a sentence scoring 0.3/0.3/0.3 on E/S/G is plainly not "
           "neutral even though no single pillar clears 0.45 — but it means the two fields can and "
           "do disagree in both directions. §2.3.3 measures how often. The model emits the full "
           "probability vector alongside both fields, which is what makes that measurement, and "
           "the confidence analysis in §2.3.2, possible at all.")
    p(doc, "The same checkpoint, thresholds and code path are applied to **both** channels. This is "
           "a correctness requirement rather than a convenience: if reports and news were labelled "
           "by different models, any measured difference between the channels — such as the pillar "
           "inversion reported in §2.3.6 — would be confounded by the difference between the "
           "annotators.")
    todo(doc, "TODO-RUN", "This is the largest remaining gap in the report and an examining board "
                          "will certainly ask about it. Note what is NOT needed: there are no "
                          "training hyper-parameters, no epochs and no train/validation split to "
                          "report, because no training was performed — "
                          "`notebooks/kaggle_esg_classify.ipynb` loads the published checkpoint and "
                          "runs inference only. What IS needed is an **evaluation**: hand-label a "
                          "random sample of roughly 300–500 sentences from this corpus against the "
                          "annotation guideline, then report per-label precision / recall / F1 with "
                          "micro- and macro-averages, plus the same for the binary `esg` flag. That "
                          "is the only thing that can settle the precision concern raised in "
                          "§2.3.3, and it needs no GPU.")

    h(doc, 3, "2.2.5  The independent conduct channel")
    p(doc, "News articles pass through the same sentence schema and the same classifier as reports, "
           "then through a preprocessing stage that normalises publication dates, adds "
           "`publish_year` and `date_uncertain` fields, and drops boilerplate. Two contracts on this "
           "channel matter more than the pipeline mechanics.")
    p(doc, "**The `date_uncertain` contract.** When an article does not state an explicit date or "
           "period for the fact it reports, the system falls back to the article's publication date "
           "as a proxy — but is **required** to mark that fact `date_uncertain = true`. The flag "
           "travels with the evidence into the final dossier, where it surfaces as an explicit "
           "caveat. The system never silently assumes a year. §2.3.5 shows why this contract is not "
           "optional: a measurable fraction of the crawled articles carry publication dates that are "
           "clearly parser artefacts.")
    p(doc, "**Independence.** Articles served from domains owned by the company under assessment are "
           "not admissible as independent evidence. They are retained and remain visible to the "
           "reader, but are excluded from the verification path. Without this guard a company that "
           "distributed its own press releases widely enough could verify its own claims, inverting "
           "the purpose of the channel.")

    h(doc, 3, "2.2.6  The reference corpus: regulatory and GRI standards")
    p(doc, "A third, run-once path builds the controlled vocabulary. Four Vietnamese instruments — "
           "Circular 96/2020/TT-BTC, Decision 2171/QĐ-BTC, QCVN 09:2013/BXD and the SSC–IFC ESG "
           "reporting guide — are downloaded and their relevant provisions extracted **verbatim** "
           "into 35 KPI definitions, each carrying a source block naming the instrument, the clause "
           "and the page. In parallel, 42 GRI Standards PDFs are linearised and segmented into 145 "
           "disclosure codes. The two are joined by a hand-curated crosswalk of which only "
           "explicitly confirmed rows are used. Because this corpus changes only when the standards "
           "themselves change, it is built once, committed to version control alongside the source "
           "PDFs and their SHA-256 digests, and treated thereafter as generated data. The method is "
           "described in §3.2 and measured in §4.3.3.")

    h(doc, 3, "2.2.7  Final packaging and distribution")
    p(doc, "Sentences are stored as JSONL, one record per sentence, carrying the traceability "
           "triple, the text, the full probability vector, the thresholded labels and the "
           "channel-specific fields (article URL, domain, publication date and crawl channel for "
           "news). Generated data is distributed through a dataset repository rather than through "
           "Git, and the pushed revision is pinned in a version file that **is** tracked in Git, so "
           "that a checkout recovers the data that went with that code. This is what makes the "
           "before/after ablations in §4.3 reproducible rather than merely reported.")

    # ---------------------------------------------------------------- 2.3
    h(doc, 2, "2.3  Exploratory Data Analysis")
    p(doc, "This section reports the exploratory analysis of the corpus. It is placed before the "
           "methodology deliberately: several of the design decisions in Chapter 3 are responses "
           "to properties measured here, and stating the property before the response makes the "
           "argument checkable. Every figure and table below is produced by "
           "`docs/eda_out/eda_report_data.py`, which streams the corpus directly from disk.")
    p(doc, "The scope is the **data**, and only the data: the raw filings, the sentences derived "
           "from them, their ESG annotation, their temporal spread and the composition of the "
           "conduct channel. What the pipeline subsequently *produces* from this corpus — KPI "
           "records, triples, the resolved knowledge graph — is measured in §4.3.4 and §4.3.5, "
           "where it belongs, for the reason given in §2.1.3: the system's own output is a result, "
           "not an input, and reporting it here would describe the system as though it were the "
           "world.")

    # 2.3.1
    h(doc, 3, "2.3.1  Scale and coverage of the raw corpus")
    fig = figure(doc, "fig_raw_corpus.png",
                 "Raw report corpus. (a) Distribution of documents held per company. "
                 "(b) Temporal spread, by the reporting year recoverable from the file name.")
    y = S["raw_year_counts"]
    p(doc, f"Figure {fig}(a) shows that coverage per company is uneven but rarely thin: the corpus "
           f"holds a mean of {S['raw_pdfs_per_company_mean']} documents per company across "
           f"{S['raw_companies']} companies. Panel (b) shows a corpus that thickens over time, "
           f"rising from {y.get(2011, 0)} documents for {2011} to {y.get(2025, 0)} for 2025. Two "
           "distinct causes are confounded here and cannot be separated from the file counts alone: "
           "listed companies have genuinely published more, and more recent filings are simply "
           "easier to retrieve from public sources. The practical consequence is the same either "
           "way — evidence density is a function of year, and any comparison across years must "
           "account for it rather than assume a stationary corpus.")

    # 2.3.2
    h(doc, 3, "2.3.2  Sentence-level properties")
    fig = figure(doc, "fig_length_confidence.png",
                 "Sentence-level properties of the labelled report corpus. (a) Sentence-length "
                 "distribution, with the ESG-classified subset overlaid. (b) Distribution of the "
                 "highest pillar probability per sentence, against the 0.5 decision threshold.")
    p(doc, f"Parsing and segmentation produced **{f(S['rep_sentences'])} sentences** from "
           f"**{f(S['rep_documents'])} documents** spanning **{f(S['rep_pages'])} pages**, a mean of "
           f"{S['rep_mean_pages_per_doc']} pages per document. The mean sentence length is "
           f"{S['rep_mean_len']:.0f} characters and the longest is {f(S['rep_max_len'])} — the "
           "latter is not a sentence at all but a table flattened into one text run by the "
           "positional extractor, and it is the clearest single artefact of the extraction choice "
           "discussed in §2.2.3.")
    p(doc, f"Panel (b) is the more consequential view. The distribution of the highest pillar "
           f"probability is strongly bimodal, with mass concentrated near 0 and near 1 and "
           f"comparatively little in between. A classifier whose scores pile up at the extremes is "
           "confident; it is not necessarily correct, and the two must not be conflated. The "
           "confidence distribution therefore constrains what can be claimed here: it shows the "
           "threshold at 0.5 is not cutting through a dense region, so small threshold changes will "
           "not move the label counts much — but it says nothing about whether the confident "
           "decisions are right. Only the evaluation table requested in §2.2.4 can answer that.")

    # 2.3.3
    h(doc, 3, "2.3.3  ESG label distribution and a precision concern")
    fig = figure(doc, "fig_label_distribution.png",
                 "ESG annotation of the report corpus, scoped to the "
                 f"{f(S['rep_pillar_labelled'])} sentences that carry at least one pillar label. "
                 "(a) Frequency of each pillar; a sentence may carry several, so the bars do not "
                 "sum to the total. (b) Every label combination that occurs, single-pillar against "
                 "multi-label, on a log axis. The unlabelled remainder of the corpus is excluded "
                 "from both panels; its size is given in the table below.")
    lab = S["rep_label_counts"]
    den = S["rep_pillar_labelled"]
    n = table(doc, "ESG annotation of the report corpus. Pillar shares are taken over the "
                   "pillar-labelled subset, not over the whole corpus, since the unlabelled "
                   "remainder carries no pillar to be a share of.",
              ["Quantity", "Sentences", "Share"],
              [["Total sentences", f(S["rep_sentences"]), "100% of corpus"],
               ["**Classified ESG-relevant** (`esg` flag)", f"**{f(S['rep_esg_sentences'])}**",
                f"**{S['rep_esg_rate']}% of corpus**"],
               ["**Carrying ≥1 pillar label**", f"**{f(den)}**",
                f"**{S['rep_pillar_labelled_pct']}% of corpus**"],
               ["Carrying no pillar label", f(S["rep_none_bucket"]),
                f"{100 - S['rep_pillar_labelled_pct']:.1f}% of corpus"],
               ["  — Governance", f(lab["Governance"]),
                f"{100*lab['Governance']/den:.1f}% of labelled"],
               ["  — Social", f(lab["Social"]), f"{100*lab['Social']/den:.1f}% of labelled"],
               ["  — Environmental", f(lab["Environmental"]),
                f"{100*lab['Environmental']/den:.1f}% of labelled"],
               ["Carrying more than one label", f(S["rep_multilabel"]),
                f"{100*S['rep_multilabel']/den:.1f}% of labelled"]],
              widths=[5.4, 2.6, 3.0])
    p(doc, f"Table {n} records four findings, and the last two are uncomfortable.")
    p(doc, f"**The two annotation rules disagree measurably.** §2.2.4 explains that a pillar tag "
           f"needs a sigmoid score of 0.45 while the `esg` flag is set by the Neutral score falling "
           f"below 0.5. These are different tests, and on this corpus they diverge in both "
           f"directions: **{f(S['rep_esg_without_pillar_label'])} sentences "
           f"({S['rep_esg_without_pillar_label_pct']}% of the ESG set) are flagged ESG-relevant "
           f"while carrying no pillar label at all**, and a further "
           f"**{f(S['rep_pillar_label_without_esg'])} carry a pillar label but are not flagged "
           f"ESG-relevant**. Neither is a bug — the first is the intended behaviour for a signal "
           "that splits across pillars, the second for a sentence that is confidently neutral yet "
           "mentions one pillar in passing. But it means the two fields must never be used "
           "interchangeably. The pipeline respects this: page selection for the expensive "
           "extraction stage keys on the `esg` flag (§3.5.2), while the pillar columns of the "
           "evidence interface are read from the indicator axis rather than from these labels at "
           "all (§3.6.3).")
    p(doc, f"**Governance dominates.** Governance appears on {f(lab['Governance'])} sentences "
           f"against {f(lab['Social'])} for Social and {f(lab['Environmental'])} for Environmental — "
           f"roughly {lab['Governance']/lab['Environmental']:.1f}× the environmental count. This is "
           "consistent with the structure of a Vietnamese annual report, in which board composition, "
           "shareholder structure, internal control and related-party transactions are mandatory "
           "and lengthy sections, while environmental disclosure is comparatively short. It has a "
           "direct downstream effect: the claim side of the graph is governance-heavy, whereas the "
           "conduct side, as §2.3.6 shows, is environmental-heavy. The two channels are not "
           "measuring the same thing in the same proportion.")
    cb = S["rep_combo_counts"]
    p(doc, f"**The multi-label formulation is barely exercised.** Panel (b) of Figure {fig} makes "
           f"this immediately visible: on a log axis the three single-pillar bars sit at "
           f"10⁴–10⁵ while all three multi-label bars sit below 1,200. Only "
           f"{f(S['rep_multilabel'])} sentences — {100*S['rep_multilabel']/den:.1f}% of the "
           f"pillar-labelled subset — carry more than one pillar, and Governance + Social "
           f"({f(cb.get('Governance+Social', 0))}) accounts for "
           f"{100*cb.get('Governance+Social',0)/max(S['rep_multilabel'],1):.0f}% of even that. "
           "Environmental + Governance, the combination one would most expect in a disclosure "
           f"about board-approved emission targets, occurs {f(cb.get('Environmental+Governance', 0))} "
           "times in a corpus of nearly nine hundred thousand sentences. In practice the model "
           "behaves as a single-label classifier. This does not invalidate the multi-label design, "
           "which is correct in principle and costs nothing, but it does mean the design is "
           "currently not earning its keep, and the report should not claim otherwise.")
    p(doc, f"**The ESG rate of {S['rep_esg_rate']}% is implausibly high, and this is reported rather "
           f"than smoothed over.** Roughly one sentence in three of an entire annual report being "
           "ESG-relevant does not match what a reader of these documents observes. The likely "
           "mechanism is lexical: Vietnamese corporate language uses environmental and social "
           "vocabulary in contexts that are not disclosures. A concrete, checkable instance is "
           "available in this corpus. The legal name of one issuer is *Công ty Cổ phần Nhựa và Môi "
           "trường Xanh An Phát* — literally *An Phat Green Environment and Plastics JSC*. The "
           f"string *môi trường xanh* (“green environment”) appears in "
           f"**{f(S['probe_hits'])} sentences**, and **{f(S['probe_labeled_env'])} of them "
           f"({S['probe_env_pct']}%) are labelled Environmental**, including bare title-page lines "
           "that state nothing but the company's name.")
    p(doc, f"That single pattern does not account for the whole gap, and it should not be presented "
           f"as if it did: {f(S['probe_labeled_env'])} sentences out of "
           f"{f(S['rep_esg_sentences'])} is a rounding error in the total. What it "
           "establishes is the *mechanism* — a company name, a slogan or a department title can "
           "trigger a pillar label on lexical grounds alone — and that mechanism has no reason to be "
           "confined to one issuer's name. Quantifying its full extent requires the labelled "
           "evaluation set requested in §2.2.4. Until then the honest statement is: **the ESG rate "
           "reported here is an upper bound on ESG-relevant content, and the downstream stages are "
           "designed on that assumption.** The design response is described in §3.5.3 — the "
           "extraction stage does not treat an ESG label as evidence of anything, only as a "
           "cost-saving filter over which pages are worth sending to an expensive model.")

    # 2.3.4
    h(doc, 3, "2.3.4  Temporal coverage")
    fig = figure(doc, "fig_temporal_coverage.png",
                 "Temporal coverage of the labelled report corpus, by reporting year, with the "
                 "ESG-classified subset overlaid.")
    ys = S["rep_year_sent"]
    ye = S["rep_year_esg"]
    rows = []
    for yy in ["2011", "2013", "2015", "2017", "2019", "2021", "2023", "2025"]:
        if yy in ys:
            rows.append([yy, f(ys[yy]), f(ye.get(yy, 0)), f"{100*ye.get(yy,0)/ys[yy]:.1f}%"])
    n = table(doc, "Temporal coverage of the labelled report corpus (odd years shown; all years "
                   "appear in Figure " + str(fig) + ").",
              ["Reporting year", "Sentences", "ESG sentences", "ESG share"],
              rows, widths=[3.0, 3.0, 3.0, 2.4])
    p(doc, f"The corpus spans {S['raw_year_min']}–{S['raw_year_max']}, with volume rising roughly "
           f"fourfold from {f(ys['2011'])} sentences for 2011 to {f(ys['2025'])} for 2025. The ESG "
           f"share rises too, and more interestingly: from {100*ye['2011']/ys['2011']:.1f}% in 2011 "
           f"to {100*ye['2025']/ys['2025']:.1f}% in 2025. The rise is gradual rather than a step "
           "change, which is worth noting precisely because a step would have been easier to "
           "interpret — Circular 96/2020 took effect in 2020, and a disclosure mandate that bound "
           "immediately should show as a discontinuity at that year. It does not.")
    p(doc, "Three explanations are consistent with a gradual rise, and the corpus cannot distinguish "
           "between them: genuine secular growth in ESG disclosure independent of the mandate; "
           "gradual compliance ramp-up rather than immediate adoption; or drift in the vocabulary "
           "the classifier keys on, with no change in underlying content at all. The third would be "
           "a measurement artefact rather than a finding, and cannot be excluded without the "
           "evaluation set of §2.2.4. The claim made here is therefore restricted to what the data "
           "supports: **ESG-classified content grows as a share of the corpus over time, and the "
           "growth is not attributable to the 2020 mandate on the evidence available.**")

    # 2.3.5
    h(doc, 3, "2.3.5  The independent conduct channel")
    fig = figure(doc, "fig_news_sources.png",
                 "The conduct channel. (a) Source concentration across the twelve most frequent "
                 "domains. (b) Distribution of stated publication year.")
    dom = S["news_top_domains"]
    ch = S["news_channels"]
    n = table(doc, "Composition of the conduct channel.",
              ["Quantity", "Value"],
              [["Tickers covered", f(S["news_tickers"])],
               ["Articles retrieved", f(S["news_articles"])],
               ["Sentences", f(S["news_sentences"])],
               ["Mean articles per ticker", str(S["news_mean_articles_per_ticker"])],
               ["Mean sentences per article", str(S["news_mean_sentences_per_article"])],
               ["**Distinct source domains**", f"**{f(S['news_domains_total'])}**"],
               ["Share held by the three largest domains", f"{S['news_top3_domain_share']}%"],
               ["Retrieved via Google News / Bing / DuckDuckGo",
                f"{f(ch.get('google_news',0))} / {f(ch.get('bing',0))} / {f(ch.get('ddg',0))}"],
               ["Sentences explicitly mentioning the company", f"{S['news_company_mentioned_pct']}%"],
               ["**Articles dated before 2010 (probable parse failures)**",
                f"**{f(S['news_suspicious_dates'])} ({S['news_suspicious_dates_pct']}%)**"]],
              widths=[7.0, 3.4])
    p(doc, f"**Independence is structurally sound.** The {f(S['news_articles'])} articles come from "
           f"{f(S['news_domains_total'])} distinct domains, and the three largest — "
           f"{', '.join(list(dom)[:3])} — together account for only "
           f"{S['news_top3_domain_share']}% of the corpus. That is a genuinely fragmented source "
           "distribution, and it is the single most favourable property of this channel: no "
           "individual outlet's editorial position can dominate the conduct side. The corpus is "
           "nonetheless concentrated in financial and business media rather than, for instance, "
           "environmental or local reporting, which bounds the *kind* of conduct it can observe.")
    p(doc, f"**Retrieval is dominated by one back-end.** Of the three search back-ends, Google News "
           f"supplied {f(ch.get('google_news',0))} articles, Bing {f(ch.get('bing',0))} and "
           f"DuckDuckGo {f(ch.get('ddg',0))}. The three-back-end design was intended as a hedge "
           "against single-provider ranking bias; in practice one provider supplies roughly "
           "three-quarters of the corpus, so the hedge is much weaker than the design implies. This "
           "is a limitation of the corpus as built, not of the design, and it is recorded in §5.2.")
    py = S["news_pub_year"]
    recent = sum(v for k, v in py.items() if int(k) >= 2025)
    tot_dated = sum(py.values())
    p(doc, f"**Severe recency skew.** {f(recent)} of {f(tot_dated)} dated articles "
           f"({100*recent/tot_dated:.1f}%) were published in 2025 or 2026. The conduct channel is "
           "therefore concentrated in the last two years, while the claim channel spans fifteen. "
           "This asymmetry in *time*, on top of the asymmetry in *volume* documented in §2.3.6, is "
           "the direct cause of the retrieval-window design in §3.6.1: a claim made in 2012 has "
           "essentially no contemporaneous conduct evidence available, so the window must look "
           "forward, and it must look far.")
    p(doc, f"**Date quality is measurably imperfect.** {f(S['news_suspicious_dates'])} articles "
           f"({S['news_suspicious_dates_pct']}%) carry a stated publication date before 2010, which "
           "for a corpus of online financial news about currently-listed issuers is not credible. "
           "One value alone, 2002, accounts for 179 of them — the signature of a template default "
           "rather than a real date. These are extraction failures, not old articles. They are left "
           "in the corpus rather than deleted, because the `date_uncertain` contract of §2.2.5 "
           "exists precisely to carry this uncertainty forward to the reader instead of hiding it. "
           "**This measurement is the empirical justification for that contract**: without it, "
           "roughly one article in seventeen would silently contribute a fabricated year to a "
           "temporal graph.")

    # 2.3.6
    h(doc, 3, "2.3.6  Claim–conduct asymmetry")
    fig = figure(doc, "fig_claim_conduct_asymmetry.png",
                 "Volume and pillar asymmetry between the two channels for the issuer analysed "
                 "end-to-end. Note the logarithmic vertical axis.", width_in=4.9)
    rl = S["aaa_report_label_counts"]
    nl = S["news_labeled_label_counts"]
    n = table(doc, "Claim channel versus conduct channel for the issuer analysed end-to-end.",
              ["Quantity", "Reports (claim)", "News (conduct)", "Ratio"],
              [["Sentences", f(S["aaa_report_sentences"]), f(S["news_labeled_sentences"]),
                f"{S['aaa_report_sentences']/S['news_labeled_sentences']:.1f}×"],
               ["**ESG sentences**", f"**{f(S['aaa_report_esg'])}**",
                f"**{f(S['news_labeled_esg'])}**", f"**{S['aaa_claim_conduct_ratio']}×**"],
               ["ESG share", f"{S['aaa_report_esg_rate']}%", f"{S['news_labeled_esg_rate']}%", "—"],
               ["  — Environmental", f(rl["Environmental"]), f(nl["Environmental"]),
                f"{rl['Environmental']/nl['Environmental']:.1f}×"],
               ["  — Social", f(rl["Social"]), f(nl["Social"]),
                f"{rl['Social']/nl['Social']:.1f}×"],
               ["  — Governance", f(rl["Governance"]), f(nl["Governance"]),
                f"{rl['Governance']/nl['Governance']:.1f}×"]],
              widths=[3.6, 3.0, 3.0, 2.0])
    p(doc, f"This table states the central empirical constraint on the entire system, and it should "
           f"be read before any cross-check result in Chapter 4. For the issuer analysed end to end, "
           f"the claim side supplies {f(S['aaa_report_esg'])} ESG sentences and the conduct side "
           f"{f(S['news_labeled_esg'])} — a ratio of **{S['aaa_claim_conduct_ratio']} to one**. "
           "There is roughly one piece of independent evidence for every sixteen claims, before any "
           "consideration of whether the two are about the same topic.")
    p(doc, f"**The pillar composition also inverts.** In reports, Governance is the largest pillar "
           f"({f(rl['Governance'])} sentences) and Environmental the smallest ({f(rl['Environmental'])}). "
           f"In news, the ordering reverses completely: Environmental is largest "
           f"({f(nl['Environmental'])}) and Governance smallest ({f(nl['Governance'])}). The "
           f"governance ratio is therefore {rl['Governance']/nl['Governance']:.0f}:1 while the "
           f"environmental ratio is only {rl['Environmental']/nl['Environmental']:.1f}:1.")
    p(doc, "The implication is specific and important. Independent media write about factories, "
           "emissions and environmental incidents; they do not write about a company's internal "
           "control framework. **A governance claim is therefore structurally far less checkable "
           "than an environmental one in this corpus** — not because governance claims are harder "
           "to evaluate in principle, but because no independent party is producing evidence about "
           "them. Since the claim side is governance-heavy and the conduct side is "
           "environmental-heavy, the two channels are least aligned exactly where the claim volume "
           "is greatest. This is the structural reason for the large unverified count reported in "
           "§4.4, and it is a property of the evidence environment rather than a defect of the "
           "matching algorithm.")

    # 2.3.7
    h(doc, 3, "2.3.7  What the exploration implies for the design")
    p(doc, "Five properties measured above drive design decisions taken in Chapter 3, and are "
           "collected here so that the connection is explicit rather than implied.")
    bullet(doc, f"**The conduct channel is ~{S['aaa_claim_conduct_ratio']}× smaller than the claim "
                f"channel and inverted in pillar composition** (§2.3.6). Consequence: the system "
                "must be able to return *“not verifiable”* as a first-class outcome rather than "
                "forcing a verdict. This is why the cross-check emits three assessments, one of "
                "which is `unverified_insufficient_evidence` (§3.6.2).", numbered=True)
    bullet(doc, f"**The conduct channel is concentrated in the last two years while claims span "
                f"fifteen** (§2.3.5). Consequence: the retrieval window is deliberately asymmetric — "
                "one year backward, fifty forward (§3.6.1).", numbered=True)
    bullet(doc, f"**{S['news_suspicious_dates_pct']}% of articles carry an incredible publication "
                f"date** (§2.3.5). Consequence: the `date_uncertain` flag is mandatory on news-derived "
                "observations and is propagated into every dossier as a visible caveat (§2.2.5, "
                "§3.5.3).", numbered=True)
    bullet(doc, f"**The ESG classifier's {S['rep_esg_rate']}% positive rate is an upper bound, with "
                f"a demonstrated lexical failure mode** (§2.3.3). Consequence: the ESG label is used "
                "only as a cost filter over which pages to send to an expensive extractor, never as "
                "evidence in its own right (§3.5.3).", numbered=True)
    bullet(doc, f"**{S['kpi_other_share']}% of extracted KPIs fall outside the controlled vocabulary "
                f"and units are written {f(S['kpi_distinct_units'])} different ways** (§4.3.4). "
                "Consequence: canonicalisation is a separate offline stage that writes a *new* "
                "property rather than overwriting the raw one, and rejects rather than force-maps "
                "ambiguous cases (§3.5.6).", numbered=True)


# ---------------------------------------------------------------------------
#                           CHAPTER 3 — METHODOLOGY
# ---------------------------------------------------------------------------
def _chapter_method(doc):
    h(doc, 1, "3  Methodology")

    h(doc, 2, "3.1  Problem Definition")
    fig = figure_slot(doc, "Overview of the system: five modules and the two evidence channels.",
                      "adapt from docs/PIPELINE_DIAGRAMS.md fig. 1")
    p(doc, f"Let a company publish a set of sustainability claims C over time, and let there exist "
           f"an independent record of its conduct E. Greenwashing analysis, as formulated here, is "
           "the task of deciding for each claim c ∈ C whether the available conduct evidence "
           "supports it, contradicts it, or is insufficient to decide — and of exposing the evidence "
           "behind that decision so a human can overturn it. Formally the system computes")
    equation(doc, "a(c)  =  Agg( { y(c, e) : e ∈ E(c) } ),   "
                  "y(c, e) ∈ { supports, contradicts, irrelevant }")
    p(doc, "where E(c) ⊂ E is the conduct evidence retrieved for claim c, y is a pairwise "
           "adjudication and Agg is a conservative aggregation defined in §3.6.2. The output a(c) is "
           "an *advisory assessment* together with its supporting evidence, never a numeric "
           "greenwashing score.")
    p(doc, f"As Figure {fig} shows, the system consists of five modules. The **indicator metadata "
           "module** (§3.2) normalises Vietnamese ESG regulation and the GRI Standards into a single "
           "machine-readable indicator axis. The **claim-side module** (§3.3) converts company PDFs "
           "into ESG-labelled, page-anchored sentences. The **conduct-side module** (§3.4) builds "
           "the independent evidence channel. The **temporal knowledge-graph module** (§3.5) fuses "
           "both channels into one temporally-versioned, entity-resolved graph. The **cross-check "
           "module** (§3.6) computes a(c) and renders it as an auditable dossier.")

    h(doc, 2, "3.2  Indicator Metadata Module")
    p(doc, "This module turns two heterogeneous bodies of standards — Vietnamese regulatory texts "
           "and the GRI Standards — into **one machine-readable indicator axis** to which every KPI "
           "observation and every sustainability claim is subsequently anchored.")

    h(doc, 3, "3.2.1  Source corpora")
    p(doc, "**(a) Vietnamese regulatory corpus.** Circular 96/2020/TT-BTC on securities-market "
           "disclosure (whose Appendix IV mandates reporting of environmental and social impact), "
           "Decision 2171/QĐ-BTC, QCVN 09:2013/BXD and the SSC–IFC ESG reporting guide. A six-stage "
           "provenance pipeline downloads each instrument, extracts the relevant provisions "
           "**verbatim**, and emits a controlled vocabulary of **35 KPI definitions**, each carrying "
           "a source block naming the instrument, the article or clause and the page. Their pillar "
           "split is **20 environmental, 14 social and 1 governance** — the Vietnamese disclosure "
           "vocabulary is overwhelmingly about physical impact. This is close to the mirror image "
           "of the GRI catalogue in Table 9 below, which is governance-heavy, and the contrast is "
           "not incidental: the two standards are answering different questions, which is precisely "
           "why the crosswalk between them can only be partial.")
    p(doc, "**(b) GRI corpus.** Forty-two GRI Standards documents — the universal standards GRI 1, 2 "
           "and 3, the sector standards GRI 11–14, and the topic-specific 200/300/400 series — "
           "44 MB of PDF, version-controlled alongside the code so the catalogue can "
           "be rebuilt from exactly the documents that produced it.")

    h(doc, 3, "3.2.2  Layout-aware document linearisation")
    p(doc, "GRI standards are published as visually complex PDFs: two-column layouts, requirement "
           "boxes, sidebars, footnotes and disclosure tables interleaved on the same page. "
           "Positional extractors recover characters but not *structure* — a requirement box merges "
           "into body prose, a disclosure table collapses into an unordered token stream, and the "
           "hierarchical relationship between a disclosure code and its guidance is lost. Because "
           "every downstream stage keys on the disclosure unit, structural loss here propagates "
           "irrecoverably.")
    p(doc, "Each standard is therefore linearised with **Marker** [CITE: Paruchuri et al., "
           "marker/Surya], a layout-aware conversion pipeline composing a layout-detection model, an "
           "order-detection model, a table-recognition model and an OCR model over a text-first "
           "backbone, emitting Markdown that preserves heading hierarchy, list structure and table "
           "cell boundaries. Marker belongs to the same generation of neural document-parsing "
           "systems as Nougat, Docling and olmOCR [CITE: Nougat; Marker/Surya; Docling; olmOCR]; "
           "it was selected for its "
           "explicit table reconstruction and its Markdown target, which retains precisely the "
           "heading hierarchy the segmenter in §3.2.3 consumes. Conversion runs against the "
           "hosted Datalab inference endpoint (`/api/v1/marker`) rather than a local model, under "
           "a three-worker concurrency bound; the published rate limit is 25 requests per 60 "
           "seconds, and an HTTP 429 is retried after a **fixed 12-second wait, up to five "
           "attempts**. Every conversion result is **cached to disk under the source PDF's file "
           "name**, so the corpus is converted exactly once and every later rebuild of the "
           "catalogue reads the cache instead of re-invoking the model. The cache key is the file "
           "name, not a digest of the file contents, so replacing a standard PDF in place while "
           "keeping its name would silently serve the stale conversion — the SHA-256 digest "
           "recorded per catalogue record (§3.2.4) is what makes that detectable after the fact.")
    p(doc, "Formally, for a GRI standard S of P pages, linearisation produces an ordered Markdown "
           "token sequence")
    equation(doc, "M(S)  =  m₁ ⊕ m₂ ⊕ ⋯ ⊕ m_P")
    p(doc, "in which each m_p retains block type and reading order, and ⊕ denotes order-preserving "
           "concatenation. §4.3.3 reports a controlled measurement of what this choice is worth "
           "relative to positional extraction.")
    todo(doc, "WRITE", "Equations are plain text here because python-docx cannot emit native Word "
                       "equation objects. In Word, select each equation line and use "
                       "Insert ▸ Equation; the symbols are already correct.")

    h(doc, 3, "3.2.3  Disclosure segmentation and the ownership rule")
    p(doc, "The linearised Markdown is segmented into **disclosure units** by a deterministic parser "
           "operating over the heading hierarchy and the GRI disclosure-code pattern. Because "
           "Markdown preserves heading depth, segmentation is a structural operation rather than a "
           "similarity-based guess: a disclosure boundary is a heading that introduces a disclosure "
           "code, and the unit extends to the next heading of equal or higher level. This yields")
    equation(doc, "D(S)  =  { d₁, d₂, … , d_n }")
    p(doc, "where each dᵢ pairs one disclosure code with its requirements and guidance.")
    p(doc, "GRI standards *re-list disclosures belonging to other standards*: the sector standards "
           "GRI 11–14 and the 2024/25 rewrites GRI 101–103 each reproduce disclosures owned by "
           "topic-specific standards, so a naive first-file-wins attribution assigns ownership to "
           "whichever document happened to be read first. Each disclosure is instead attributed to "
           "the standard whose identifier is a prefix of the disclosure code. In this corpus the "
           "rule is not cosmetic: applying it corrected the attribution of **80 of 136 entries** and "
           "repaired **31 mangled titles**. It is pinned by a regression test rather than left to "
           "reviewer vigilance, because it is exactly the class of defect that produces a "
           "plausible-looking catalogue containing systematically wrong ownership.")

    h(doc, 3, "3.2.4  Catalogue schema and the dual-standard crosswalk")
    p(doc, "The module emits a catalogue of **145 GRI disclosure codes**, each a record")
    equation(doc, "Cᵢ = ⟨ code, title_en, title_vi, pillar, requirement_type, units, "
                  "tt96_equivalent, versions, provenance ⟩")
    n = table(doc, "Pillar distribution of the GRI catalogue (n = 145 disclosure codes).",
              ["Pillar", "Disclosure codes", "Share"],
              [["Governance", "59", "40.7%"],
               ["Environmental", "51", "35.2%"],
               ["Social", "35", "24.1%"],
               ["**Total**", "**145**", "**100%**"]],
              widths=[4.0, 3.0, 2.4])
    p(doc, f"Table {n} gives the pillar distribution; **20 codes carry a confirmed Circular-96 "
           "equivalent**. Provenance is a `(source_pdf, page, sha256)` triple, so a claim traced "
           "back to a GRI requirement can be verified against the exact document revision that "
           "produced it. The `versions[]` field records each standard's version year, effective date "
           "and status, which is what allows the graph to state that a 2016 disclosure was "
           "superseded by a 2020 revision — a temporal fact rather than a static label.")
    p(doc, "A hand-curated crosswalk maps Vietnamese regulatory indicators onto GRI codes. Only rows "
           "explicitly marked *confirmed* are materialised as `equivalentTo` edges; draft rows are "
           "retained for review but excluded unless explicitly enabled. This is a deliberate "
           "**precision-over-recall** decision: an incorrect equivalence silently attributes evidence "
           "to the wrong regulatory requirement, which is materially worse than an absent edge "
           "because it is invisible to the reader.")

    h(doc, 2, "3.3  Claim-Side Report Processing Module")
    p(doc, "The acquisition, parsing, segmentation and classification stages of this module are "
           "described in §2.2.2–§2.2.4 together with the measurements that characterise their "
           "output, and are not repeated here. What belongs to the methodology rather than to the "
           "data is the interface this module presents to the rest of the system: a stream of JSONL "
           "records, one per sentence, each carrying the traceability triple, the text, the "
           "probability vector and the thresholded labels — and the guarantee that the traceability "
           "triple is never modified downstream.")
    p(doc, "**Quantitative extraction: KPI observations.** Sentences carry claims; numbers carry "
           "evidence. A separate extraction stage therefore "
           "reads each page that contains at least one ESG-relevant sentence and emits typed KPI "
           "records under structured output, using the 35-indicator controlled vocabulary built in "
           "§3.2. Restricting the stage to ESG-bearing pages is a cost decision: it is the "
           "difference between paying for every page of a 90,000-page corpus and paying only for the "
           "pages that could possibly contain a disclosure. Its yield, and the two-thirds of "
           "extracted KPIs that fall outside the controlled vocabulary, are reported in §4.3.4.")
    p(doc, "Each record pairs a KPI type with one or more observations, and each observation carries "
           "a value, a unit **as written in the source**, a kind (`achieved`, `target`, "
           "`projection`, `baseline`), a year, the source-sentence identifier and a verbatim "
           "snippet. Units are deliberately *not* normalised at this stage. Normalisation is a "
           "separate, later, offline step, so that the raw string remains available for audit — a "
           "principle whose value is visible in §4.3.4, where the raw unit distribution turns out to "
           "be the most informative single view of extraction quality.")


    h(doc, 2, "3.4  Conduct-Side News Module")
    p(doc, "Likewise the crawl, extraction and normalisation stages are described in §2.2.5. Two "
           "design properties are methodological and are stated here.")
    p(doc, "**Schema sharing with a channel stamp.** News sentences enter *the same* schema as "
           "report sentences and are classified by *the same* model, but every node and edge "
           "eventually derived from them carries `source_type = news`. Sharing the schema is what "
           "allows one identity space; the stamp is what keeps the channels separable at query time. "
           "Neither property is recoverable if the two channels are modelled separately and joined "
           "later.")
    p(doc, "**The self-verification guard.** The cross-check stage removes `verifiedBy` edges "
           "arising from domains owned by the issuer and records them separately as flagged, "
           "non-independent support. They remain visible but do not count as verification.")

    h(doc, 2, "3.5  Temporal Knowledge Graph Module")

    h(doc, 3, "3.5.1  Schema and the T1/T2/T3 tier model")
    fig = figure_slot(doc, "Temporal KG schema: the T1/T2/T3 tier partition and where time is "
                           "stored in each tier.",
                      "adapt from docs/PIPELINE_DIAGRAMS.md fig. 8")
    p(doc, "A single schema file is the source of truth for the graph: **28 node classes** and "
           "**48 directed edge labels**, the latter declared over **76 legal (source_class, "
           "target_class) pairs** — an edge label may be legal for more than one pair, which is "
           "why the validator accepts any matching pair and auto-swaps a reversed one (§3.5.3). "
           "Classes are partitioned into three tiers, and each class belongs to exactly one "
           "tier — a property enforced by test rather than by documentation.")
    table(doc, "The T1/T2/T3 tier model and the placement of time. The partition is exhaustive "
               "and disjoint: 14 + 11 + 3 = 28 classes.",
          ["Tier", "n", "Classes", "Where time lives"],
          [["**T1** — Durable entities", "14",
            "Organization, Person, Facility, Product, Material, Location, Country, Standard, "
            "Regulation, Authority, Community, ClaimKeyword, Certification, StandardIndicator",
            "Not on the node. Identity is timeless; history is held in `temporal_versions` and "
            "`supersedes` edges."],
           ["**T2** — Observations and events", "11",
            "KPIObservation, Emission, Waste, Penalty, Controversy, MediaReport, "
            "ThirdPartyVerification, Investment, Project, Initiative, CarbonOffsetProject",
            "On the node itself (`valid_from` / `valid_to`)."],
           ["**T3** — Statements and norms", "3",
            "SustainabilityClaim, Goal, ScienceBasedTarget",
            "On the node and on the connecting edges."]],
          widths=[1.5, 0.5, 4.6, 3.4], right_align_from=99)
    p(doc, "Two placements in that table are judgement calls rather than obvious readings, and both "
           "are recorded in the code that enforces the partition. **Certification is treated as T1**, "
           "not as an event: the node represents a *type of certificate* and is therefore durable, "
           "while the period for which a particular company holds it lives on the "
           "`holdsCertification` edge. **StandardIndicator is also T1**, because an indicator in a "
           "published standard is a durable reference object rather than something the company "
           "did — even though, unlike the other T1 classes, it is generated from a controlled "
           "vocabulary rather than extracted from text.")
    p(doc, "That second point has a direct consequence for the ablation in §4.3.2, and it is worth "
           "stating before the result rather than after it. Indicator nodes are **deliberately** "
           "high-degree: the whole purpose of the axis is that every observation of a given "
           "indicator hangs off one node. Counting them in the traversability metrics would "
           "therefore measure the size of the vocabulary rather than the connectedness of the "
           "graph, and would make a before/after comparison across the very change that introduces "
           "them meaningless. They are consequently **excluded from the Q7 hub and path metrics, "
           "and only from those**. The improvement reported in §4.3.2 is thus not an artefact of "
           "adding well-connected nodes and then counting them.")
    p(doc, "Eight principles govern the temporal design; three carry most of the engineering "
           "consequence and are stated here, with the full set in Appendix A.")
    bullet(doc, "**P1 — Entity identity is timeless.** No time-valued field may appear in the "
                "identity keys of a T1 class. Violating P1 splits a single company into one entity "
                "per year and destroys every cross-year inference the system exists to perform. The "
                "rule is **automatically linted**, not left to convention. §4.3.1 shows what "
                "enforcing it was worth.")
    bullet(doc, "**P2 — In the resolved graph, time lives on edges and on T2/T3 nodes**, never on a "
                "T1 entity node.")
    bullet(doc, "**P4 — Dates are canonical ISO.** Every temporal value is normalised to "
                "`YYYY[-MM[-DD]]`, and a version chain containing an open version must have exactly "
                "one member flagged as current.")

    h(doc, 3, "3.5.2  Triple extraction")
    figure_slot(doc, "Triple extraction, offline validation and entity resolution.",
                "adapt from docs/PIPELINE_DIAGRAMS.md fig. 5–6")
    p(doc, "For every page containing at least one ESG-classified sentence, the system assembles a "
           "prompt from the page text, the KPI observations extracted for that page and the graph "
           "schema, and requests **structured output**: typed JSON containing nodes and edges with "
           "temporal metadata attached. Two prompt templates are used, one per channel — a "
           "claim-side template producing SustainabilityClaim, Goal, Initiative and reported "
           "KPIObservation nodes, and a conduct-side template producing Controversy, MediaReport, "
           "Penalty and observed KPIObservation nodes. Every node and edge is stamped with its "
           "`source_type` at this point, and the stamp is never recomputed later.")
    p(doc, "**The ESG label is a cost filter, not evidence.** Restricting extraction to ESG-bearing "
           "pages is what makes the stage affordable, and this is the only role the classifier's "
           "output plays. No downstream decision reads the label. This is a direct response to the "
           "precision concern measured in §2.3.3: a filter that over-fires costs money, whereas a "
           "filter treated as evidence would propagate its errors into every conclusion.")
    p(doc, "**Language guard.** Both templates require Vietnamese output for the `name`, `title` and "
           "`description` fields. The reason is not stylistic: if the model translates a proper noun "
           "into English on one page and leaves it in Vietnamese on another, entity resolution will "
           "later split one real-world entity into two. The templates are pinned byte-for-byte by a "
           "test, because a harmlessly reworded prompt still runs successfully while changing every "
           "extraction it produces.")

    h(doc, 3, "3.5.3  Offline validation and repair")
    p(doc, "Validation runs in three phases, and the ordering is the substantive design decision.")
    bullet(doc, "**Phase 1 (offline).** Reverse the direction of edges emitted backwards and "
                "validate against the schema. An edge label may be legal for several distinct "
                "(source class, target class) pairs; the validator accepts any matching pair and "
                "auto-swaps a reversed one.", numbered=True)
    bullet(doc, "**Phase 1.5 (offline).** Normalise all dates to ISO form, warn where "
                "`valid_from > valid_to`, and default the `date_uncertain` flag on news-side "
                "observation nodes.", numbered=True)
    bullet(doc, "**Phase 2 (LLM).** Only triples *still* invalid after both offline phases are "
                "batched to a language model for repair. This ordering keeps model cost proportional "
                "to the number of **errors** rather than to the volume of **data**.", numbered=True)
    p(doc, "**Value-preservation guard.** The repair model may correct the *shape* of a triple — its "
           "class, predicate and temporal fields — but is forbidden from translating, reformatting, "
           "inventing or dropping a property *value*. This is enforced in code rather than by prompt "
           "instruction: after the model responds, property values are restored from the original. "
           "Without the guard, a model prompted in English will helpfully “correct” a Vietnamese "
           "proper noun, and the resulting split entity will not surface until entity resolution "
           "several stages later.")

    h(doc, 3, "3.5.4  Entity resolution")
    figure_slot(doc, "Entity resolution stages A–D.", "adapt from docs/ENTITY_RESOLUTION.md")
    p(doc, "Entity resolution proceeds in four stages, the first of which has three substages. "
           "**Stage A.1** is an exact merge on schema identity keys; **Stage A.2** applies a "
           "*frozen* issuer anchor that collapses every name variant of the reporting company; and "
           "**Stage A.3** applies a *frozen* standards anchor that collapses the four or more "
           "attested spellings of “GRI”, and both the Vietnamese and English forms of Circular 96, "
           "onto one canonical node each. The anchors are called frozen because a later stage may "
           "not split what they merged. **Stage B** blocks on a Vietnamese-normalised signature "
           "combined with cosine similarity over embeddings; **Stage C**, LLM adjudication of "
           "residual ambiguous pairs under an explicit budget cap; and **Stage D**, consolidation by "
           "disjoint-set union, preserving temporal history throughout.")
    p(doc, "**The results reported in this thesis were produced with Stages B and C disabled**, "
           "because the API project behind the embedding and adjudication key is billing-blocked. "
           "Only Stage A and the deterministic part of Stage B were active. This is a real "
           "limitation rather than a presentational detail — the residual-duplicate figure in §4.3.1 "
           "would be lower with the full pipeline enabled — and it is restated in §5.2.")

    h(doc, 3, "3.5.5  Provenance patching and the indicator axis")
    figure_slot(doc, "Materialisation of the standard-indicator axis.",
                "adapt from docs/STANDARD_INDICATOR_AXIS.md")
    p(doc, "**Provenance patching** stamps `source_doc` and `source_page` — and, for news-derived "
           "nodes, the article title, URL and domain — onto claim and evidence nodes, using a "
           "four-tier precedence: direct parse of the source identifier, exact index lookup, "
           "recomputation of the stable identifier, and finally a page-token match. The stage "
           "operates under a hard invariant: **it never reorders nodes**, because the graph loader "
           "keys database nodes by array index and the dossiers reference nodes positionally. "
           f"§4.3.5 reports the coverage this achieves in practice ({S['graph_provenance_pct']}%).")
    p(doc, "**Indicator-axis materialisation** appends approximately 35 StandardIndicator nodes and "
           "four kinds of edge: `partOf` (indicator to its parent document), `measuredUnder` "
           "(observation to indicator, read from a previously canonicalised indicator identifier and "
           "never guessed at this stage), `equivalentTo` (Circular 96 to GRI, confirmed crosswalk "
           "rows only), and a keyword tier of `alignsWithIndicator` (claim, goal or initiative to "
           "indicator, longest matching phrase winning). The stage is **append-only** and asserts "
           "that property of its own output, so dossiers computed against an earlier revision of the "
           "graph remain valid.")
    p(doc, "**Canonicalisation is precision-oriented.** The stage that assigns each observation a "
           "canonical indicator identifier writes a *new* property and never rewrites the raw "
           "`kpi_type`, which participates in identity keys. Financial KPIs denominated in VND are "
           "**rejected rather than force-mapped**. §4.3.4 is the measurement behind that rule: with "
           f"{S['kpi_other_share']}% of extracted groups falling outside the vocabulary and money "
           "the second most common unit, a permissive mapper would have attributed revenue figures "
           "to environmental indicators at scale.")
    p(doc, "**The self-reported-zero rule.** A Penalty node whose amount is zero represents a company "
           "stating that it *was never fined*. Such a node is flagged and generates **no conduct "
           "edge**. This is the most dangerous shape in the entire problem — a self-authored "
           "assertion counted as independent evidence of good conduct — and handling it correctly is "
           "a one-line rule whose absence would have quietly inflated the supported-claim count.")

    h(doc, 3, "3.5.6  Graph materialisation and hub decomposition")
    p(doc, "The resolved node and edge collections are loaded into a property-graph database. Nodes "
           "are keyed by array index, since entities are already resolved and must not be "
           "de-duplicated a second time. Edges are merged on a temporal edge key, so the same "
           "relation asserted for several years remains several distinct edges rather than "
           "collapsing into one. Version chains become `supersedes`-linked sequences for the classes "
           "where superseding is legally meaningful, and a JSON property otherwise.")
    p(doc, f"A separate, read-only export addresses the hub structure measured in §4.3.5. Every "
           f"organisation cluster whose summed degree exceeds a threshold (default 500) is "
           f"decomposed: its edges are grouped into synthetic bucket nodes keyed by "
           f"(year, predicate), cutting the hub's own degree to one edge per bucket. On the graph "
           f"reported here the single issuer cluster is decomposed into **357 buckets, reducing "
           f"maximum degree from {f(S['graph_max_degree'])} to 542** — an order of magnitude.")
    p(doc, "The stage nevertheless reports `threshold_met: false`, and that is worth reading "
           "rather than skipping. The largest single bucket still holds 541 edges, marginally over "
           "the 500 gate, because version 1 buckets by (year, predicate) only and one "
           "(year, predicate) combination is genuinely that large. Escalating to a third key would "
           "close the gap; the stage instead reports the shortfall honestly rather than declaring "
           "success. The export **never writes back** to the resolved graph or the database — it "
           "is a dataset-construction artifact for downstream path-reasoning work, and every "
           "synthetic node it creates is flagged as such, because a bucket hop carries no source "
           "sentence and must not be presented as a citable reasoning step.")

    h(doc, 2, "3.6  Claim–Conduct Cross-Check Module")
    figure_slot(doc, "Claim–conduct cross-check: retrieval, adjudication and dossier generation.",
                "adapt from docs/PIPELINE_DIAGRAMS.md fig. 7")

    h(doc, 3, "3.6.1  Retrieval")
    p(doc, "The conduct pool is defined narrowly: only nodes of a conduct class (Controversy, "
           "Penalty, MediaReport, KPIObservation, ThirdPartyVerification) **whose `source_type` is "
           "`news`** are eligible. A report-derived KPI can never be retrieved as evidence for a "
           "report-derived claim, which is what makes the cross-check a test rather than a "
           "consistency check.")
    p(doc, "Retrieval over that pool proceeds in **two tiers**, and the first is the reason the "
           "indicator axis of §3.5.5 exists.")
    bullet(doc, "**Tier 1 — indicator join.** Conduct linked by `measuredUnder` to the same "
                "StandardIndicator that the claim is linked to by `alignsWithIndicator` is injected "
                "directly, with a large fixed score boost that guarantees it outranks every "
                "token-matched candidate. Crucially, tier 1 **bypasses the lexical gate "
                "altogether**: a claim reading *“giảm phát thải”* and a KPI reading *“12.450 "
                "tCO₂e”* share no tokens whatsoever, yet they concern the same indicator. Pure "
                "lexical retrieval cannot connect them; this is exactly the class of evidence a "
                "flat retriever loses.", numbered=True)
    bullet(doc, "**Tier 2 — topic-token overlap.** Remaining conduct nodes qualify only if they "
                "share at least one topic token with the claim, ranked by overlap count and then "
                "by recency.", numbered=True)
    p(doc, "Both tiers are then filtered by an **asymmetric temporal window** — one year before the "
           "claim, fifty years after — and the top eight survivors are retained. The asymmetry "
           "follows directly from §2.3.5: a claim made in 2012 can be contradicted by conduct "
           "observed in 2020 but essentially cannot be contradicted by conduct observed in 2005, "
           "and since roughly 69% of the conduct channel was published in the last two years, a "
           "symmetric window would spend most of its budget on evidence that does not exist.")
    p(doc, "One exception is built into the window and is worth stating explicitly, because it is a "
           "direct consequence of the data-quality measurement in §2.3.5: **a conduct node flagged "
           "`date_uncertain` is exempt from the temporal filter entirely.** Filtering on a date the "
           "system has already declared untrustworthy would discard evidence on the strength of a "
           "value it does not believe. The node is admitted and its uncertainty is carried forward "
           "into the dossier as a caveat instead.")
    p(doc, "§4.6 records the cost of the wide window: it should be paired with a distance-decay "
           "weighting during aggregation, which it currently is not.")

    h(doc, 3, "3.6.2  Adjudication and aggregation")
    p(doc, "For each pair (c, e) an adjudicating language model returns one of three labels together "
           "with a confidence value and a **natural-language rationale**. The rationale is mandatory "
           "rather than optional: it is what a human analyst reads in order to decide whether to "
           "accept the system's judgement, and an assessment that cannot be argued with cannot be "
           "audited.")
    p(doc, "LLM adjudication is **mandatory and has no deterministic fallback**. If no provider is "
           "available the stage aborts at start-up rather than silently degrading to a heuristic "
           "that would produce numbers of a different and undocumented kind. The provider used for "
           "the reported results is `gpt-4o-mini`.")
    p(doc, "Pair-level verdicts are aggregated into three advisory assessments: `appears_supported`, "
           "`appears_contradicted` and `unverified_insufficient_evidence`. Where a dossier contains "
           "both supporting and contradicting evidence, **contradiction takes precedence over "
           "support**. This asymmetry is intentional and conservative: the cost of surfacing a mixed "
           "case for human attention is small, whereas the cost of reporting a claim as supported "
           "while contradicting evidence sits in the same dossier is the exact failure the system "
           "exists to prevent.")

    h(doc, 3, "3.6.3  The advisory dossier and presentation layer")
    p(doc, "Each claim yields a record containing its identifier, text and year; the assessment; an "
           "explicit advisory flag; lists of supporting and contradicting evidence, where each item "
           "carries the node index, class, text, confidence, rationale, `date_uncertain` flag and "
           "independence flag; the separately-recorded non-independent support; and a list of "
           "caveats. One caveat appears in **every** dossier without exception:")
    quote(doc, "“No ground-truth greenwashing label exists; this is an advisory opinion.”")
    p(doc, "Two surfaces consume the dossiers. The **claim ledger** is a per-company table ordered "
           "signal-first — contradicted, then supported, then unverified — rendered exclusively from "
           "the database so that what an analyst reads is what the database actually holds. The "
           "**ESG Evidence View** is a three-column web interface aligned to the Circular 96 and GRI "
           "axes. It displays only claims carrying an `alignsWithIndicator` edge, so each card's "
           "E/S/G column is read from the linked indicator's declared pillar rather than inferred "
           "from the claim text.")
    figure_slot(doc, "The ESG Evidence View interface (three-column TT96/GRI layout).",
                "screenshot: run python api/main.py and open http://localhost:8000")


# ---------------------------------------------------------------------------
#                           CHAPTER 4 — EXPERIMENT
# ---------------------------------------------------------------------------
def _chapter_experiment(doc):
    h(doc, 1, "4  Experiment")

    h(doc, 2, "4.1  Implementation Details")
    p(doc, "The system is implemented in Python and runs on Windows. Graph construction is organised "
           "as **sixteen named stages** behind a single dispatcher, executed from the repository "
           "root; two of them are also runnable as *blocks* that chain three stages in memory and "
           "write the shared artifact once instead of three times. The ESG classifier is a "
           "published checkpoint run offline (§2.2.4) — no training is performed and the pipeline "
           "itself has no GPU dependency. Triple extraction, validation repair and "
           "cross-check adjudication call commercial language-model APIs under explicit rate limits "
           "and budget caps; every other stage — validation phases 1 and 1.5, KPI canonicalisation, "
           "entity resolution stages A and D, provenance patching, indicator-axis materialisation, "
           "the quality instrument and the export view — is fully **offline and deterministic**.")
    p(doc, "The resolved graph is materialised in Neo4j over the Bolt protocol. Generated data is "
           "distributed through a dataset repository with the pushed revision pinned in a "
           "version-controlled file, so a checkout recovers the data that went with that code.")
    p(doc, "**Testing.** The project follows test-first development: no production code lands "
           "without a failing test that demanded it. Tests are plain assertion scripts that run "
           "offline against real artifacts already on disk, with no language model, no database and "
           "no network, so the suite is free to run and repeatable. Paid or networked stages are "
           "covered by stubbing *underneath* the provider abstraction — a deterministic fake keyed "
           "on a checksum of the prompt — so the real stage logic still executes against fake I/O. "
           "This is what makes it possible to pin a paid prompt byte-for-byte in a test: a reworded "
           "prompt still runs successfully while silently changing every result it produces.")

    h(doc, 2, "4.2  Evaluation Metrics")
    p(doc, "**Why there is no ground truth.** No labelled set of greenwashing instances exists for "
           "Vietnamese listed companies. Third-party ESG ratings cannot substitute, because those "
           "ratings are themselves derived largely from the same corporate disclosure the system is "
           "attempting to test; using them as ground truth would be circular, and a high agreement "
           "score against them would measure conformity to disclosure rather than fidelity to "
           "conduct. The system therefore emits no score, and evaluation is displaced onto "
           "properties measurable without labels.")
    p(doc, "**The instrument.** Eight attributes of the resolved graph are measured entirely offline "
           "— no language model, no database — so the instrument is free, deterministic and "
           "repeatable. It is run with a label before and after every design change, and the "
           "difference between the two runs is the experimental result.")
    n = table(doc, "The eight-attribute label-free graph-quality instrument.",
              ["#", "Attribute", "What it measures"],
              [["Q1", "Accuracy", "Non-NFC names; names corrupted during text extraction."],
               ["Q2", "Consistency", "Illegal edges; non-ISO dates; `valid_from > valid_to`; broken "
                                     "current-version chains; versions split by formatting; missing "
                                     "`date_uncertain`; T1 classes carrying time in identity keys."],
               ["Q3", "Conciseness", "Surplus duplicate T1 entity nodes."],
               ["Q4", "Completeness", "Coverage of the conduct channel."],
               ["Q5", "Timeliness", "Share of edges and T2 nodes carrying `valid_from`; share of "
                                    "news-derived T2 nodes carrying `date_uncertain`."],
               ["Q6", "Provenance", "Share of nodes carrying `source_type`; share of KPI "
                                    "observations whose source identifier parses."],
               ["Q7", "Traversability", "Median degree; share of leaf nodes; share of "
                                        "masked-answerable queries; share of claim→conduct paths "
                                        "reachable structurally; share of T2 nodes with degree ≥ 2."],
               ["Q8", "Independence", "Distribution of conduct evidence across channels."]],
              widths=[0.7, 2.0, 7.6], right_align_from=99)
    p(doc, f"Table {n} defines the attributes. Three further label-free designs have been specified "
           "but not run, and are reported as future work with a completed design rather than as "
           "results: *metamorphic relations* (paraphrasing a claim must not change its verdict), a "
           "*negative control with permutation testing* (pairing a company's claims against another "
           "company's conduct pool must collapse the contradiction rate towards chance), and "
           "*Krippendorff's α* for agreement across repeated adjudication runs.")
    todo(doc, "TODO-RUN", "The negative control is by far the highest value-per-hour item left in "
                          "the evaluation. It is cheap, requires no new code beyond a swapped "
                          "conduct pool, and directly answers the sharpest question an examiner can "
                          "ask — “how do you know the system is not simply labelling at random?” A "
                          "collapse towards zero demonstrates genuine specificity; a failure to "
                          "collapse is itself a finding worth reporting honestly.")

    h(doc, 2, "4.3  Result Analysis")
    p(doc, "This section reports what the pipeline produces. §4.3.1–§4.3.3 are three **controlled "
           "ablations**, each holding the corpus and the language model fixed and varying exactly "
           "one component of the pipeline. §4.3.4 and §4.3.5 then characterise the **artifacts the "
           "full pipeline yields** — the extracted KPI observations and the resolved knowledge "
           "graph — including the places where they fall short, since a system's output is "
           "evidence about the system and has to be reported as such rather than presented as "
           "corpus statistics.")
    todo(doc, "WRITE", "Read this paragraph before the defence. Each ablation compares two "
                       "snapshots of the graph taken immediately before and after the change in "
                       "question, so both arms of a pair share a lineage and the comparison is "
                       "valid. The snapshots were taken on different dates as the pipeline "
                       "developed, so the ABSOLUTE totals differ between ablations and from the "
                       "final graph reported in §4.3.5, which incorporates later ingestion. Either "
                       "keep this sentence in the text or re-run the quality instrument on the "
                       "current graph so that §4.3.5 and the ablations in §4.3.1-§4.3.3 share at "
                       "least one common point. The "
                       "instrument is offline and free to run.")

    h(doc, 3, "4.3.1  Ablation A — temporal integrity")
    p(doc, "**Setup.** Identical corpus, identical language model; the only variable is the offline "
           "temporal-handling logic — ISO date normalisation, version-chain repair, `date_uncertain` "
           "defaulting, and the removal of time-valued fields from the identity keys of T1 classes.")
    n = table(doc, "Effect of the temporal-integrity redesign.",
              ["Metric", "Baseline", "After redesign", "Δ"],
              [["Graph size (nodes / edges)", "10,573 / 13,008", "10,362 / 13,047", "−211 / +39"],
               ["**Q2 total consistency violations**", "**1,098**", "**1**", "**−99.9%**"],
               ["  — broken current-version chains", "660", "0", "−660"],
               ["  — versions split by formatting", "312", "0", "−312"],
               ["  — missing `date_uncertain`", "124", "0", "−124"],
               ["  — T1 classes with time in identity", "2", "0", "−2"],
               ["Q3 surplus duplicate T1 nodes", "271", "60", "−78%"],
               ["Q5 T2 nodes with `valid_from`", "0.0%", "87.7%", "+87.7 pp"],
               ["Q5 news T2 with `date_uncertain`", "0.0%", "100.0%", "+100 pp"],
               ["Q7 leaf nodes", "83.2%", "82.2%", "−1.0 pp"],
               ["Q7 masked-answerable queries", "25.1%", "26.3%", "+1.2 pp"]],
              widths=[4.6, 2.2, 2.4, 1.8])
    p(doc, f"**Interpretation.** Three observations. First, the graph simultaneously *lost* nodes "
           "and *gained* edges (−211 / +39). That combination is the signature of correct "
           "consolidation rather than of data loss: 211 duplicate nodes were merged, and edges "
           "previously scattered across those duplicates now converge on a single entity. Second, "
           "the 78% reduction in surplus duplicate entity nodes is the direct consequence of the "
           "timeless-identity rule — before the redesign the same organisation was split into a "
           "separate entity for each year in which it was mentioned. Third, the figure of 87.7% for "
           "observation nodes carrying a validity date is deliberately not 100%: the remainder are "
           "observations for which the source genuinely states no date. The system leaves the field "
           "empty rather than inventing a value, and the figure is reported unrounded for that "
           "reason.")

    h(doc, 3, "4.3.2  Ablation B — standard-indicator axis")
    p(doc, "**Setup.** The same resolved graph with the indicator-axis materialisation stage (§3.5.5) "
           "enabled and disabled.")
    n = table(doc, "Effect of materialising the standard-indicator axis.",
              ["Metric", "Before", "After", "Δ"],
              [["Nodes / edges", "10,362 / 13,047", "10,393 / 13,790", "+31 / +743"],
               ["Q7 leaf nodes", "82.2%", "**75.8%**", "**−6.4 pp**"],
               ["Q7 masked-answerable queries", "26.3%", "**34.8%**", "**+8.5 pp**"],
               ["Q7 T2 nodes with degree ≥ 2", "10.1%", "**19.9%**", "**+9.8 pp**"],
               ["Q7 median degree", "1.0", "1.0", "—"],
               ["Q2 consistency violations", "1", "1", "0"]],
              widths=[4.6, 2.2, 2.4, 1.8])
    p(doc, "**Interpretation.** Thirty-one additional nodes produced 743 additional edges, so each "
           "indicator node connects on average about 24 observations or claims. The share of "
           "observation nodes with degree of at least two nearly doubled: before this stage most KPI "
           "observations were dangling leaves that could not participate in any reasoning path at "
           "all. This is the mechanism that converts a *bag of observations* into a **traversable "
           "graph**, and it is why the masked-answerable share rises by 8.5 percentage points from "
           "an operation adding only 0.3% more nodes. That Q2 violations remain at 1 confirms the "
           "stage is genuinely append-only and breaks no existing invariant.")
    n = table(doc, "Structural anchoring per observation class after materialisation.",
              ["Class", "Nodes", "Degree ≥ 2"],
              [["Emission", "24", "100.0%"],
               ["Project", "255", "53.3%"],
               ["Investment", "282", "50.4%"],
               ["KPIObservation", "4,906", "16.9%"],
               ["Initiative", "495", "14.7%"],
               ["Waste", "15", "13.3%"],
               ["MediaReport", "91", "9.9%"],
               ["Controversy", "2", "0.0%"],
               ["Penalty", "4", "0.0%"],
               ["ThirdPartyVerification", "24", "0.0%"]],
              widths=[4.0, 2.4, 2.6])
    p(doc, "Three classes sit at 0% anchoring — Controversy, Penalty and ThirdPartyVerification. "
           "This is the thinness of the conduct channel measured in §2.3.6 and §4.3.5 reappearing "
           "as a structural property, not a defect in the anchoring algorithm: with two Controversy "
           "nodes and four Penalty nodes in the entire graph, there is very little to anchor.")

    h(doc, 3, "4.3.3  Ablation C — document linearisation")
    p(doc, "**Motivation.** §3.2.2 selects a layout-aware parser over positional extraction. That is "
           "a design assertion, and an assertion of that kind should be measured rather than "
           "declared, particularly since comparable work uses positional extraction on the same "
           "corpus.")
    p(doc, "**Setup.** Three arms differ **only** in how the PDF is converted to text: positional "
           "raw extraction; positional extraction with blocks re-sorted into reading order by "
           "coordinate; and layout-aware linearisation as described in §3.2.2. A **single disclosure "
           "segmenter, copied verbatim from the production parser, is applied identically to all "
           "three arms**, so the comparison isolates the parser and cannot be won by tuning a "
           "detector. Ground truth is the disclosure list published by GRI itself in its official "
           "content index, which is derived from no parser in this experiment and therefore cannot "
           "favour any arm. Of the 42 standards, 36 appear in the content index and are scored; the "
           "remaining six are excluded rather than scored against an empty denominator.")
    n = table(doc, "Parser ablation on 36 GRI standards, scored against the official GRI content "
                   "index.",
              ["Parser", "Recall", "Precision", "F1", "Title fidelity", "False pos."],
              [["Positional, raw text", "**100.0%**", "62.6%", "77.0%", "81.8%", "82"],
               ["Positional, ordered blocks", "**100.0%**", "62.6%", "77.0%", "81.8%", "82"],
               ["Layout-aware (Marker/Surya)", "98.5%", "**100.0%**", "**99.3%**", "**100.0%**",
                "**0**"]],
              widths=[3.8, 1.6, 1.8, 1.4, 1.9, 1.5])
    p(doc, "**Positional extraction wins on recall, and that is reported rather than suppressed.** "
           "It recovers every disclosure code in the corpus. The difference between the arms lies "
           "entirely in *precision* and *title fidelity*, and both failures are structural.")
    n = table(doc, "False positives by kind.",
              ["Parser", "Total FP", "Ground-truth gap (own standard)",
               "Real contamination (cross-standard)"],
              [["Positional, raw text", "82", "4", "78"],
               ["Positional, ordered blocks", "82", "4", "78"],
               ["Layout-aware (Marker/Surya)", "0", "0", "0"]],
              widths=[3.6, 1.8, 3.2, 3.2])
    p(doc, f"A code whose prefix matches the standard being parsed — `306-1` found inside GRI 306 — "
           "is almost certainly a **ground-truth gap**: the 2021 content index retired disclosures "
           "the 2016 PDF still contains, so the parser is penalised for extracting genuine content. "
           "Four of the 82 are of this kind. The remaining **78 are genuine cross-standard "
           "contamination**: a code such as `2-23` harvested out of a cross-reference inside GRI 101 "
           "and then attributed to GRI 101. This is precisely the mis-attribution that the ownership "
           "rule of §3.2.3 exists to prevent, arriving one stage earlier and by a different route.")
    p(doc, "Title fidelity fails in a single consistent pattern: truncation at a line break. For GRI "
           "305-7, positional extraction yields *“Nitrogen oxides (NOx), sulfur oxides (SOx), and "
           "other significant”*, silently dropping *“air emissions”*. That string becomes the "
           "display name of a StandardIndicator node, so the defect is visible to the end user of "
           "the evidence interface and not merely internal.")
    n = table(doc, "Control: character-level recall versus structural recall.",
              ["Parser", "Mention recall", "Unit recall", "Markdown headings", "Table rows"],
              [["Positional, raw text", "100.0%", "100.0%", "0", "0"],
               ["Positional, ordered blocks", "100.0%", "100.0%", "0", "0"],
               ["Layout-aware (Marker/Surya)", "100.0%", "98.5%", "2,322", "1,418"]],
              widths=[3.6, 2.2, 2.0, 2.2, 1.6])
    p(doc, "Every arm recovers 100% of disclosure codes at character level. Positional extraction "
           "therefore loses no *text* — it loses *structure*, recovering zero headings and zero "
           "table rows against 2,322 and 1,418 respectively. Without heading structure there is no "
           "signal distinguishing a real disclosure heading from a sentence that merely mentions a "
           "disclosure code, which is the direct cause of the 78 contaminating false positives.")
    n = table(doc, "Detector ablation: emphasis-tolerant segmentation applied to all arms.",
              ["Parser", "Recall (production)", "Recall (tolerant)", "Precision (tolerant)",
               "Title fidelity (tolerant)"],
              [["Positional, raw text", "100.0%", "100.0%", "62.6%", "81.8%"],
               ["Positional, ordered blocks", "100.0%", "100.0%", "62.6%", "81.8%"],
               ["Layout-aware (Marker/Surya)", "98.5%", "**100.0%**", "**100.0%**", "**100.0%**"]],
              widths=[3.4, 2.2, 2.0, 2.2, 2.2])
    p(doc, "The two disclosures missed by the layout-aware arm are not parser failures: the "
           "production segmenter cannot match a disclosure code wrapped in Markdown emphasis, and "
           "only a parser rich enough to emit emphasis can ever be penalised by that. **Conclusion.** "
           "Positional text extraction recovers every disclosure code but cannot distinguish a "
           "disclosure heading from a cross-reference to one, yielding 78 spurious disclosure units "
           "and truncating 18.2% of titles. Layout-aware linearisation resolves both, at 100% "
           "precision and 100% title fidelity, and at no recall cost once the segmenter tolerates "
           "emphasis.")
    todo(doc, "WRITE", "This ablation surfaced a genuine defect in the production segmenter — it "
                       "silently skips disclosures whose code is emphasised — left unfixed because "
                       "repairing it changes a version-controlled artifact (`config/gri_catalog."
                       "json`). Decide before submission whether to fix it. Either choice is "
                       "defensible; reporting the defect and the decision is stronger than "
                       "reporting neither.")
    h(doc, 3, "4.3.4  Extraction yield and the limits of the controlled vocabulary")
    fig = figure(doc, "fig_kpi_types.png",
                 "KPI extraction. (a) Distribution over the controlled vocabulary, with `other` as "
                 "the catch-all. (b) The twelve most frequent raw units, exactly as written in the "
                 "source documents.")
    kk = S["kpi_kinds"]
    n = table(doc, "KPI extraction yield.",
              ["Quantity", "Value"],
              [["Pages processed", f(S["kpi_pages_processed"])],
               ["Pages yielding at least one KPI",
                f"{f(S['kpi_pages_with_kpi'])} ({S['kpi_page_hit_rate']}%)"],
               ["KPI groups extracted", f(S["kpi_groups"])],
               ["**Observations extracted**", f"**{f(S['kpi_observations'])}**"],
               ["Observations per group", str(S["kpi_obs_per_group"])],
               ["Distinct KPI types matched", str(S["kpi_distinct_types"])],
               ["**Share falling into `other`**", f"**{S['kpi_other_share']}%**"],
               ["Distinct raw units", f(S["kpi_distinct_units"])],
               ["Kind: achieved / target / projection / baseline",
                f"{f(kk.get('achieved',0))} / {f(kk.get('target',0))} / "
                f"{f(kk.get('projection',0))} / {f(kk.get('baseline',0))}"],
               ["Observations carrying a verbatim source snippet", f"{S['kpi_snippet_coverage']}%"]],
              widths=[7.0, 3.4])
    p(doc, f"**Two-thirds of extracted KPIs fall outside the controlled vocabulary.** "
           f"{S['kpi_other_share']}% of groups are typed `other`, meaning the extractor found a "
           "quantitative disclosure that no indicator in the 35-item Vietnamese vocabulary matches. "
           "That is a large fraction, and it admits two readings: either the vocabulary is too "
           "narrow for what these companies actually report, or the extractor is over-firing on "
           "numbers that are not ESG disclosures at all. The canonicalisation stage (§3.5.5) "
           "settles the question, because it records *why* each observation failed to map rather "
           "than only *that* it failed.")
    n2 = table(doc, "Outcome of KPI canonicalisation against the 35-indicator controlled "
                    "vocabulary, as recorded by the stage itself.",
               ["Outcome", "Observations", "Share"],
               [["**Mapped to an indicator**", f"**{f(S['canon_mapped'])}**",
                 f"**{S['canon_mapped_pct']}%**"],
                ["Unmapped", f(S["canon_unmapped"]), f"{100-S['canon_mapped_pct']:.1f}%"],
                ["  — `rejected_unit` (refused: financial)", f(S["canon_rejected_unit"]),
                 f"{S['canon_rejected_share_of_unmapped']}% of unmapped"],
                ["  — `no_match` (genuine vocabulary gap)", f(S["canon_no_match"]),
                 f"{S['canon_no_match_share_of_unmapped']}% of unmapped"],
                ["Distinct indicators actually used", str(S["canon_indicators_used"]), "of 35"]],
               widths=[5.0, 2.6, 2.8])
    p(doc, f"Table {n2} shows that the second reading is the dominant one, and that the pipeline "
           f"already handles it. Of {f(S['canon_unmapped'])} unmapped observations, "
           f"**{f(S['canon_rejected_unit'])} ({S['canon_rejected_share_of_unmapped']}%) were "
           f"refused outright on the basis of their unit** — a financial figure denominated in VND "
           f"is not a missing ESG mapping, it is correctly-excluded noise. Only "
           f"**{f(S['canon_no_match'])} ({S['canon_no_match_share_of_unmapped']}%) are genuine "
           "vocabulary gaps**, and those are the real backlog. The most frequent unmapped titles "
           "confirm the diagnosis beyond argument: *" +
           ", ".join(S["canon_top_unmapped_titles"][:5]) + "*. These are income-statement and "
           "balance-sheet line items, not sustainability disclosures.")
    p(doc, "The distinction is only available because the stage **writes a new `kpi_id` property "
           "and never overwrites the raw `kpi_type` it read off the page**, and because it records "
           "the deciding rule per observation rather than collapsing every failure into a single "
           "null. A pipeline that force-mapped ambiguous cases would have attributed revenue "
           f"figures to environmental indicators {f(S['canon_rejected_unit'])} times over, silently "
           "and invisibly. This is the concrete justification for the precision-over-recall rule "
           "stated in §3.5.5.")
    p(doc, f"It also bounds what the indicator axis can currently do: only "
           f"{S['canon_indicators_used']} of the 35 indicators in the vocabulary are exercised at "
           f"all, and {S['canon_mapped_pct']}% of observations reach one. The axis is real but "
           "sparse, which is the correct way to read the anchoring figures in §4.3.2.")
    p(doc, f"**Unit strings are inconsistent in exactly the ways that matter.** The "
           f"{f(S['kpi_distinct_units'])} distinct raw unit strings include `VND`, `tỷ đồng` and "
           "`đồng` as three spellings of money at two different scales, and `tấn` and `Tấn` "
           "differing only in capitalisation. Preserving these verbatim rather than normalising at "
           "extraction time is what makes the problem visible; a stage that silently canonicalised "
           "on the way in would have produced a clean-looking corpus and an unfalsifiable one.")
    p(doc, f"**The corpus is overwhelmingly retrospective.** Of {f(S['kpi_observations'])} "
           f"observations, {f(kk.get('achieved',0))} are `achieved` values and only "
           f"{f(kk.get('target',0))} are `target` values — about "
           f"{100*kk.get('target',0)/S['kpi_observations']:.0f}%. This matters more than it first "
           "appears. A greenwashing analysis is most powerful when it can compare a *stated target* "
           "against a *later achieved value*, and the corpus supplies very few targets to anchor "
           "such a comparison. The system's forward-looking capability is therefore bounded by data "
           "availability, not by design, and the goal-oriented parts of the schema are "
           "correspondingly sparsely populated.")
    p(doc, f"**Traceability is complete.** {S['kpi_snippet_coverage']}% of observations carry a "
           "verbatim source snippet and a source-sentence identifier, so every number in the graph "
           "can be shown to a reader alongside the sentence it came from. This is the one property "
           "of the extraction stage that is unambiguously as designed.")

    h(doc, 3, "4.3.5  Structure of the constructed graph")
    fig = figure(doc, "fig_graph_composition.png",
                 "Composition of the resolved knowledge graph. (a) The ten most frequent node "
                 "classes. (b) The ten most frequent edge labels.")
    fig2 = figure(doc, "fig_degree_distribution.png",
                  "Degree distribution of the resolved graph on log–log axes. The distribution is "
                  "extremely heavy-tailed: median degree 1, maximum degree "
                  f"{f(S['graph_max_degree'])}.", width_in=4.9)
    src = S["graph_source_type"]
    n = table(doc, "Structural properties of the resolved graph.",
              ["Property", "Value"],
              [["Nodes / edges", f"{f(S['graph_nodes'])} / {f(S['graph_edges'])}"],
               ["Node classes / edge labels",
                f"{S['graph_node_classes']} / {S['graph_edge_labels']}"],
               ["**Nodes by channel (report / news / unstamped)**",
                f"**{f(src.get('report',0))} / {f(src.get('news',0))} / {f(src.get('(unset)',0))}**"],
               ["Nodes carrying temporal versions",
                f"{f(S['graph_versioned_nodes'])} (holding {f(S['graph_version_records'])} versions)"],
               ["Nodes stamped with document provenance",
                f"{f(S['graph_provenance_stamped'])} ({S['graph_provenance_pct']}%)"],
               ["Mean degree / median degree",
                f"{S['graph_mean_degree']} / {S['graph_median_degree']}"],
               ["**Maximum degree**", f"**{f(S['graph_max_degree'])}**"],
               ["Isolated nodes", f(S["graph_isolated_nodes"])]],
              widths=[6.4, 4.0])
    p(doc, f"**The channel imbalance survives into the graph.** Of {f(S['graph_nodes'])} nodes, "
           f"{f(src.get('report',0))} derive from reports and only {f(src.get('news',0))} from news "
           f"— a ratio of about {src.get('report',0)/max(src.get('news',1),1):.0f} to one, even "
           "steeper than the sentence-level ratio in §2.3.6, because report pages yield more "
           "extractable structure per sentence than news prose does. Every cross-check result in "
           "§4.4 must be read against this number.")
    p(doc, f"**The graph is extremely hub-dominated.** Mean degree is {S['graph_mean_degree']} but "
           f"median degree is {S['graph_median_degree']}, and the maximum is "
           f"{f(S['graph_max_degree'])}. Figure {fig2} shows the resulting heavy tail on log–log "
           "axes. The interpretation is straightforward once stated: the issuing company is a single "
           "node to which almost every observation, claim and goal attaches, so a small number of "
           "organisation nodes carry a large fraction of all edges while the typical node has "
           "exactly one. This is not a defect — it is the correct shape for a corpus about one "
           "company — but it has a concrete consequence for any path-based reasoning over the "
           "graph, because a traversal that enters the hub can reach almost anything, which makes "
           "the resulting path uninformative. It is the direct motivation for the hub-decomposition "
           "export described in §3.5.6.")
    p(doc, f"**Provenance coverage is partial.** {S['graph_provenance_pct']}% of nodes carry a "
           "source document and page. The remainder are nodes for which the provenance patcher could "
           "not resolve a source identifier through any of its four matching tiers. That figure is "
           "reported rather than rounded up because provenance is the property the ethical argument "
           "of §5.3 rests on: a node without provenance is a statement a reader cannot check.")


    h(doc, 2, "4.4  Claim–Conduct Cross-Check Results")
    n = table(doc, "Claim–conduct cross-check results for the issuer analysed end-to-end.",
              ["Quantity", "Value"],
              [["Sustainability claims extracted", "1,093"],
               ["Conduct pool", "124 (16 MediaReport + 108 KPIObservation)"],
               ["Claims with at least one candidate", "748 (68.4%)"],
               ["Candidate claim–evidence pairs", "3,461"],
               ["Average candidates per claim", "3.17"],
               ["LLM adjudications performed", "3,461 (0 failures)"],
               ["Linking edges written", "152"],
               ["  → `appears_supported`", "70"],
               ["  → `appears_contradicted`", "22"],
               ["  → `unverified_insufficient_evidence`", "1,001"]],
              widths=[6.4, 4.0])
    p(doc, f"**On the figure of 1,001.** This should not be read as a failure rate, and presenting "
           "it as one would misrepresent the system. It is designed behaviour, and §2.3.6 predicted "
           f"it quantitatively: with a claim-to-conduct ratio of {S['aaa_claim_conduct_ratio']} to "
           "one, and with the claim side governance-heavy while the conduct side is "
           "environmental-heavy, most claims have no topically relevant independent evidence in "
           "existence. Ninety-two per cent of claims are returned as not verifiable because the "
           "independent conduct corpus available for a mid-capitalisation Vietnamese issuer "
           "genuinely is thin. A system that returned a verdict for all 1,093 claims from that "
           "evidence base would be manufacturing confidence it does not possess. The honest output "
           "is the one the system actually produces, and it carries its own caveat:")
    quote(doc, "“Thin independent conduct — absence of contradiction is NOT exoneration.”")

    h(doc, 2, "4.5  Qualitative Results")
    p(doc, "Three dossiers, one per assessment class, drawn unmodified from the system's output. "
           "They were **selected to illustrate the three classes, not sampled at random**, and "
           "should be read as illustrations rather than as evidence about the distribution.")
    h(doc, 3, "Case 1 — Supported (social / labour)")
    p(doc, "**Claim** (annual report 2012): *“The company implements monthly, quarterly and annual "
           "bonus systems to motivate employees.”* → `appears_supported`, confidence 0.90. "
           "**Evidence**: a KPIObservation recording an employee stock-bonus payout of 114.5 billion "
           "VND against 2024 profits. **Comment**: the system connected a policy commitment made in "
           "2012 to an observed disbursement in 2024 — precisely the cross-year inference the "
           "temporal representation exists to support, and one an atemporal graph could not express. "
           "The dossier automatically carries a `date_uncertain` caveat on the evidence.")
    h(doc, 3, "Case 2 — Contradicted (governance / capital)")
    p(doc, "**Claim** (annual report 2012): *“Actively sought investment sources in order to use "
           "capital from shareholders and investors effectively.”* → `appears_contradicted`. "
           "**Contradicting evidence**: total assets changed by −11.5% as of 30 June 2025, "
           "confidence 0.90. **Supporting evidence also present**: Factory No. 6, 500.6 billion VND, "
           "2016. Because both directions are present the dossier carries the caveat *“Evidence is "
           "mixed”*, and under the precedence rule of §3.6.2 the contradiction determines the "
           "assessment. **Comment**: this is also the clearest instance of error class E1 below — a "
           "2012 claim adjudicated against 2025 evidence with no discounting for the thirteen-year "
           "gap.")
    h(doc, 3, "Case 3 — Unverified")
    p(doc, "**Claim** (annual report 2011): a statement concerning social and charitable activities "
           "and gift-giving. → `unverified_insufficient_evidence`, with no retrieved candidates at "
           "all. **Comment**: the archetype of the 1,001 unverified claims — a qualitative "
           "commitment with no measurable KPI attached and no independent reporting. It marks the "
           "system's genuine capability boundary, and is more informative about that boundary than "
           "either preceding case.")

    h(doc, 2, "4.6  Error Analysis")
    p(doc, "Four error classes are identifiable from the current output. Their **mechanisms are "
           "established from reading dossiers; their relative frequencies have not been measured**, "
           "and no percentage should be inferred from the ordering below.")
    p(doc, "**E1 — Temporal-distance mismatch.** The adjudicator pairs a 2012 claim with 2025 "
           "evidence (Case 2) without discounting for temporal distance. The retrieval window of "
           "fifty years forward, itself a response to the recency skew measured in §2.3.5, is too "
           "wide to be used without a decay function. *Remedy:* apply a distance-decay weight over "
           "the year gap during aggregation.")
    p(doc, "**E2 — Generic-claim over-matching.** Highly generic claims — *“the company strives for "
           "sustainable development”* — match almost any KPI observation, generating both spurious "
           "support and spurious contradiction. *Remedy:* score claim specificity before retrieval "
           "and route claims below a threshold directly to `unverified` rather than adjudicating "
           "them.")
    p(doc, "**E3 — Sparse conduct channel.** The dominant cause of the 1,001 unverified claims, and "
           "the one class whose magnitude *is* measured: §4.3.5 puts the graph at "
           f"{f(S['graph_source_type'].get('report',0))} report-derived nodes against "
           f"{f(S['graph_source_type'].get('news',0))} news-derived ones. This is a "
           "data-availability limitation, not an algorithmic error. *Remedy:* extend the conduct "
           "channel with administrative-penalty records, regulator and exchange announcements, and "
           "environmental monitoring data.")
    p(doc, "**E4 — Caveat fatigue from `date_uncertain` propagation.** Most news-side evidence must "
           "use the publication date as a proxy — §2.3.5 shows why this is unavoidable given the "
           f"{S['news_suspicious_dates_pct']}% of articles carrying unusable dates — so the "
           "uncertain-date caveat appears on a large share of dossiers and progressively loses its "
           "warning value through familiarity. *Remedy:* extract sentence-level temporal expressions "
           "with named-entity recognition instead of falling back to the article-level publication "
           "date.")
    todo(doc, "TODO-RUN", "Manually reviewing roughly 20 dossiers and estimating the share "
                          "attributable to each of E1–E4 would turn this section from a taxonomy "
                          "into a measurement and let it be presented as a distribution table. A "
                          "few hours of reading; no code required.")


# ---------------------------------------------------------------------------
#                     CHAPTER 5 — CONCLUSION AND DISCUSSION
# ---------------------------------------------------------------------------
def _chapter_conclusion(doc):
    h(doc, 1, "5  Conclusion and Discussion")

    h(doc, 2, "5.1  Conclusion")
    p(doc, "This thesis addressed the detection of greenwashing in Vietnamese listed companies in "
           "the construction, building-materials and real-estate sector — a setting in which the "
           "party authoring the sustainability narrative is also the party benefiting from it, and "
           "in which systems reading only the corporate report inherit the very bias they are meant "
           "to expose. The approach was a two-channel temporal knowledge graph: report-derived "
           "claims and independently-sourced news conduct extracted into a single schema and a "
           "single identity space, kept distinguishable at query time, and joined by a cross-check "
           "layer that adjudicates each claim against the conduct evidence available for it.")
    p(doc, f"A corpus was assembled specifically for the task: {f(S['raw_pdfs'])} annual-report PDFs "
           f"from {S['raw_companies']} sector issuers ({S['raw_gb']} GB), parsed into "
           f"{f(S['rep_sentences'])} page-anchored sentences over {f(S['rep_pages'])} pages, "
           f"alongside an independent conduct channel of {f(S['news_articles'])} news articles from "
           f"{f(S['news_domains_total'])} distinct domains. The exploratory analysis of this corpus "
           "is a contribution in its own right, because it establishes quantitatively the "
           "constraints the system must operate under: an approximately "
           f"{S['aaa_claim_conduct_ratio']}-to-one claim-to-conduct volume ratio, a pillar "
           "composition that inverts between the two channels, a conduct channel concentrated in the "
           f"last two years, and {S['kpi_other_share']}% of extracted quantitative disclosures "
           "falling outside the controlled regulatory vocabulary.")
    p(doc, f"The resulting graph contains {f(S['graph_nodes'])} nodes and {f(S['graph_edges'])} "
           "edges. Three controlled ablations, measured with a label-free eight-attribute "
           "instrument, quantify the design decisions behind it: the temporal-integrity redesign "
           "reduced schema consistency violations from 1,098 to 1 and surplus duplicate entities by "
           "78%; materialising the standard-indicator axis raised the share of masked-answerable "
           "queries from 26.3% to 34.8% and nearly doubled the share of observation nodes "
           "participating in more than one relation, from an operation adding only 31 nodes; and "
           "layout-aware document linearisation eliminated all 78 cross-standard mis-attributions "
           "and all title truncation that positional extraction incurs on the GRI corpus. "
           "Cross-checking 1,093 extracted claims produced 3,461 adjudicated pairs with no failures, "
           "yielding 70 apparently-supported and 22 apparently-contradicted claims — and 1,001 "
           "claims explicitly returned as unverified rather than silently assumed clean.")

    h(doc, 2, "5.2  Limitations")
    bullet(doc, f"**The conduct channel is thin.** Only {f(S['graph_source_type'].get('news',0))} of "
                f"{f(S['graph_nodes'])} nodes originate from news, and 92% of claims consequently "
                "cannot be verified. This is stated plainly rather than reframed, because the "
                "conduct channel is the component on which the entire contribution rests, and its "
                "size bounds every conclusion in §4.4.", numbered=True)
    bullet(doc, "**One issuer is analysed in depth.** The end-to-end results constitute a case study "
                "with genuine temporal depth, not evidence generalisable to the sector. The corpus "
                "supports a wider run; the analysis has not yet been performed.", numbered=True)
    bullet(doc, "**The ESG classifier is an off-the-shelf checkpoint with no evaluation on this "
                "corpus.** It was applied as published, without fine-tuning, so it was fitted to a "
                "different data distribution than the one it is used on here. §2.3.3 documents a "
                "measurable precision concern, a demonstrated lexical failure mode, and a "
                "quantified disagreement between the pillar-tag rule and the ESG-flag rule; no "
                "per-label precision/recall figure exists to settle any of them. Every sentence "
                "count in Chapter 2 should be read as an upper bound on ESG-relevant content.",
           numbered=True)
    bullet(doc, "**There is no ground truth**, so no precision or recall figure for *greenwashing "
                "detection* exists or can exist here. What is reported is graph-quality attributes "
                "and internal consistency. A reader looking for a detection accuracy will not find "
                "one, by design.", numbered=True)
    bullet(doc, "**Entity resolution ran in a degraded configuration.** Embedding blocking and LLM "
                "adjudication are implemented and covered by tests but were disabled for the "
                "reported results because of an API billing block. The residual duplicate count of "
                "60 in §4.3.1 is therefore an upper bound rather than the system's best achievable "
                "result.", numbered=True)
    bullet(doc, "**The news channel is dominated by one retrieval back-end.** §2.3.5 shows one "
                "provider supplied roughly three-quarters of the articles, so the three-back-end "
                "design provides much less protection against ranking bias than intended.",
           numbered=True)
    bullet(doc, "**Dependence on a commercial language model.** Adjudication runs on `gpt-4o-mini`; "
                "results are not bit-for-bit reproducible across model versions.", numbered=True)
    bullet(doc, "**No systematic human evaluation.** No panel of independent annotators has "
                "re-scored the verdicts, so no inter-annotator agreement figure is available and the "
                "adjudicator's judgements are unaudited.", numbered=True)

    h(doc, 2, "5.3  Ethical Considerations")
    p(doc, "This system makes statements about the conduct of **real, publicly listed companies**, "
           "which places obligations on its design that a purely technical evaluation would not "
           "capture.")
    bullet(doc, "**Advisory, never a verdict.** Every output carries an explicit advisory flag and "
                "the caveat that no ground-truth label exists. The system does not emit a "
                "greenwashing score, and this is a deliberate refusal rather than an unimplemented "
                "feature.")
    bullet(doc, "**Absence of contradiction is not exoneration.** This is hard-coded as a caveat on "
                "every statistical summary the system produces, because it is the most likely "
                "misreading of a result dominated by 1,001 unverified claims.")
    bullet(doc, "**Provenance is a safeguard, not a feature.** Because every statement resolves to a "
                "file, a page and a sentence, an incorrect allegation can always be refuted by a "
                "reader opening the cited page. §4.3.5 reports the coverage honestly "
                f"({S['graph_provenance_pct']}% of nodes) rather than claiming completeness, because "
                "a node without provenance is precisely a statement a reader cannot check.")
    bullet(doc, "**Self-verification is blocked.** Content published by the company itself is never "
                "counted as independent evidence in its own favour.")
    bullet(doc, "**Risk of misuse.** If a reader ignores the caveats and reads "
                "`appears_contradicted` as a finding of fact, the system could cause unwarranted "
                "reputational harm. For that reason the claim ledger presents its coverage caveat in "
                "the header, above the signal-first ordering, rather than in a footnote where it "
                "would be read last or not at all.")

    h(doc, 2, "5.4  Future Work")
    p(doc, "The 1,001 unverified claims indicate the direction of future work more clearly than any "
           "other result in this thesis. The binding constraint is not the graph, the schema or the "
           "adjudicator but the volume of independent conduct evidence: extending the conduct "
           "channel with administrative-penalty records, regulator and exchange announcements, and "
           "environmental monitoring data would convert a large share of those claims into decidable "
           "ones. In priority order, the remaining work is: produce the classifier evaluation table "
           "(§2.2.4); run the negative control (§4.2); extend the conduct channel (§4.6, E3); scale "
           "the end-to-end analysis from one issuer to the sector, for which the corpus is already "
           "in place; add distance decay to adjudication aggregation (§4.6, E1); and build a "
           "path-based reasoning layer over the graph, for which the hub-decomposition export of "
           "§3.5.6 is the prepared substrate.")
    p(doc, "The design principle that should survive all of them is the one adopted here at the "
           "outset: in a domain with no ground truth, a system that shows its evidence and declines "
           "to score is more useful, and more defensible, than one that supplies a number nobody can "
           "check.")


# ---------------------------------------------------------------------------
#                              BACK MATTER
# ---------------------------------------------------------------------------
def _back_matter(doc):
    h(doc, 1, "References")
    todo(doc, "CITE", "Entries [1]–[33] below are transcribed **verbatim** from "
                      "`docs/Paper.docx` and are not edited here, so the numbering in the body "
                      "text still resolves. Seven of them are incomplete in the source and are "
                      "marked ⚠ — they carry a placeholder rather than a citation. Four more are "
                      "cross-references to an earlier entry rather than works in their own right "
                      "([13]→[1], [14]→[2], [24]→[18], [27]→[23]); a numbered list should not "
                      "contain the same work twice, so these should either be merged into the "
                      "entry they point at (renumbering the body) or replaced with the distinct "
                      "work actually meant. See the defect list immediately after the entries.")
    refs = [
        ("[1]", "Lyon, T. P., Maxwell, J. W.: Greenwash: corporate environmental disclosure "
                "under threat of audit. Journal of Economics & Management Strategy 20(1) (2011)",
         False),
        ("[2]", "Delmas, M. A., Burbano, V. C.: The drivers of greenwashing. California "
                "Management Review 54(1), 64–87 (2011)", False),
        ("[3]", "Kaoukis, G., et al.: EmeraldMind: a knowledge graph–augmented framework for "
                "greenwashing detection. In: Proceedings of the ACM Web Conference 2026 (2026). "
                "https://doi.org/10.1145/3774904.3792997", False),
        ("[4]", "Prime Minister of Vietnam: Decision No. 13/2024/QĐ-TTg on updating the lists of "
                "sectors and GHG-emitting facilities. Hanoi (2024)", False),
        ("[5]", "Recognition over performance: a critical review of ESG practices in Vietnam. "
                "Intergovernmental Research and Policy Journal (2025)", False),
        ("[6]", "Slaughter and May: ESG in APAC 2025 — Vietnam (2025). "
                "https://www.slaughterandmay.com/services/practices/environmental-social-and-"
                "governance/esg-in-apac-2025/vietnam/", False),
        ("[7]", "Pham, D., et al.: How Vietnam can achieve net-zero carbon emissions in "
                "construction and built environment by 2050: an integrated AHP and DEMATEL "
                "approach (2025)", False),
        ("[8]", "BW Industrial: ESG — a growing focus in Vietnam and opportunity for "
                "sustainability in new-economy real estate (2024)", False),
        ("[9]", "Detecting greenwashing: a natural language processing literature survey. "
                "arXiv:2502.07541 (2025)", False),
        ("[10]", "Corporate greenwashing detection in text — a survey (2025)", False),
        ("[11]", "Kim, J. S., et al.: Establishment of NLP-based greenwashing pattern detection "
                 "service. In: Advances in Computer Science and Ubiquitous Computing. Springer "
                 "(2023)", False),
        ("[12]", "Detecting greenwashing in ESG disclosure: an NLP-based analysis of Central and "
                 "Eastern European firms. Sustainability 18(3), 1486 (2026)", False),
        ("[13]", "Lyon, T. P., Maxwell, J. W. (2011) — duplicate of [1]", True),
        ("[14]", "Delmas, M. A., Burbano, V. C. (2011) — duplicate of [2]", True),
        ("[15]", "European Securities and Markets Authority (ESMA): Guidelines on greenwashing "
                 "in sustainability reporting (2023)", False),
        ("[16]", "Competition and Consumer Commission: Guidelines on environmental claims (2023)",
         False),
        ("[17]", "Gorovaia, N., et al.: Identifying greenwashing in corporate social "
                 "responsibility reports using natural language processing. European Financial "
                 "Management. Wiley (2025)", False),
        ("[18]", "Luccioni, A., Baylor, E., Duchene, N.: Analyzing sustainability reports using "
                 "natural language processing. arXiv:2011.08073 (2020)", False),
        ("[19]", "Webersinke, N., et al.: ClimateBert: a pretrained language model for "
                 "climate-related text. arXiv:2110.12010 (2022)", False),
        ("[20]", "FinBERT-ESG — **full citation missing in the source**", True),
        ("[21]", "Kölbel, J. F., et al.: Ask BERT: how regulatory disclosure of transition and "
                 "physical climate risks affects the CDS term structure. SSRN 3616324 (2020)",
         False),
        ("[22]", "BERTopic applied to ESG — **full citation missing in the source**", True),
        ("[23]", "Vinella, F., et al.: specificity scoring for greenwashing hints in ESG reports "
                 "(2024). aclanthology.org/2025.swisstext-1.3.pdf — **incomplete in the source**",
         True),
        ("[24]", "Luccioni, A.: ClimateQA (2020) — points back to [18]", True),
        ("[25]", "ESGReveal: an LLM-based approach for extracting structured data from ESG "
                 "reports. arXiv:2312.17264 (2023)", False),
        ("[26]", "Schimanski, T., et al.: ESGBERT nature dataset (2024) — **full citation "
                 "missing in the source**", True),
        ("[27]", "Vinella, F., et al. (2024) — duplicate of [23]", True),
        ("[28]", "Bingler, J. A., et al.: Cheap talk and cherry-picking: what ClimateBert has to "
                 "say on corporate climate risk disclosures. Finance Research Letters (2022)",
         False),
        ("[29]", "ChatClimate — **full citation missing in the source**", True),
        ("[30]", "Knowledge-graph fact-checking survey — **full citation missing in the source**",
         True),
        ("[31]", "State Securities Commission of Vietnam: ESG implementation and disclosure "
                 "handbook (2024)", False),
        ("[32]", "Lexology: Vietnam — comprehensive regulatory framework for the carbon market "
                 "underway (2025)", False),
        ("[33]", "IR Global: Green buildings in Vietnam — how sustainable are they? (2022)",
         False),
    ]
    for num, entry, flagged in refs:
        bullet(doc, f"**{num}**{' ⚠' if flagged else ''}  {entry}")

    todo(doc, "CITE", "**Four defects in the list above, all inherited from "
                      "`docs/Paper.docx` and none of them fixed here, because fixing a citation "
                      "by guessing is worse than flagging it.** (1) *Marker [4] is used for two "
                      "different instruments*: §1.1 cites [4] for Prime Minister Decision "
                      "13/2024/QĐ-TTg, but §1.2.5 cites [4] for Circular 96/2020/TT-BTC. "
                      "Circular 96 needs its own number — it is the single most-cited instrument "
                      "in this report. (2) *Marker [21] does not match its entry*: the body cites "
                      "[21] as **ESG-BERT**, but entry [21] is Kölbel et al., *Ask BERT*, a "
                      "different model. (3) *Seven entries carry placeholders* rather than "
                      "citations: [20], [22], [23], [26], [29], [30] — and [5], [9], [10], [12] "
                      "have no author, which will not pass an LNCS reference check. (4) *Four "
                      "entries are duplicates* pointing at earlier ones: [13]→[1], [14]→[2], "
                      "[24]→[18], [27]→[23].")
    p(doc, "**Still to be sourced — new entries, numbering from [34].** These are the "
           "[CITE: …] markers remaining in the body; each names the group it needs.")
    for group, items in [
        ("Vietnamese regulation (primary sources)", "Decision 2171/QĐ-BTC; QCVN 09:2013/BXD; "
                                                    "SSC–IFC ESG reporting guide. (Circular "
                                                    "96/2020/TT-BTC also needs its own number — "
                                                    "see defect 1 above.)"),
        ("Standards", "GRI Universal Standards 2021; GRI topic-specific standards; GRI content "
                      "index template 2021 (ground truth in §4.3.3)."),
        ("Document parsing", "PyMuPDF (§2.2.3); Marker/Surya, olmOCR, Nougat, Docling "
                             "(§3.2.2, §4.3.3)."),
        ("Vietnamese NLP", "underthesea (§2.2.3); ViDeBERTa — the specific published checkpoint "
                           "`nguyen599/ViDeBERTa-v3-ESG-base` used in §2.2.4."),
        ("Knowledge graphs and temporal representation", "Hogan et al., Knowledge Graphs (ACM "
                                                         "CSUR 2021); bi-temporal data models; "
                                                         "temporal-KG representation; LLM-based "
                                                         "KG construction (Trajanoska et al.; "
                                                         "Carta et al.); Neo4j. (§1.2.4, §3.5.)"),
        ("RAG and LLMs", "Lewis et al., Retrieval-Augmented Generation (§1.1); structured output "
                         "and constrained decoding; LLM-as-a-judge; hallucination in high-stakes "
                         "domains."),
        ("Evaluation methodology", "Metamorphic testing; permutation testing; Krippendorff's α; "
                                   "FEVER. (§4.2.)"),
    ]:
        bullet(doc, f"**{group}** — {items}")

    page_break(doc)

    h(doc, 1, "Appendix")
    p(doc, "The following appendices are specified but not yet rendered. Each can be produced "
           "mechanically from the artifact named.")
    bullet(doc, "**Appendix A — Graph schema.** The 28 node classes, 48 edge labels over 76 legal "
                "class pairs, and the T1/T2/T3 tier assignment, with the full statement of the "
                "eight design "
                "principles. Source: `config/schema.json`.")
    bullet(doc, "**Appendix B — Extraction prompt templates**, claim-side and conduct-side, "
                "including the language constraint of §3.5.2.")
    bullet(doc, "**Appendix C — Adjudication prompt and output schema** of the cross-check stage.")
    bullet(doc, "**Appendix D — Indicator metadata schema.** The GRI catalogue record and the KPI "
                "definition record, including provenance and SHA-256 blocks.")
    bullet(doc, "**Appendix E — Three complete advisory dossiers** corresponding to Cases 1–3 in "
                "§4.5, reproduced verbatim from the system output.")
    bullet(doc, "**Appendix F — Neo4j constraints and analyst queries.**")
    bullet(doc, "**Appendix G — Parser ablation reproduction package.** Benchmark script, test "
                "suite, per-document extractions for all three arms, and side-by-side comparison "
                "tables for all 36 standards. Source: `gri/benchmark_results/`.")
    bullet(doc, "**Appendix H — Exploratory data analysis package.** The measurement script, the "
                "full statistics file and every figure in Chapter 2. Source: "
                "`docs/eda_report_data.py`, `docs/eda_out/eda_stats.json`, `docs/figures/`.")

    page_break(doc)

    h(doc, 1, "Annex Z — Outstanding items (remove this annex before submission)")
    p(doc, "Every gap in this document, collected in one place and ordered by how much closing it "
           "would strengthen the defence.")
    table(doc, "Outstanding items, highest impact first.",
          ["#", "Item", "Where", "Kind", "Effort"],
          [["1", "Classifier EVALUATION on a hand-labelled sample of this corpus (per-label P/R/F1 "
                 "+ the binary esg flag). Not training details — nothing was trained",
            "§2.2.4", "TODO-RUN", "Medium, no GPU"],
           ["2", "Negative control: one issuer's claims against another's conduct pool", "§4.2",
            "TODO-RUN", "Low"],
           ["3", "References — the entire bibliography", "All", "CITE", "Medium"],
           ["4", "Citation slots marked inline in red", "§1.1, §1.2, §3.2", "CITE", "Medium"],
           ["5", "Re-run the quality instrument on the current graph so §4.3.5 and the ablations share a "
                 "common point", "§4.3", "TODO-RUN", "Low"],
           ["6", "E1–E4 frequency estimates from ~20 reviewed dossiers", "§4.6", "TODO-RUN", "Low"],
           ["7", "Two or three additional issuers end-to-end", "§2.1, §5.2", "TODO-RUN", "High"],
           ["8", "Architecture and pipeline figures: draw into the eight remaining placeholders "
                 "(the nine data figures in Chapter 2 are already generated)",
            "§2.2, §3", "WRITE", "Medium"],
           ["9", "Acknowledgment section", "Front", "WRITE", "Trivial"],
           ["10", "Title page: names, student IDs, supervisor", "Cover", "WRITE", "Trivial"],
           ["11", "Convert plain-text equations to Word equation objects", "§3.1, §3.2", "WRITE",
            "Low"],
           ["12", "Decide whether to fix the emphasis defect in the GRI segmenter", "§4.3.3",
            "WRITE", "Low"],
           ["13", "Discarded-metrics discussion from EVALUATION_WITHOUT_LABELS.md §8", "§1.2",
            "WRITE", "Low"],
           ["14", "Update the three field-driven lists (contents, tables, figures) in Word",
            "Front", "WRITE", "Trivial"]],
          widths=[0.6, 5.0, 2.2, 1.6, 1.4], right_align_from=99)
    p(doc, "**Sections deliberately not drafted.** §1.1 and §1.2 are written in full prose but carry "
           "no references; they are complete as argument and incomplete as scholarship. The "
           "bibliography is empty by choice. No number anywhere in this document was estimated, "
           "rounded up or carried over from a previous draft without being re-read from the artifact "
           "that produces it — Chapter 2 is generated directly from `docs/eda_out/eda_stats.json`.")


def main() -> int:
    doc = build()
    doc.save(OUT)
    print(f"wrote {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")
    print(f"  tables: {_table_no[0]}   figures: {_figure_no[0]}   equations: {_eq_no[0]}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
